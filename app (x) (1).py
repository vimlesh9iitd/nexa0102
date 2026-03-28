"""
NEXA — Unified AI Platform
Home  →  INFOCHAT (model3)  |  IMAGINE (Text-to-Image model1)
Single-file, session-based navigation, zero broken features.
Run:  streamlit run app.py
"""

import streamlit as st
import requests
import base64
import re
import math
import os

# ─────────────────────────────────────────────────────────────────────────────
#  PAGE CONFIG  — called ONCE at the very top
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="NEXA — AI Platform",
    layout="wide",
    initial_sidebar_state="collapsed",
    page_icon="⚡",
)

# ─────────────────────────────────────────────────────────────────────────────
#  NAV STATE
# ─────────────────────────────────────────────────────────────────────────────
if "page" not in st.session_state:
    st.session_state.page = "home"

# ─────────────────────────────────────────────────────────────────────────────
#  SHARED LOGO HELPER
# ─────────────────────────────────────────────────────────────────────────────
def _get_logo_b64() -> str:
    """Return base64 string of logo/l8.jpeg if it exists, else ''."""
    for p in ["logo/l9.jpeg", "logo/l9.jpg", "logo/l8.jpeg", "logo/l8.jpg", "l8.jpeg"]:
        if os.path.exists(p):
            with open(p, "rb") as f:
                return base64.b64encode(f.read()).decode()
    return ""


# ═════════════════════════════════════════════════════════════════════════════
#
#   H O M E   P A G E
#
# ═════════════════════════════════════════════════════════════════════════════
def render_home():
    logo_b64 = _get_logo_b64()

    # ── CSS ──────────────────────────────────────────────────────────────────
    st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@700;900&family=Syne:wght@700;800&family=DM+Sans:wght@400;500;600&display=swap');

:root{
  --bg:#05030f; --accent:#00e87a; --accent2:#00c96a;
  --purple:#bf5fff; --cyan:#00d4ff; --pink:#ff2d9b; --gold:#ffd700;
  --text:#e8d5ff; --muted:rgba(232,213,255,.55);
  --border-p:rgba(191,95,255,.22); --border-g:rgba(0,232,122,.22);
  --card-bg:rgba(13,6,30,.92);
}
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0;}
html,body,[data-testid="stAppViewContainer"],[data-testid="stApp"]{
  background:var(--bg)!important; color:var(--text);
  font-family:'DM Sans',sans-serif; -webkit-font-smoothing:antialiased;
}
[data-testid="stAppViewContainer"]>.main{background:var(--bg)!important;}
.block-container{padding:0!important;max-width:100%!important;}
#MainMenu,footer,header{visibility:hidden;}

/* ambient */
[data-testid="stAppViewContainer"]::before{
  content:'';position:fixed;inset:0;pointer-events:none;z-index:0;
  background:
    radial-gradient(ellipse 80% 55% at 15% 8%,  rgba(191,95,255,.14) 0%,transparent 60%),
    radial-gradient(ellipse 65% 45% at 85% 85%,  rgba(0,212,255,.09)  0%,transparent 55%),
    radial-gradient(ellipse 50% 40% at 50% 50%,  rgba(0,232,122,.05)  0%,transparent 65%);
}

/* ── NAV ── */
.hm-nav{
  position:relative;z-index:10;
  display:flex;align-items:center;justify-content:space-between;
  padding:16px 52px;
  border-bottom:1px solid var(--border-p);
  background:rgba(5,3,15,.86);
  backdrop-filter:blur(20px);
}
.hm-logo{display:flex;align-items:center;gap:13px;}
.hm-logo-icon{
  width:44px;height:44px;background:var(--accent);
  border-radius:12px;display:flex;align-items:center;
  justify-content:center;font-size:22px;flex-shrink:0;
  box-shadow:0 0 20px rgba(0,232,122,.4);
}
.hm-logo-img{
  width:44px;height:44px;border-radius:12px;object-fit:cover;
  box-shadow:0 0 20px rgba(0,232,122,.45);flex-shrink:0;
}
@keyframes hm-glow{
  0%,100%{text-shadow:0 0 8px rgba(0,232,122,.55),0 0 18px rgba(0,232,122,.3);}
  50%    {text-shadow:0 0 16px rgba(0,232,122,.9),0 0 32px rgba(0,232,122,.5);}
}
.hm-logo-name{
  font-family:'Orbitron',monospace;font-size:21px;font-weight:900;letter-spacing:4px;
  background:linear-gradient(135deg,#00e87a 0%,#00ffcc 50%,#00c96a 100%);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;
  animation:hm-glow 2.5s ease-in-out infinite;
}
.hm-logo-tag{font-size:10.5px;letter-spacing:3px;color:rgba(0,232,122,.5);font-family:'DM Sans',sans-serif;}

/* ── HERO ── */
@keyframes hm-fadeup{
  0%{opacity:0;transform:translateY(26px);filter:blur(8px);}
  100%{opacity:1;transform:translateY(0);filter:blur(0);}
}
@keyframes hm-gradflow{
  0%{background-position:0% 50%;}50%{background-position:100% 50%;}100%{background-position:0% 50%;}
}
.hm-title{
  font-family:'Orbitron',monospace;
  font-size:clamp(44px,8vw,96px);font-weight:900;letter-spacing:12px;
  background:linear-gradient(135deg,#bf5fff 0%,#00d4ff 33%,#00e87a 66%,#ffd700 100%);
  background-size:300% 300%;
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;
  animation:hm-fadeup .8s ease-out both, hm-gradflow 7s ease infinite;
  margin-bottom:18px;
}
.hm-sub{
  font-size:clamp(14px,1.6vw,19px);color:var(--muted);
  font-family:'DM Sans',sans-serif;letter-spacing:.4px;line-height:1.7;
  animation:hm-fadeup .8s .2s ease-out both;max-width:560px;margin-bottom:22px;
}
.hm-pill{
  display:inline-flex;align-items:center;gap:8px;
  padding:7px 22px;border-radius:40px;
  background:rgba(0,232,122,.07);border:1px solid rgba(0,232,122,.28);
  font-size:11.5px;letter-spacing:2.5px;color:rgba(0,232,122,.82);
  font-family:'Orbitron',monospace;
  animation:hm-fadeup .8s .4s ease-out both;
}

/* ── DIVIDER ── */
.hm-divider{
  height:1px;max-width:880px;margin:0 auto 60px;
  background:linear-gradient(90deg,transparent,rgba(191,95,255,.4),rgba(0,212,255,.3),transparent);
  border:none;position:relative;z-index:5;
}

/* ── CARDS GRID ── */
.hm-grid{
  position:relative;z-index:5;
  display:grid;grid-template-columns:repeat(2,1fr);gap:32px;
  max-width:860px;margin:0 auto;padding:0 32px 90px;
}
@media(max-width:700px){.hm-grid{grid-template-columns:1fr;}}

/* ── CARD ── */
@keyframes hm-cardin{0%{opacity:0;transform:translateY(36px);}100%{opacity:1;transform:translateY(0);}}
.hm-card{
  background:var(--card-bg);border-radius:24px;
  padding:38px 32px 32px;position:relative;overflow:hidden;
  cursor:pointer;
  transition:transform .28s cubic-bezier(.34,1.56,.64,1),
             box-shadow .28s ease,border-color .28s ease;
  animation:hm-cardin .6s ease-out both;
}
.hm-card::before{
  content:'';position:absolute;top:0;left:0;right:0;height:3px;
  background:var(--c-grad);
}

/* purple card */
.card-purple{border:1px solid rgba(191,95,255,.28);
  --c-grad:linear-gradient(90deg,#bf5fff,#8b5cf6,#00d4ff);}
.card-purple:hover{border-color:rgba(191,95,255,.75);transform:translateY(-9px);
  box-shadow:0 28px 70px rgba(191,95,255,.2),0 0 0 1px rgba(191,95,255,.16);}

/* green card */
.card-green{border:1px solid rgba(0,232,122,.28);
  --c-grad:linear-gradient(90deg,#00e87a,#00ffcc,#00d4ff);}
.card-green:hover{border-color:rgba(0,232,122,.75);transform:translateY(-9px);
  box-shadow:0 28px 70px rgba(0,232,122,.18),0 0 0 1px rgba(0,232,122,.14);}

/* card internals */
.hm-icon{font-size:44px;display:block;margin-bottom:18px;}
.hm-badge{
  display:inline-flex;align-items:center;gap:5px;
  padding:4px 13px;border-radius:20px;
  font-size:9.5px;letter-spacing:2px;font-family:'Orbitron',monospace;margin-bottom:14px;
}
.badge-p{background:rgba(191,95,255,.12);border:1px solid rgba(191,95,255,.35);color:#bf5fff;}
.badge-g{background:rgba(0,232,122,.10);border:1px solid rgba(0,232,122,.32);color:#00e87a;}
.hm-ctitle{
  font-family:'Orbitron',monospace;font-size:18px;font-weight:900;
  letter-spacing:2px;margin-bottom:11px;
}
.ctitle-p{color:#bf5fff;}
.ctitle-g{color:#00e87a;}
.hm-cdesc{
  font-family:'DM Sans',sans-serif;font-size:14px;color:var(--muted);
  line-height:1.65;margin-bottom:26px;
}
.hm-feats{display:flex;flex-direction:column;gap:8px;}
.hm-feat{display:flex;align-items:center;gap:10px;font-family:'DM Sans',sans-serif;font-size:13px;color:rgba(232,213,255,.72);}
.dot{width:6px;height:6px;border-radius:50%;flex-shrink:0;}
.dot-p{background:#bf5fff;box-shadow:0 0 7px rgba(191,95,255,.65);}
.dot-g{background:#00e87a;box-shadow:0 0 7px rgba(0,232,122,.65);}
.dot-y{background:#ffd700;box-shadow:0 0 7px rgba(255,215,0,.65);}
.dot-c{background:#00ff88;box-shadow:0 0 7px rgba(0,255,136,.65);}

/* teal/green-medical card — Dr. NEXA */
.card-teal{border:1px solid rgba(0,255,136,.28);
  --c-grad:linear-gradient(90deg,#00ff88,#00d4ff,#bf5fff);}
.card-teal:hover{border-color:rgba(0,255,136,.75);transform:translateY(-9px);
  box-shadow:0 28px 70px rgba(0,255,136,.18),0 0 0 1px rgba(0,255,136,.14);}
.ctitle-c{color:#00ff88;}
.badge-c{background:rgba(0,255,136,.10);border:1px solid rgba(0,255,136,.32);color:#00ff88;}

/* gold card */
.card-gold{border:1px solid rgba(255,215,0,.28);
  --c-grad:linear-gradient(90deg,#ffd700,#ff9500,#ffcc00);}
.card-gold:hover{border-color:rgba(255,215,0,.75);transform:translateY(-9px);
  box-shadow:0 28px 70px rgba(255,215,0,.18),0 0 0 1px rgba(255,215,0,.14);}
.ctitle-y{color:#ffd700;}
.badge-y{background:rgba(255,215,0,.10);border:1px solid rgba(255,215,0,.32);color:#ffd700;}

/* pink card — Your's NEXA */
.card-pink{border:1px solid rgba(255,107,157,.28);
  --c-grad:linear-gradient(90deg,#ff6b9d,#c44569,#9b59b6);}
.card-pink:hover{border-color:rgba(255,107,157,.75);transform:translateY(-9px);
  box-shadow:0 28px 70px rgba(255,107,157,.22),0 0 0 1px rgba(255,107,157,.16);}
.ctitle-pk{color:#ff6b9d;}
.badge-pk{background:rgba(255,107,157,.10);border:1px solid rgba(255,107,157,.32);color:#ff6b9d;}
.dot-pk{background:#ff6b9d;box-shadow:0 0 7px rgba(255,107,157,.65);}

/* orange card — India Tour */
.card-orange{border:1px solid rgba(255,140,0,.28);
  --c-grad:linear-gradient(90deg,#ff8c00,#ff4500,#ffd700);}
.card-orange:hover{border-color:rgba(255,140,0,.75);transform:translateY(-9px);
  box-shadow:0 28px 70px rgba(255,140,0,.22),0 0 0 1px rgba(255,140,0,.16);}
.ctitle-or{color:#ff8c00;}
.badge-or{background:rgba(255,140,0,.10);border:1px solid rgba(255,140,0,.32);color:#ff8c00;}
.dot-or{background:#ff8c00;box-shadow:0 0 7px rgba(255,140,0,.65);}

/* ── HERO: full-width centered ── */
.hm-hero{
  position:relative;z-index:5;
  display:flex;flex-direction:column;align-items:center;
  padding:72px 24px 36px;text-align:center;
}

/* ── BUTTON ROW WRAPPER ── */
.hm-btn-row-wrap{
  position:relative;z-index:10;
  max-width:1100px;margin:0 auto 64px;padding:0 28px;
}

/* ── KEYFRAMES ── */
@keyframes nx-float{
  0%,100%{transform:translateY(0);}
  50%{transform:translateY(-6px);}
}
@keyframes nx-glow-p{
  0%,100%{box-shadow:0 0 14px rgba(191,95,255,.4),0 0 32px rgba(191,95,255,.15),inset 0 1px 0 rgba(255,255,255,.07);}
  50%{box-shadow:0 0 26px rgba(191,95,255,.75),0 0 60px rgba(191,95,255,.28),inset 0 1px 0 rgba(255,255,255,.13);}
}
@keyframes nx-glow-g{
  0%,100%{box-shadow:0 0 14px rgba(0,232,122,.4),0 0 32px rgba(0,232,122,.15),inset 0 1px 0 rgba(255,255,255,.07);}
  50%{box-shadow:0 0 26px rgba(0,232,122,.75),0 0 60px rgba(0,232,122,.28),inset 0 1px 0 rgba(255,255,255,.13);}
}
@keyframes nx-glow-y{
  0%,100%{box-shadow:0 0 14px rgba(255,215,0,.4),0 0 32px rgba(255,215,0,.15),inset 0 1px 0 rgba(255,255,255,.07);}
  50%{box-shadow:0 0 26px rgba(255,215,0,.75),0 0 60px rgba(255,215,0,.28),inset 0 1px 0 rgba(255,255,255,.13);}
}
@keyframes nx-glow-c{
  0%,100%{box-shadow:0 0 14px rgba(0,255,136,.4),0 0 32px rgba(0,255,136,.15),inset 0 1px 0 rgba(255,255,255,.07);}
  50%{box-shadow:0 0 26px rgba(0,255,136,.75),0 0 60px rgba(0,255,136,.28),inset 0 1px 0 rgba(255,255,255,.13);}
}
@keyframes nx-glow-pk{
  0%,100%{box-shadow:0 0 14px rgba(255,107,157,.4),0 0 32px rgba(255,107,157,.15),inset 0 1px 0 rgba(255,255,255,.07);}
  50%{box-shadow:0 0 26px rgba(255,107,157,.75),0 0 60px rgba(255,107,157,.28),inset 0 1px 0 rgba(255,255,255,.13);}
}
@keyframes nx-glow-or{
  0%,100%{box-shadow:0 0 14px rgba(255,140,0,.4),0 0 32px rgba(255,140,0,.15),inset 0 1px 0 rgba(255,255,255,.07);}
  50%{box-shadow:0 0 26px rgba(255,140,0,.75),0 0 60px rgba(255,140,0,.28),inset 0 1px 0 rgba(255,255,255,.13);}
}
@keyframes nx-shine{
  0%{left:-100%;}
  100%{left:200%;}
}

/* ── BASE BUTTON — all 6 identical size ── */
div[data-testid="column"] .stButton>button{
  position:relative!important;inset:auto!important;
  width:100%!important;
  min-height:96px!important;
  height:96px!important;
  border-radius:18px!important;
  cursor:pointer!important;
  font-family:'Orbitron',monospace!important;
  font-size:11px!important;font-weight:800!important;
  letter-spacing:1.8px!important;text-transform:uppercase!important;
  padding:10px 8px!important;
  display:flex!important;align-items:center!important;justify-content:center!important;
  flex-direction:column!important;
  gap:6px!important;
  opacity:1!important;z-index:20!important;
  backdrop-filter:blur(18px)!important;
  overflow:hidden!important;
  white-space:pre-line!important;
  text-align:center!important;
  line-height:1.35!important;
  transition:transform .22s cubic-bezier(.34,1.56,.64,1),filter .22s ease!important;
}
/* shimmer sweep pseudo-element via ::after on the wrapper — use background trick */
div[data-testid="column"] .stButton>button::after{
  content:''!important;
  position:absolute!important;
  top:0!important;left:-100%!important;
  width:60%!important;height:100%!important;
  background:linear-gradient(105deg,transparent 30%,rgba(255,255,255,.12) 50%,transparent 70%)!important;
  animation:nx-shine 3.5s ease-in-out infinite!important;
  pointer-events:none!important;
}
div[data-testid="column"] .stButton>button:hover{
  transform:translateY(-7px) scale(1.04)!important;
  filter:brightness(1.22)!important;
}

/* ── COL 1 — INFOCHAT purple ── */
div[data-testid="column"]:nth-child(1) .stButton>button{
  background:linear-gradient(145deg,rgba(191,95,255,.18) 0%,rgba(139,92,246,.10) 100%)!important;
  border:1.5px solid rgba(191,95,255,.70)!important;
  color:#d8a8ff!important;
  animation:nx-float 3.6s ease-in-out infinite, nx-glow-p 3.6s ease-in-out infinite!important;
  animation-delay:0s,0s!important;
}
div[data-testid="column"]:nth-child(1) .stButton>button::after{animation-delay:0s!important;}

/* ── COL 2 — IMAGINE green ── */
div[data-testid="column"]:nth-child(2) .stButton>button{
  background:linear-gradient(145deg,rgba(0,232,122,.16) 0%,rgba(0,212,255,.10) 100%)!important;
  border:1.5px solid rgba(0,232,122,.70)!important;
  color:#00ffaa!important;
  animation:nx-float 3.6s ease-in-out infinite, nx-glow-g 3.6s ease-in-out infinite!important;
  animation-delay:.42s,.42s!important;
}
div[data-testid="column"]:nth-child(2) .stButton>button::after{animation-delay:.6s!important;}

/* ── COL 3 — STUDY PLANNER gold ── */
div[data-testid="column"]:nth-child(3) .stButton>button{
  background:linear-gradient(145deg,rgba(255,215,0,.15) 0%,rgba(255,149,0,.10) 100%)!important;
  border:1.5px solid rgba(255,215,0,.70)!important;
  color:#ffe566!important;
  animation:nx-float 3.6s ease-in-out infinite, nx-glow-y 3.6s ease-in-out infinite!important;
  animation-delay:.84s,.84s!important;
}
div[data-testid="column"]:nth-child(3) .stButton>button::after{animation-delay:1.2s!important;}

/* ── COL 4 — DR. NEXA teal ── */
div[data-testid="column"]:nth-child(4) .stButton>button{
  background:linear-gradient(145deg,rgba(0,255,136,.15) 0%,rgba(0,212,255,.10) 100%)!important;
  border:1.5px solid rgba(0,255,136,.70)!important;
  color:#66ffbb!important;
  animation:nx-float 3.6s ease-in-out infinite, nx-glow-c 3.6s ease-in-out infinite!important;
  animation-delay:1.26s,1.26s!important;
}
div[data-testid="column"]:nth-child(4) .stButton>button::after{animation-delay:1.8s!important;}

/* ── COL 5 — YOUR'S NEXA pink ── */
div[data-testid="column"]:nth-child(5) .stButton>button{
  background:linear-gradient(145deg,rgba(255,107,157,.16) 0%,rgba(196,69,105,.10) 100%)!important;
  border:1.5px solid rgba(255,107,157,.70)!important;
  color:#ffaacc!important;
  animation:nx-float 3.6s ease-in-out infinite, nx-glow-pk 3.6s ease-in-out infinite!important;
  animation-delay:1.68s,1.68s!important;
}
div[data-testid="column"]:nth-child(5) .stButton>button::after{animation-delay:2.4s!important;}

/* ── COL 6 — INDIA TOUR orange ── */
div[data-testid="column"]:nth-child(6) .stButton>button{
  background:linear-gradient(145deg,rgba(255,140,0,.15) 0%,rgba(255,69,0,.10) 100%)!important;
  border:1.5px solid rgba(255,140,0,.70)!important;
  color:#ffcc66!important;
  animation:nx-float 3.6s ease-in-out infinite, nx-glow-or 3.6s ease-in-out infinite!important;
  animation-delay:2.1s,2.1s!important;
}
div[data-testid="column"]:nth-child(6) .stButton>button::after{animation-delay:3s!important;}

/* ── FOOTER ── */
.hm-footer{
  text-align:center;padding:22px;
  font-size:10.5px;letter-spacing:2.5px;
  color:rgba(232,213,255,.2);font-family:'DM Sans',sans-serif;
  position:relative;z-index:5;
}



::-webkit-scrollbar{width:4px;}
::-webkit-scrollbar-track{background:var(--bg);}
::-webkit-scrollbar-thumb{background:rgba(191,95,255,.35);border-radius:2px;}
</style>
""", unsafe_allow_html=True)

    # ── NAV — logo + NEXA text ────────────────────────────────────────────────
    logo_tag = (
        f'<img class="hm-logo-img" src="data:image/jpeg;base64,{logo_b64}" '
        f'style="width:44px;height:44px;border-radius:12px;object-fit:cover;'
        f'box-shadow:0 0 20px rgba(0,232,122,.45);flex-shrink:0;"/>'
        if logo_b64 else '<div class="hm-logo-icon">⚡</div>'
    )
    st.markdown(
        f'<div class="hm-nav">'
        f'<div class="hm-logo" style="display:flex;align-items:center;gap:13px;">'
        f'{logo_tag}'
        f'<div><div class="hm-logo-name">NEXA</div>'
        f'<div class="hm-logo-tag">AI PLATFORM · POWERED BY NEXA-1.o</div>'
        f'</div></div></div>',
        unsafe_allow_html=True,
    )

    # ── HERO — full-width center ──────────────────────────────────────────────
    st.markdown("""
<div class="hm-hero">
  <div class="hm-title">NEXA</div>
  <p class="hm-sub">Your unified AI platform — chat, generate images, plan your studies, and get AI medical advice.</p>
  <div class="hm-pill">✦ Powered by NEXA-1.o</div>
</div>
""", unsafe_allow_html=True)

    # ── 6 BUTTONS BELOW HERO — equal columns ─────────────────────────────────
    st.markdown('<div class="hm-btn-row-wrap">', unsafe_allow_html=True)
    b1, b2, b3, b4, b5, b6 = st.columns(6, gap="small")
    with b1:
        if st.button("💬\nINFOCHAT", key="btn_ic"):
            st.session_state.page = "infochat"; st.rerun()
    with b2:
        if st.button("🎨\nIMAGINE", key="btn_im"):
            st.session_state.page = "imagine"; st.rerun()
    with b3:
        if st.button("📚\nSTUDY\nPLANNER", key="btn_sp"):
            st.session_state.page = "study_planner"; st.rerun()
    with b4:
        if st.button("🩺\nDR. NEXA", key="btn_dn"):
            st.session_state.page = "dr_nexa"; st.rerun()
    with b5:
        if st.button("💕\nYOUR'S\nNEXA", key="btn_yn"):
            st.session_state.page = "yours_nexa"; st.rerun()
    with b6:
        if st.button("🇮🇳\nINDIA\nTOUR", key="btn_ivt"):
            st.session_state.page = "india_tour"; st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)


# ═════════════════════════════════════════════════════════════════════════════
#
#   I N F O C H A T   ( m o d e l 3 . t x t  —  exact original code )
#
# ═════════════════════════════════════════════════════════════════════════════
def render_infochat():

    logo_b64 = _get_logo_b64()

    # ── CSS (original from model3) ────────────────────────────────────────────
    css = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700;900&family=Rajdhani:wght@300;400;600&display=swap');

:root {
    --neon-purple: #bf5fff;
    --neon-blue:   #00d4ff;
    --neon-pink:   #ff2d9b;
    --dark-bg:     #05030f;
    --card-bg:     rgba(15, 8, 35, 0.85);
    --border:      rgba(191, 95, 255, 0.25);
}

html, body, [data-testid="stAppViewContainer"] {
    background: var(--dark-bg) !important;
    font-family: 'Rajdhani', sans-serif !important;
    color: #e8d5ff !important;
}

[data-testid="stAppViewContainer"]::before {
    content: '';
    position: fixed;
    inset: 0;
    background:
        radial-gradient(ellipse 80% 60% at 20% 10%, rgba(191,95,255,0.18) 0%, transparent 60%),
        radial-gradient(ellipse 60% 50% at 80% 80%, rgba(0,212,255,0.12) 0%, transparent 55%),
        radial-gradient(ellipse 50% 40% at 50% 50%, rgba(255,45,155,0.07) 0%, transparent 60%);
    pointer-events: none;
    z-index: 0;
}

[data-testid="stSidebar"] {
    background: linear-gradient(180deg, rgba(10,4,28,0.97) 0%, rgba(5,2,15,0.99) 100%) !important;
    border-right: 1px solid var(--border) !important;
    box-shadow: 4px 0 40px rgba(191,95,255,0.12) !important;
}
[data-testid="stSidebar"] * { color: #d4b8ff !important; }

#MainMenu, footer, header { visibility: hidden; }
[data-testid="chatAvatarIcon-assistant"],
[data-testid="chatAvatarIcon-user"] { display: none !important; }

@keyframes pulse {
    0%,100% { opacity:1; transform:scale(1); }
    50%      { opacity:0.35; transform:scale(0.88); }
}
@keyframes wave {
    0%,60%,100% { transform:rotate(0deg); }
    10%  { transform:rotate(14deg); }
    20%  { transform:rotate(-8deg); }
    30%  { transform:rotate(14deg); }
    40%  { transform:rotate(-4deg); }
    50%  { transform:rotate(10deg); }
}
@keyframes glowPulse {
    0%,100% { box-shadow: 0 0 12px rgba(191,95,255,0.5), 0 0 24px rgba(191,95,255,0.2); }
    50%      { box-shadow: 0 0 20px rgba(191,95,255,0.9), 0 0 40px rgba(191,95,255,0.4), 0 0 60px rgba(0,212,255,0.2); }
}
@keyframes fadeSlideIn {
    from { opacity:0; transform:translateY(16px); }
    to   { opacity:1; transform:translateY(0); }
}
@keyframes borderFlow {
    0%   { background-position: 0% 50%; }
    50%  { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}

.ai-avatar {
    width: 38px; height: 38px;
    border-radius: 10px;
    object-fit: cover;
    flex-shrink: 0;
    animation: glowPulse 3s ease-in-out infinite;
    border: 1.5px solid rgba(191,95,255,0.6);
}
.thinking-logo {
    width: 38px; height: 38px;
    border-radius: 10px;
    animation: pulse 1s ease-in-out infinite, glowPulse 2s ease-in-out infinite;
    border: 1.5px solid rgba(191,95,255,0.6);
}
.wave-hand {
    display: inline-block;
    animation: wave 2s infinite;
    transform-origin: 70% 70%;
}
.ai-row {
    display: flex;
    align-items: flex-start;
    gap: 12px;
    margin-bottom: 8px;
    animation: fadeSlideIn 0.4s ease-out;
}
.ai-bubble {
    font-size: 1em;
    line-height: 1.7;
    font-family: 'Rajdhani', sans-serif;
    letter-spacing: 0.3px;
    color: #e8d5ff;
    background: var(--card-bg);
    border: 1px solid var(--border);
    border-radius: 0 16px 16px 16px;
    padding: 12px 16px;
    max-width: 820px;
    backdrop-filter: blur(12px);
    box-shadow: 0 4px 24px rgba(191,95,255,0.08), inset 0 1px 0 rgba(255,255,255,0.05);
    position: relative;
    overflow: hidden;
}
.ai-bubble::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(191,95,255,0.6), rgba(0,212,255,0.4), transparent);
}

.nexa-header {
    text-align: center;
    padding: 28px 0 10px;
}
.nexa-title {
    font-family: 'Orbitron', monospace !important;
    font-size: 3.2em !important;
    font-weight: 900 !important;
    letter-spacing: 6px;
    background: linear-gradient(135deg, #bf5fff 0%, #00d4ff 50%, #ff2d9b 100%);
    background-size: 200% 200%;
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    animation: borderFlow 4s ease infinite;
    margin: 0;
    display: block;
}
.nexa-sub {
    font-family: 'Rajdhani', sans-serif;
    font-size: 0.85em;
    letter-spacing: 4px;
    color: rgba(191,95,255,0.7);
    text-transform: uppercase;
    margin-top: 4px;
}
.nexa-divider {
    height: 1px;
    background: linear-gradient(90deg, transparent, rgba(191,95,255,0.5), rgba(0,212,255,0.3), transparent);
    margin: 12px 0 20px;
    border: none;
}
.visual-card {
    background: linear-gradient(135deg, rgba(15,8,35,0.95) 0%, rgba(8,4,20,0.98) 100%);
    border: 1px solid rgba(191,95,255,0.3);
    border-radius: 16px;
    padding: 20px;
    margin-top: 14px;
    position: relative;
    overflow: hidden;
    box-shadow: 0 8px 40px rgba(191,95,255,0.12), inset 0 1px 0 rgba(255,255,255,0.04);
}
.visual-card::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 2px;
    background: linear-gradient(90deg, #bf5fff, #00d4ff, #ff2d9b, #bf5fff);
    background-size: 300% 100%;
    animation: borderFlow 3s linear infinite;
}
.visual-label {
    font-family: 'Orbitron', monospace;
    font-size: 0.65em;
    letter-spacing: 3px;
    color: rgba(191,95,255,0.8);
    text-transform: uppercase;
    margin-bottom: 12px;
    display: flex;
    align-items: center;
    gap: 8px;
}
.visual-label::after {
    content: '';
    flex: 1;
    height: 1px;
    background: linear-gradient(90deg, rgba(191,95,255,0.4), transparent);
}

[data-testid="stChatInput"] {
    background: rgba(15,8,35,0.9) !important;
    border: 1px solid rgba(191,95,255,0.35) !important;
    border-radius: 14px !important;
    box-shadow: 0 0 20px rgba(191,95,255,0.1) !important;
}
[data-testid="stChatInput"]:focus-within {
    border-color: rgba(191,95,255,0.7) !important;
    box-shadow: 0 0 30px rgba(191,95,255,0.25) !important;
}
[data-testid="stChatInput"] textarea {
    color: #e8d5ff !important;
    font-family: 'Rajdhani', sans-serif !important;
    font-size: 1em !important;
}
.stButton > button {
    background: linear-gradient(135deg, rgba(191,95,255,0.15), rgba(0,212,255,0.1)) !important;
    border: 1px solid rgba(191,95,255,0.4) !important;
    color: #d4b8ff !important;
    font-family: 'Rajdhani', sans-serif !important;
    letter-spacing: 1px !important;
    border-radius: 10px !important;
    transition: all 0.2s !important;
}
.stButton > button:hover {
    border-color: rgba(191,95,255,0.8) !important;
    box-shadow: 0 0 16px rgba(191,95,255,0.3) !important;
}

::-webkit-scrollbar { width: 4px; }
::-webkit-scrollbar-track { background: var(--dark-bg); }
::-webkit-scrollbar-thumb { background: rgba(191,95,255,0.4); border-radius: 2px; }

/* ── EXIT BUTTON (inline top-bar) ── */
.exit-bar {
    position: relative; z-index: 20;
    display: flex; align-items: center; justify-content: space-between;
    padding: 10px 28px 0 28px;
}
.exit-bar-title {
    font-family: 'Orbitron', monospace;
    font-size: 0.72em; letter-spacing: 3px;
    color: rgba(191,95,255,0.55); text-transform: uppercase;
}
/* Streamlit button inside exit-bar — styled as pill */
.exit-bar .stButton > button {
    background: rgba(255,45,155,0.10) !important;
    border: 1px solid rgba(255,45,155,0.45) !important;
    color: #ff6bb3 !important;
    font-family: 'Orbitron', monospace !important;
    font-size: 0.68em !important;
    letter-spacing: 2px !important;
    border-radius: 30px !important;
    padding: 6px 20px !important;
    transition: all 0.22s ease !important;
    width: auto !important;
    position: static !important;
    opacity: 1 !important;
    height: auto !important;
    inset: unset !important;
    z-index: auto !important;
}
.exit-bar .stButton > button:hover {
    background: rgba(255,45,155,0.22) !important;
    border-color: rgba(255,45,155,0.85) !important;
    color: #ff2d9b !important;
    box-shadow: 0 0 18px rgba(255,45,155,0.35) !important;
    transform: translateY(-1px) !important;
}
</style>
"""
    st.markdown(css, unsafe_allow_html=True)

    # ── API CONFIG (original model3) ──────────────────────────────────────────
    GROQ_API_KEY = "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"
    GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
    GROQ_MODEL   = "llama-3.1-8b-instant"

    # ── SESSION STATE (original model3 keys, namespaced to avoid conflicts) ──
    if "ic_messages" not in st.session_state:
        st.session_state.ic_messages = [{"role": "assistant", "content": "namaste"}]
    if "ic_lang" not in st.session_state:
        st.session_state.ic_lang = "en"

    # ── SIDEBAR (original model3 sidebar + Back button added) ────────────────
    if logo_b64:
        st.sidebar.markdown(
            f'<img src="data:image/jpeg;base64,{logo_b64}" width="64" '
            'style="border-radius:12px;border:1.5px solid rgba(191,95,255,0.5);'
            'box-shadow:0 0 20px rgba(191,95,255,0.4);margin-bottom:8px"/>',
            unsafe_allow_html=True
        )
    st.sidebar.markdown(
        '<h2 style="font-family:Orbitron,monospace;letter-spacing:3px;'
        'color:#bf5fff;margin:0">NEXA</h2>', unsafe_allow_html=True)
    st.sidebar.markdown(
        '<p style="font-size:0.75em;letter-spacing:2px;'
        'color:rgba(191,95,255,0.6);margin-top:2px">Powered By NEXA-1.o</p>',
        unsafe_allow_html=True)
    st.sidebar.markdown("---")

    # ← Back to Home
    if st.sidebar.button("🏠  Back to Home", key="ic_back", use_container_width=True):
        st.session_state.page = "home"
        st.rerun()

    st.sidebar.markdown("---")
    st.sidebar.subheader("🌐 Language")
    lang_choice = st.sidebar.radio(
        "", ["English", "Hindi", "Hinglish"], label_visibility="collapsed",
        key="ic_lang_radio"
    )
    st.session_state.ic_lang = {"English": "en", "Hindi": "hi", "Hinglish": "hinglish"}[lang_choice]

    st.sidebar.markdown("---")
    st.sidebar.markdown("""
<div style="font-size:0.78em;color:rgba(191,95,255,0.65);line-height:1.8;letter-spacing:0.5px">
⚡ <b>Visual Mode</b><br>
Use keywords like:<br>
<span style="color:#00d4ff">create · generate · draw · visualize<br>
explain visually · diagram · flowchart</span><br><br>
NEXA will respond with <b>text + visual diagram</b>.
</div>
""", unsafe_allow_html=True)

    st.sidebar.markdown("---")
    with st.sidebar:
        if st.button("🗑️ Clear Chat", use_container_width=True, key="ic_clear"):
            st.session_state.ic_messages = [{"role": "assistant", "content": "namaste"}]
            st.rerun()
    st.sidebar.markdown(
        '<div style="font-size:0.72em;text-align:center;color:rgba(191,95,255,0.4);'
        'margin-top:20px;letter-spacing:1px">NEXA · Powered By NEXA-1.o</div>',
        unsafe_allow_html=True
    )

    # ── VISUAL TRIGGER DETECTION (original model3) ───────────────────────────
    VISUAL_KEYWORDS = [
        "create", "generate", "draw", "visualize", "visual", "diagram",
        "flowchart", "show me", "illustrate", "chart", "explain visually",
        "make a", "design", "sketch", "map out", "plot", "architecture",
        "structure", "layout", "workflow", "process", "steps visually",
        "mind map", "mindmap", "timeline", "cycle"
    ]

    def needs_visual(prompt: str) -> bool:
        p = prompt.lower()
        return any(kw in p for kw in VISUAL_KEYWORDS)

    # ── API CALL (original model3) ────────────────────────────────────────────
    def get_ai_response(prompt: str, visual: bool = False) -> str:
        lang = st.session_state.ic_lang
        if lang == "hi":
            lang_instr = "Reply in Hindi only."
        elif lang == "hinglish":
            lang_instr = (
                "Reply in Hinglish — natural Hindi-English mix in Roman script. "
                "Example: 'Yaar, yeh concept bahut interesting hai!'"
            )
        else:
            lang_instr = "Reply in English."

        if visual:
            extra = (
                " The user wants a visual explanation. "
                "First give a clear TEXT explanation (2-4 paragraphs). "
                "Then on a new line write exactly '===VISUAL===' (nothing else on that line). "
                "After that, describe the visual as structured data:\n"
                "TYPE: [flowchart|mindmap|timeline|comparison|architecture|process|cycle]\n"
                "TITLE: <short title>\n"
                "NODES: node1 | node2 | node3 | node4 | node5 (max 8 short labels)\n"
                "CONNECTIONS: node1->node2 | node2->node3 (for flowchart/process/architecture)\n"
                "COLOR_THEME: purple\n"
                "Keep node labels under 20 characters each."
            )
        else:
            extra = ""

        system_prompt = (
            "You are NEXA, a powerful intelligent assistant powered by NEXA-1.o. "
            "Answer clearly and helpfully. " + lang_instr + extra
        )
        headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
        payload = {
            "model": GROQ_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                *[{"role": m["role"], "content": m["content"]}
                  for m in st.session_state.ic_messages[-8:]
                  if m["content"] != "namaste"],
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 900,
            "temperature": 0.7
        }
        try:
            r = requests.post(GROQ_API_URL, json=payload, headers=headers, timeout=30)
            if r.status_code == 200:
                return r.json()["choices"][0]["message"]["content"].strip()
            elif r.status_code == 401:
                return "❌ Invalid API key."
            elif r.status_code == 429:
                return "⏳ Rate limit. Please wait."
            else:
                return f"❌ Error {r.status_code}: {r.text[:150]}"
        except requests.exceptions.Timeout:
            return "⏰ Request timed out."
        except Exception as e:
            return f"💡 Error: {str(e)}"

    # ── PARSE VISUAL DATA (original model3) ──────────────────────────────────
    def parse_visual_data(visual_block: str) -> dict:
        data = {}
        for line in visual_block.strip().splitlines():
            if ":" in line:
                key, _, val = line.partition(":")
                data[key.strip().upper()] = val.strip()
        return data

    # ── BUILD SVG VISUAL (original model3) ───────────────────────────────────
    def build_svg(vdata: dict) -> str:
        vtype     = vdata.get("TYPE", "flowchart").lower()
        title     = vdata.get("TITLE", "Visual Diagram")
        nodes_raw = [n.strip() for n in vdata.get("NODES", "").split("|") if n.strip()]
        conns_raw = [c.strip() for c in vdata.get("CONNECTIONS", "").split("|") if c.strip()]

        W, H = 800, 460
        colors = ["#bf5fff","#00d4ff","#ff2d9b","#ffd700","#7fff00","#ff7f50","#00ffcc","#ff9500"]

        def esc(s):
            return s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

        def wrap(s, limit=18):
            return s[:limit] + ("…" if len(s) > limit else "")

        parts = [
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" '
            f'style="background:#05030f;border-radius:12px;font-family:Rajdhani,sans-serif">',
            '<defs>',
            '<filter id="glow"><feGaussianBlur stdDeviation="3" result="b"/>'
            '<feMerge><feMergeNode in="b"/><feMergeNode in="SourceGraphic"/></feMerge></filter>',
            '<filter id="sg"><feGaussianBlur stdDeviation="7" result="b"/>'
            '<feMerge><feMergeNode in="b"/><feMergeNode in="SourceGraphic"/></feMerge></filter>',
            '<marker id="arr" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">'
            '<polygon points="0 0,10 3.5,0 7" fill="#bf5fff"/></marker>',
            '<radialGradient id="bg" cx="30%" cy="20%" r="75%">'
            '<stop offset="0%" stop-color="#120520"/>'
            '<stop offset="100%" stop-color="#05030f"/></radialGradient>',
            '</defs>',
            f'<rect width="{W}" height="{H}" fill="url(#bg)" rx="12"/>',
            f'<line x1="0" y1="0" x2="40" y2="0" stroke="#bf5fff" stroke-width="2"/>',
            f'<line x1="0" y1="0" x2="0" y2="40" stroke="#bf5fff" stroke-width="2"/>',
            f'<line x1="{W}" y1="0" x2="{W-40}" y2="0" stroke="#00d4ff" stroke-width="2"/>',
            f'<line x1="{W}" y1="0" x2="{W}" y2="40" stroke="#00d4ff" stroke-width="2"/>',
            f'<line x1="0" y1="{H}" x2="40" y2="{H}" stroke="#ff2d9b" stroke-width="2"/>',
            f'<line x1="0" y1="{H}" x2="0" y2="{H-40}" stroke="#ff2d9b" stroke-width="2"/>',
            f'<line x1="{W}" y1="{H}" x2="{W-40}" y2="{H}" stroke="#ffd700" stroke-width="2"/>',
            f'<line x1="{W}" y1="{H}" x2="{W}" y2="{H-40}" stroke="#ffd700" stroke-width="2"/>',
            f'<rect x="0" y="0" width="{W}" height="48" fill="rgba(191,95,255,0.07)" rx="12"/>',
            f'<text x="{W//2}" y="30" text-anchor="middle" font-size="14" font-weight="700" '
            f'fill="#bf5fff" filter="url(#glow)" letter-spacing="3" font-family="Orbitron,monospace">'
            f'{esc(title)}</text>',
            f'<line x1="40" y1="48" x2="{W-40}" y2="48" stroke="rgba(191,95,255,0.3)" stroke-width="1"/>',
        ]

        # ---- FLOWCHART / PROCESS / ARCHITECTURE / WORKFLOW ----
        if vtype in ("flowchart","process","architecture","workflow") and nodes_raw:
            n = len(nodes_raw)
            cols = min(n, 4)
            rows_count = math.ceil(n / cols)
            bw, bh = 155, 48
            hg, vg = 38, 58
            total_w = cols * bw + (cols - 1) * hg
            start_x = (W - total_w) // 2
            start_y = 70

            positions = {}
            for i, node in enumerate(nodes_raw):
                col = i % cols
                row = i // cols
                x = start_x + col * (bw + hg)
                y = start_y + row * (bh + vg)
                cx = x + bw // 2
                cy = y + bh // 2
                positions[node] = (cx, cy, x, y)
                c = colors[i % len(colors)]
                parts += [
                    f'<rect x="{x+2}" y="{y+4}" width="{bw}" height="{bh}" rx="10" '
                    f'fill="{c}" opacity="0.08"/>',
                    f'<rect x="{x}" y="{y}" width="{bw}" height="{bh}" rx="10" '
                    f'fill="rgba(10,4,25,0.92)" stroke="{c}" stroke-width="1.8" filter="url(#glow)"/>',
                    f'<rect x="{x+8}" y="{y+3}" width="{bw-16}" height="3" rx="2" '
                    f'fill="{c}" opacity="0.3"/>',
                    f'<text x="{cx}" y="{cy}" text-anchor="middle" dominant-baseline="middle" '
                    f'font-size="12" fill="{c}" font-weight="600">'
                    f'{esc(wrap(node))}</text>',
                ]

            for conn in conns_raw:
                m = re.split(r"->|→|>>", conn)
                if len(m) == 2:
                    src, dst = m[0].strip(), m[1].strip()
                    sp = next((positions[k] for k in positions if k.lower()==src.lower()), None)
                    dp = next((positions[k] for k in positions if k.lower()==dst.lower()), None)
                    if sp and dp:
                        x1, y1 = sp[0], sp[1] + bh // 2
                        x2, y2 = dp[0], dp[1] - bh // 2
                        parts.append(
                            f'<path d="M{x1},{y1} C{x1},{y1+20} {x2},{y2-20} {x2},{y2}" '
                            f'stroke="rgba(191,95,255,0.55)" stroke-width="1.6" fill="none" '
                            f'marker-end="url(#arr)"/>'
                        )

        # ---- MINDMAP / COMPARISON ----
        elif vtype in ("mindmap","comparison") and nodes_raw:
            cx, cy = W // 2, H // 2
            parts += [
                f'<ellipse cx="{cx}" cy="{cy}" rx="80" ry="36" fill="rgba(191,95,255,0.12)" '
                f'stroke="#bf5fff" stroke-width="2.5" filter="url(#sg)"/>',
                f'<text x="{cx}" y="{cy}" text-anchor="middle" dominant-baseline="middle" '
                f'font-size="13" fill="#bf5fff" font-weight="700">{esc(wrap(title, 16))}</text>',
            ]
            for i, node in enumerate(nodes_raw):
                angle = (2 * math.pi * i / len(nodes_raw)) - math.pi / 2
                r = 165
                nx = int(cx + r * math.cos(angle))
                ny = int(cy + r * math.sin(angle))
                c = colors[i % len(colors)]
                parts += [
                    f'<line x1="{cx}" y1="{cy}" x2="{nx}" y2="{ny}" '
                    f'stroke="{c}" stroke-width="1.2" stroke-opacity="0.35"/>',
                    f'<ellipse cx="{nx}" cy="{ny}" rx="70" ry="27" fill="rgba(10,4,25,0.92)" '
                    f'stroke="{c}" stroke-width="1.8" filter="url(#glow)"/>',
                    f'<text x="{nx}" y="{ny}" text-anchor="middle" dominant-baseline="middle" '
                    f'font-size="11" fill="{c}" font-weight="600">{esc(wrap(node, 18))}</text>',
                ]

        # ---- TIMELINE / CYCLE ----
        elif vtype in ("timeline","cycle") and nodes_raw:
            n = len(nodes_raw)
            step = (W - 140) // max(n, 1)
            ym = H // 2
            parts.append(
                f'<line x1="60" y1="{ym}" x2="{W-60}" y2="{ym}" '
                f'stroke="rgba(191,95,255,0.4)" stroke-width="2.5"/>'
            )
            parts.append(
                f'<polygon points="{W-55},{ym-5} {W-40},{ym} {W-55},{ym+5}" fill="#bf5fff"/>'
            )
            for i, node in enumerate(nodes_raw):
                x = 70 + i * step + step // 2
                c = colors[i % len(colors)]
                above = i % 2 == 0
                ly = ym - 70 if above else ym + 70
                ly2 = ym - 26 if above else ym + 26
                parts += [
                    f'<line x1="{x}" y1="{ym}" x2="{x}" y2="{ly2}" '
                    f'stroke="{c}" stroke-width="1.4" stroke-dasharray="4,3"/>',
                    f'<circle cx="{x}" cy="{ym}" r="9" fill="{c}" filter="url(#glow)"/>',
                    f'<circle cx="{x}" cy="{ym}" r="4" fill="white" opacity="0.6"/>',
                    f'<rect x="{x-58}" y="{ly-22}" width="116" height="44" rx="10" '
                    f'fill="rgba(10,4,25,0.92)" stroke="{c}" stroke-width="1.5"/>',
                    f'<text x="{x}" y="{ly}" text-anchor="middle" dominant-baseline="middle" '
                    f'font-size="11" fill="{c}" font-weight="600">{esc(wrap(node, 18))}</text>',
                ]

        # ---- FALLBACK GRID ----
        else:
            for i, node in enumerate(nodes_raw[:8]):
                col = i % 4
                row = i // 4
                x = 55 + col * 185
                y = 75 + row * 170
                c = colors[i % len(colors)]
                parts += [
                    f'<rect x="{x}" y="{y}" width="160" height="55" rx="10" '
                    f'fill="rgba(10,4,25,0.92)" stroke="{c}" stroke-width="1.8" filter="url(#glow)"/>',
                    f'<text x="{x+80}" y="{y+28}" text-anchor="middle" dominant-baseline="middle" '
                    f'font-size="12" fill="{c}" font-weight="600">{esc(wrap(node))}</text>',
                ]

        parts.append('</svg>')
        return "\n".join(parts)

    # ── RENDER ASSISTANT (original model3) ───────────────────────────────────
    def render_assistant(content, show_visual=False, visual_block=""):
        if logo_b64:
            img_tag = f'<img class="ai-avatar" src="data:image/jpeg;base64,{logo_b64}"/>'
        else:
            img_tag = ('<div style="width:38px;height:38px;border-radius:10px;flex-shrink:0;'
                       'background:linear-gradient(135deg,#bf5fff,#00d4ff);'
                       'animation:glowPulse 3s ease-in-out infinite"></div>')

        if content == "namaste":
            bubble = (
                'Namaste <span class="wave-hand">🙏</span>'
                ' &nbsp;I\'m <b style="color:#bf5fff">NEXA</b> powered by '
                '<b style="color:#00d4ff">NEXA-1.o</b>. Ask me anything!<br>'
                '<span style="font-size:0.85em;color:rgba(191,95,255,0.6);letter-spacing:0.5px">'
                '✦ Try: <i>"create a flowchart of machine learning"</i> or '
                '<i>"generate a mindmap of Python basics"</i></span>'
            )
        else:
            bubble = content

        st.markdown(
            f'<div class="ai-row">{img_tag}'
            f'<div class="ai-bubble">{bubble}</div></div>',
            unsafe_allow_html=True
        )

        if show_visual and visual_block:
            vdata = parse_visual_data(visual_block)
            if vdata.get("NODES") or vdata.get("TITLE"):
                svg_code = build_svg(vdata)
                st.markdown(
                    f'<div class="visual-card">'
                    f'<div class="visual-label">⬡ NEXA VISUAL ENGINE · {vdata.get("TYPE","DIAGRAM").upper()}</div>'
                    f'{svg_code}'
                    f'</div>',
                    unsafe_allow_html=True
                )

    # ── EXIT BUTTON BAR ───────────────────────────────────────────────────────
    st.markdown('<div class="exit-bar">', unsafe_allow_html=True)
    _ec1, _ec2 = st.columns([6, 1])
    with _ec1:
        st.markdown(
            '<span class="exit-bar-title">💬 INFOCHAT &nbsp;·&nbsp; Visual Intelligence</span>',
            unsafe_allow_html=True,
        )
    with _ec2:
        if st.button("✕  Exit", key="ic_exit_top", help="Return to NEXA Home"):
            st.session_state.page = "home"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    # ── HEADER (original model3) ──────────────────────────────────────────────
    st.markdown(
        '<div class="nexa-header">'
        '<span class="nexa-title">NEXA</span>'
        '<div class="nexa-sub">Powered By NEXA-1.o · Visual Intelligence</div>'
        '</div>',
        unsafe_allow_html=True
    )
    st.markdown('<hr class="nexa-divider">', unsafe_allow_html=True)

    # ── CHAT HISTORY (original model3) ───────────────────────────────────────
    for msg in st.session_state.ic_messages:
        if msg["role"] == "assistant":
            render_assistant(
                msg["content"],
                show_visual=msg.get("has_visual", False),
                visual_block=msg.get("visual_block", "")
            )
        else:
            with st.chat_message("user"):
                st.markdown(
                    f'<span style="color:#e8d5ff;font-family:Rajdhani,sans-serif">{msg["content"]}</span>',
                    unsafe_allow_html=True
                )

    # ── CHAT INPUT (original model3) ─────────────────────────────────────────
    prompt = st.chat_input("✦ Ask NEXA anything... or say 'create a flowchart of...'")

    if prompt:
        st.session_state.ic_messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(
                f'<span style="color:#e8d5ff;font-family:Rajdhani,sans-serif">{prompt}</span>',
                unsafe_allow_html=True
            )

        visual_mode = needs_visual(prompt)

        if logo_b64:
            img_src = f'data:image/jpeg;base64,{logo_b64}'
            thinking_img = f'<img class="thinking-logo" src="{img_src}"/>'
        else:
            thinking_img = ('<div style="width:38px;height:38px;border-radius:10px;flex-shrink:0;'
                            'background:linear-gradient(135deg,#bf5fff,#00d4ff);'
                            'animation:pulse 1s ease-in-out infinite"></div>')

        thinking = st.empty()
        thinking.markdown(
            f'<div class="ai-row">{thinking_img}'
            f'<div class="ai-bubble" style="opacity:0.6;font-style:italic;color:#bf5fff">'
            f'{"🎨 NEXA Visual Engine is generating..." if visual_mode else "⚡ NEXA-1.o is thinking..."}'
            f'</div></div>',
            unsafe_allow_html=True
        )

        raw = get_ai_response(prompt, visual=visual_mode)
        thinking.empty()

        has_visual    = False
        visual_block  = ""
        text_response = raw

        if visual_mode and "===VISUAL===" in raw:
            parts_split   = raw.split("===VISUAL===", 1)
            text_response = parts_split[0].strip()
            visual_block  = parts_split[1].strip()
            has_visual    = True

        render_assistant(text_response, show_visual=has_visual, visual_block=visual_block)
        st.session_state.ic_messages.append({
            "role":         "assistant",
            "content":      text_response,
            "has_visual":   has_visual,
            "visual_block": visual_block
        })
        st.rerun()


# ═════════════════════════════════════════════════════════════════════════════
#
#   I M A G I N E   ( T e x t _ t o _ m o d e l _ 1 . t x t  —  exact code )
#
# ═════════════════════════════════════════════════════════════════════════════
def render_imagine():

    logo_b64 = _get_logo_b64()

    # ── MODEL LOADER (original) ───────────────────────────────────────────────
    @st.cache_resource
    def load_model():
        from diffusers import StableDiffusionPipeline
        import torch
        pipe = StableDiffusionPipeline.from_pretrained(
            "runwayml/stable-diffusion-v1-5",
            torch_dtype=torch.float32,
        )
        pipe.enable_attention_slicing()
        pipe.to("cpu")
        return pipe

    # ── STYLE MAP (original) ──────────────────────────────────────────────────
    style_suffix_map = {
        "Photorealistic": "photorealistic, 8k, ultra detailed, sharp focus",
        "Anime":          "anime style, vibrant colors, Studio Ghibli",
        "Oil Painting":   "oil painting, classical art, textured canvas",
        "Cinematic":      "cinematic, dramatic lighting, movie still, anamorphic lens",
        "Sketch":         "pencil sketch, hand drawn, detailed linework",
    }
    styles = list(style_suffix_map.keys())

    # ── CSS (original Text_to_model_1.txt) ───────────────────────────────────
    st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@700;800&family=DM+Sans:wght@400;500;600&display=swap');

/* ── DESIGN TOKENS ────────────────────────────────── */
:root {
    --bg:           #0a0f0d;
    --border:       rgba(0,255,140,0.15);
    --accent:       #00e87a;
    --accent2:      #00c96a;
    --text:         #e8f5ee;
    --muted:        #8da898;
    --box-bg:       rgba(17,26,20,0.88);
    --box-border:   2px solid var(--accent);
    --box-radius:   16px;
    --box-shadow:   0 0 22px rgba(0,232,122,0.14), inset 0 0 14px rgba(0,232,122,0.04);
    --box-width:    100%;
    --bar-height:   50px;
    --bar-pad-v:    14px;
    --bar-pad-h:    20px;
    --space-xs:  clamp(6px,  1vw, 10px);
    --space-sm:  clamp(10px, 2vw, 16px);
    --space-md:  clamp(16px, 3vw, 28px);
    --space-lg:  clamp(28px, 5vw, 48px);
    --space-xl:  clamp(40px, 7vw, 72px);
}

/* ── RESET + BASE ─────────────────────────────────── */
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

html, body,
[data-testid="stAppViewContainer"],
[data-testid="stApp"] {
    background-color: var(--bg) !important;
    color: var(--text);
    font-family: 'DM Sans', sans-serif;
    -webkit-font-smoothing: antialiased;
}
[data-testid="stAppViewContainer"] > .main { background-color: var(--bg) !important; }
.block-container { padding: 0 !important; max-width: 100% !important; }
#MainMenu, footer, header { visibility: hidden; }

/* ── SIDEBAR ── */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, rgba(10,15,13,0.97) 0%, rgba(5,8,6,0.99) 100%) !important;
    border-right: 1px solid var(--border) !important;
}
[data-testid="stSidebar"] * { color: #c8e8d8 !important; }

/* ── NAV ── */
.nav {
    position: relative; z-index: 10;
    display: flex; align-items: center; justify-content: space-between;
    padding: clamp(12px, 2vw, 18px) clamp(16px, 4vw, 44px);
    border-bottom: 1px solid var(--border);
    background: rgba(10,15,13,0.82);
    backdrop-filter: blur(16px);
    -webkit-backdrop-filter: blur(16px);
}
.logo { display: flex; align-items: center; gap: 10px; }
.logo-icon {
    width: 38px; height: 38px; background: var(--accent);
    border-radius: 10px; display: flex; align-items: center;
    justify-content: center; font-size: 18px; flex-shrink: 0;
}
.logo-icon-img {
    width: 38px; height: 38px; border-radius: 10px;
    object-fit: cover; box-shadow: 0 0 14px rgba(0,232,122,0.45);
    flex-shrink: 0;
}

@keyframes nexaGlowPulse {
    0%, 100% { text-shadow: 0 0 8px rgba(0,232,122,0.55), 0 0 18px rgba(0,232,122,0.35), 0 0 32px rgba(0,255,204,0.20); }
    50%       { text-shadow: 0 0 14px rgba(0,232,122,0.85), 0 0 28px rgba(0,232,122,0.55), 0 0 50px rgba(0,255,204,0.35); }
}
.logo-name {
    font-family: 'Syne', sans-serif;
    font-size: clamp(13px, 1.2vw, 16px);
    font-weight: 800; letter-spacing: 2px;
    background: linear-gradient(135deg, #00e87a 0%, #00ffcc 50%, #00c96a 100%);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;
    animation: nexaGlowPulse 2.5s ease-in-out infinite;
}

/* ── HERO ── */
@keyframes heroEnter {
    0%   { opacity: 0; transform: translateY(20px); }
    100% { opacity: 1; transform: translateY(0); }
}
.hero-wrap {
    position: relative; z-index: 5;
    display: flex; flex-direction: column; align-items: center;
    padding: var(--space-xl) clamp(16px, 6vw, 80px) var(--space-lg);
    text-align: center;
    animation: heroEnter 0.8s cubic-bezier(0.22, 1, 0.36, 1) both;
}
@keyframes wordFadeUp {
    0%   { opacity: 0; transform: translateY(26px); filter: blur(10px); }
    60%  { opacity: 1; filter: blur(0); }
    100% { opacity: 1; transform: translateY(0); filter: blur(0); }
}
@keyframes headingBreathe {
    0%, 100% { filter: drop-shadow(0 0 6px rgba(0,232,122,0.18)) drop-shadow(0 0 18px rgba(0,232,122,0.08)); }
    50%       { filter: drop-shadow(0 0 14px rgba(0,232,122,0.42)) drop-shadow(0 0 32px rgba(0,255,204,0.16)); }
}
.hero-title-animated {
    font-family: 'Syne', sans-serif;
    font-size: clamp(22px, 4.2vw, 54px);
    font-weight: 800; line-height: 1.12; letter-spacing: -0.025em;
    margin-bottom: var(--space-md);
    display: flex; flex-wrap: wrap; justify-content: center;
    align-items: baseline; gap: 0 0.27em;
    max-width: min(860px, 92vw); width: 100%;
    animation: headingBreathe 4s ease-in-out infinite;
    animation-delay: 2.8s; animation-fill-mode: both; will-change: filter;
}
.hero-word {
    display: inline-block; opacity: 0;
    animation: wordFadeUp 0.7s cubic-bezier(0.22, 0.61, 0.36, 1) forwards;
    will-change: transform, opacity; color: #ffffff; -webkit-text-fill-color: #ffffff;
    cursor: default;
    transition: transform 0.24s cubic-bezier(0.34, 1.56, 0.64, 1), filter 0.24s ease;
}
.hero-word:hover { transform: translateY(-4px) scale(1.04); filter: brightness(1.25) drop-shadow(0 0 12px rgba(255,255,255,0.5)); }
.hero-word.accent {
    background: linear-gradient(135deg, #00e87a 0%, #00ffcc 55%, #7fffcb 100%);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;
}
.hero-word.accent:hover { filter: brightness(1.3) drop-shadow(0 0 14px rgba(0,232,122,0.75)); }
@media (prefers-reduced-motion: reduce) {
    .hero-wrap { animation: none !important; }
    .hero-word { animation: none !important; opacity: 1 !important; transform: none !important; filter: none !important; }
    .hero-title-animated { animation: none !important; }
    .powered-badge { animation: none !important; opacity: 1 !important; }
    .hero-sub { animation: none !important; opacity: 1 !important; }
}
@keyframes badgeEnter {
    0%   { opacity: 0; transform: translateY(14px) scale(0.96); }
    100% { opacity: 1; transform: translateY(0) scale(1); }
}
@keyframes badgeBreathe {
    0%, 100% { box-shadow: 0 0 0 1px rgba(0,232,122,0.20), 0 0 12px rgba(0,232,122,0.10), inset 0 0 8px rgba(0,232,122,0.04); }
    50%       { box-shadow: 0 0 0 1px rgba(0,232,122,0.38), 0 0 22px rgba(0,232,122,0.22), inset 0 0 14px rgba(0,232,122,0.08); }
}
.powered-badge {
    display: inline-flex; align-items: center; gap: 8px;
    padding: 7px 18px; border-radius: 999px;
    font-size: clamp(11px, 1vw, 13px); font-weight: 500;
    color: var(--accent); letter-spacing: 0.3px;
    margin-bottom: var(--space-sm);
    background: rgba(0, 232, 122, 0.07);
    border: 1px solid rgba(0, 232, 122, 0.25);
    backdrop-filter: blur(12px); -webkit-backdrop-filter: blur(12px);
    animation: badgeEnter 0.7s cubic-bezier(0.22, 1, 0.36, 1) 2.2s both, badgeBreathe 5s ease-in-out 3.2s infinite;
    will-change: transform, box-shadow;
    transition: transform 0.24s cubic-bezier(0.34, 1.56, 0.64, 1), box-shadow 0.24s ease;
}
.powered-badge:hover {
    transform: scale(1.04);
    box-shadow: 0 0 0 1px rgba(0,232,122,0.55), 0 0 28px rgba(0,232,122,0.32), inset 0 0 16px rgba(0,232,122,0.10);
}
@keyframes subEnter {
    0%   { opacity: 0; transform: translateY(12px); }
    100% { opacity: 1; transform: translateY(0); }
}
.hero-sub {
    color: var(--muted); font-size: clamp(13px, 1.2vw, 15px); line-height: 1.75;
    max-width: min(460px, 90vw);
    animation: subEnter 0.7s cubic-bezier(0.22, 1, 0.36, 1) 2.5s both;
    will-change: transform, opacity;
}
.section-label {
    font-size: clamp(10px, 0.9vw, 11px); letter-spacing: 2px;
    text-transform: uppercase; color: var(--muted); white-space: nowrap;
}
.style-tags { display: flex; flex-wrap: nowrap; gap: 8px; align-items: center; overflow-x: auto; scrollbar-width: none; }
.style-tags::-webkit-scrollbar { display: none; }
.tag {
    padding: 4px 14px; border-radius: 999px;
    border: 1px solid rgba(255,255,255,0.10); background: rgba(255,255,255,0.03);
    font-size: clamp(11px, 0.9vw, 12px); color: var(--muted);
    cursor: pointer; font-family: 'DM Sans', sans-serif;
    user-select: none; white-space: nowrap;
    transition: border-color 0.2s ease, color 0.2s ease, background 0.2s ease, transform 0.2s cubic-bezier(0.34,1.56,0.64,1);
}
.tag:hover { border-color: var(--accent); color: var(--accent); background: rgba(0,232,122,0.08); transform: translateY(-1px); }
.tag.active { border-color: var(--accent); color: var(--accent); background: rgba(0,232,122,0.10); }
[data-testid="stTextArea"] { margin-top: 6px !important; margin-bottom: 12px !important; width: var(--box-width) !important; }
[data-testid="stTextArea"] label { display: none !important; }
[data-testid="stTextArea"] > div {
    background: var(--box-bg) !important; border: var(--box-border) !important;
    border-radius: var(--box-radius) !important; box-shadow: var(--box-shadow) !important;
    padding: 0 !important; height: var(--bar-height) !important;
    min-height: var(--bar-height) !important; display: flex !important;
    align-items: center !important; transition: box-shadow 0.25s ease !important;
}
[data-testid="stTextArea"] > div:focus-within {
    box-shadow: 0 0 0 2px rgba(0,232,122,0.35), 0 0 22px rgba(0,232,122,0.18) !important;
}
[data-testid="stTextArea"] textarea {
    background: transparent !important; border: none !important;
    border-radius: var(--box-radius) !important; color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important; font-size: clamp(13px, 1.1vw, 14px) !important;
    padding: var(--bar-pad-v) var(--bar-pad-h) !important; caret-color: var(--accent) !important;
    resize: none !important; height: var(--bar-height) !important;
    min-height: var(--bar-height) !important; max-height: var(--bar-height) !important;
    width: 100% !important; box-shadow: none !important; outline: none !important;
    overflow-y: auto !important; line-height: 1.4 !important;
}
[data-testid="stTextArea"] textarea:focus { background: rgba(0,232,122,0.04) !important; }
[data-testid="stTextArea"] textarea::placeholder { color: rgba(0,232,122,0.45) !important; opacity: 1 !important; font-style: italic !important; }
.card-top {
    background: var(--box-bg); border: var(--box-border); border-radius: var(--box-radius);
    box-shadow: var(--box-shadow); backdrop-filter: blur(14px); -webkit-backdrop-filter: blur(14px);
    width: var(--box-width); margin-bottom: 12px;
    display: flex; align-items: center; gap: 16px;
    height: var(--bar-height); min-height: var(--bar-height);
    padding: var(--bar-pad-v) var(--bar-pad-h); overflow: hidden;
}
[data-testid="stExpander"] {
    background: var(--box-bg) !important; border: var(--box-border) !important;
    border-radius: var(--box-radius) !important; box-shadow: var(--box-shadow) !important;
    margin-top: 0 !important; margin-bottom: 12px !important;
    width: var(--box-width) !important; overflow: hidden !important;
}
[data-testid="stExpander"] summary {
    color: var(--accent) !important; font-size: clamp(12px, 1vw, 13px) !important;
    font-weight: 600 !important; padding: var(--bar-pad-v) var(--bar-pad-h) !important;
    letter-spacing: 0.5px !important; min-height: var(--bar-height) !important;
    display: flex !important; align-items: center !important; transition: color 0.2s ease !important;
}
[data-testid="stExpander"] summary:hover { color: var(--text) !important; }
[data-testid="stExpander"] > div > div { padding: 8px 24px 18px !important; }
[data-testid="stSlider"] > div > div > div > div { background: var(--accent) !important; }
[data-testid="stSlider"] label { color: var(--muted) !important; font-size: 13px !important; }
[data-testid="stButton"] > button {
    width: 100% !important; padding: 15px !important;
    background: var(--accent) !important; color: #0a0f0d !important;
    border: none !important; border-radius: 0 0 16px 16px !important;
    font-size: clamp(13px, 1.1vw, 15px) !important; font-weight: 700 !important;
    font-family: 'DM Sans', sans-serif !important; letter-spacing: 0.5px !important;
    transition: background 0.22s ease, box-shadow 0.22s ease, transform 0.22s cubic-bezier(0.34,1.56,0.64,1) !important;
    box-shadow: none !important; margin-top: 0 !important;
}
[data-testid="stButton"] > button:hover {
    background: var(--accent2) !important;
    box-shadow: 0 6px 32px rgba(0,232,122,0.40) !important;
    transform: translateY(-2px) !important;
}
[data-testid="stButton"] > button:active { transform: translateY(0) !important; }
[data-testid="stVerticalBlock"] > [data-testid="stVerticalBlockBorderWrapper"],
div[data-testid="stVerticalBlock"] > div { gap: 0 !important; }
[data-testid="stImage"] img {
    border-radius: 16px !important; border: 1px solid var(--border) !important;
    box-shadow: 0 0 60px rgba(0,232,122,0.12) !important;
}
[data-testid="stDownloadButton"] > button {
    background: transparent !important; border: 1px solid var(--border) !important;
    color: var(--accent) !important; border-radius: 10px !important;
    font-family: 'DM Sans', sans-serif !important; font-weight: 600 !important;
    width: 100% !important; margin-top: 8px !important;
    transition: border-color 0.2s ease, box-shadow 0.2s ease !important;
}
[data-testid="stDownloadButton"] > button:hover {
    border-color: var(--accent) !important; box-shadow: 0 0 18px rgba(0,232,122,0.22) !important;
}
[data-testid="stAlert"] {
    background: rgba(0,232,122,0.06) !important; border: 1px solid var(--border) !important;
    border-radius: 12px !important; color: var(--text) !important;
}
[data-testid="stColumn"] > div { width: 100% !important; }
@media (max-width: 768px) { .hero-title-animated { gap: 0 0.22em; letter-spacing: -0.018em; } .nav { flex-wrap: nowrap; } }
@media (max-width: 480px) { .hero-title-animated { gap: 0 0.16em; letter-spacing: -0.01em; } .hero-wrap { padding: var(--space-lg) 16px var(--space-md); } }
@media (min-width: 1600px) { .hero-title-animated { max-width: 960px; } .hero-sub { max-width: 520px; } }

/* ── EXIT BUTTON in NAV (IMAGINE) ── */
.nav-exit-btn > button,
.im-exit .stButton > button {
    background: rgba(255,45,155,0.10) !important;
    border: 1px solid rgba(255,45,155,0.45) !important;
    color: #ff6bb3 !important;
    font-family: 'Syne', sans-serif !important;
    font-size: 0.72em !important;
    font-weight: 700 !important;
    letter-spacing: 2px !important;
    border-radius: 30px !important;
    padding: 7px 22px !important;
    transition: all 0.22s ease !important;
    width: auto !important;
    position: static !important;
    opacity: 1 !important;
    height: auto !important;
    inset: unset !important;
    z-index: auto !important;
    cursor: pointer !important;
    margin: 0 !important;
}
.im-exit .stButton > button:hover {
    background: rgba(255,45,155,0.22) !important;
    border-color: rgba(255,45,155,0.85) !important;
    color: #ff2d9b !important;
    box-shadow: 0 0 18px rgba(255,45,155,0.35) !important;
    transform: translateY(-1px) !important;
}
</style>""", unsafe_allow_html=True)

    # ── SIDEBAR (original + Back button) ─────────────────────────────────────
    with st.sidebar:
        if logo_b64:
            st.markdown(
                f'<img src="data:image/jpeg;base64,{logo_b64}" width="64" '
                'style="border-radius:12px;border:1.5px solid rgba(0,232,122,0.5);'
                'box-shadow:0 0 20px rgba(0,232,122,0.4);margin-bottom:8px"/>',
                unsafe_allow_html=True
            )
        st.markdown(
            '<h2 style="font-family:Syne,sans-serif;letter-spacing:3px;color:#00e87a;margin:0">NEXA</h2>',
            unsafe_allow_html=True)
        st.markdown(
            '<p style="font-size:.75em;letter-spacing:2px;color:rgba(0,232,122,.6);margin-top:2px">IMAGINE · Text to Image</p>',
            unsafe_allow_html=True)
        st.markdown("---")

        # ← Back to Home
        if st.button("🏠  Back to Home", key="im_back", use_container_width=True):
            st.session_state.page = "home"
            st.rerun()

        st.markdown("---")
        st.markdown("""
<div style="font-size:.78em;color:rgba(0,232,122,.65);line-height:1.9">
🎨 <b>How to use:</b><br>
<span style="color:#00d4ff">① Type your image description</span><br>
<span style="color:#00e87a">② Choose a visual style</span><br>
<span style="color:#ffd700">③ Adjust settings if needed</span><br>
<span style="color:#ff2d9b">④ Click Generate & download!</span>
</div>""", unsafe_allow_html=True)
        st.markdown(
            '<div style="font-size:.72em;text-align:center;color:rgba(0,232,122,.4);'
            'margin-top:20px;letter-spacing:1px">NEXA · IMAGINE · NEXA-1.o</div>',
            unsafe_allow_html=True)

    # ── SESSION STATE (original) ──────────────────────────────────────────────
    if "im_style" not in st.session_state:
        st.session_state.im_style = "Photorealistic"

    # Handle style selection from URL params (original logic)
    params = st.query_params
    if "set_style" in params:
        val = params["set_style"]
        if val in styles:
            st.session_state.im_style = val
        st.query_params.clear()
        st.rerun()

    # ── PARTICLE CANVAS + JS (original model1) ────────────────────────────────
    st.markdown("""<canvas id="bg"></canvas>
<script>
(function(){
    var c=document.getElementById('bg');
    if(!c)return;
    var x=c.getContext('2d'),W,H,N=[];
    function resize(){W=c.width=innerWidth;H=c.height=innerHeight;}
    resize();
    window.addEventListener('resize',resize);
    for(var i=0;i<80;i++){
        N.push({
            x:Math.random()*innerWidth,y:Math.random()*innerHeight,
            vx:(Math.random()-.5)*.35,vy:(Math.random()-.5)*.35,
            r:Math.random()*2+1
        });
    }
    function draw(){
        x.clearRect(0,0,W,H);
        for(var i=0;i<N.length;i++){
            var n=N[i];
            n.x+=n.vx;n.y+=n.vy;
            if(n.x<0||n.x>W)n.vx*=-1;
            if(n.y<0||n.y>H)n.vy*=-1;
            x.beginPath();x.arc(n.x,n.y,n.r,0,Math.PI*2);
            x.fillStyle='#00e87a';x.fill();
            for(var j=i+1;j<N.length;j++){
                var m=N[j],dx=n.x-m.x,dy=n.y-m.y,d=Math.sqrt(dx*dx+dy*dy);
                if(d<130){
                    x.beginPath();x.moveTo(n.x,n.y);x.lineTo(m.x,m.y);
                    x.strokeStyle='rgba(0,232,122,'+(0.15*(1-d/130))+')';
                    x.lineWidth=0.8;x.stroke();
                }
            }
        }
        requestAnimationFrame(draw);
    }
    draw();
})();
</script>""", unsafe_allow_html=True)

    # ── NAV BAR (original model1 + Exit button) ──────────────────────────────
    logo_html = '<div class="logo-icon">&#9889;</div>'
    if logo_b64:
        logo_html = f'<img src="data:image/jpeg;base64,{logo_b64}" class="logo-icon-img"/>'

    st.markdown(
        '<div class="nav">'
        '<div class="logo">'
        + logo_html +
        '<span class="logo-name">NEXA</span>'
        '</div>'
        '<div class="im-exit" style="display:flex;align-items:center;">',
        unsafe_allow_html=True,
    )
    if st.button("✕  Exit", key="im_exit_nav", help="Return to NEXA Home"):
        st.session_state.page = "home"
        st.rerun()
    st.markdown('</div></div>', unsafe_allow_html=True)

    # ── HERO (original model1) ────────────────────────────────────────────────
    st.markdown("""<div class="hero-wrap">

  <h1 class="hero-title-animated" aria-label="Turn Your Words Into Your Thoughts">
    <span class="hero-word"        style="animation-delay:0.10s" aria-hidden="true">Turn</span>
    <span class="hero-word"        style="animation-delay:0.45s" aria-hidden="true">Your</span>
    <span class="hero-word"        style="animation-delay:0.80s" aria-hidden="true">Words</span>
    <span class="hero-word accent" style="animation-delay:1.15s" aria-hidden="true">Into</span>
    <span class="hero-word accent" style="animation-delay:1.50s" aria-hidden="true">Your</span>
    <span class="hero-word accent" style="animation-delay:1.85s" aria-hidden="true">Thoughts</span>
  </h1>

  <div class="powered-badge" role="text" aria-label="Powered by model NEXA-1.o">
    &#10022; Powered by model NEXA-1.o &rarr;
  </div>

  <p class="hero-sub">
    Describe anything you can imagine. NEXA transforms your
    text into stunning, high-resolution visuals ✨ in seconds.
  </p>

</div>""", unsafe_allow_html=True)

    # ── BUILD STYLE TAGS (original model1) ───────────────────────────────────
    tags_html = ""
    for s in styles:
        active = "active" if s == st.session_state.im_style else ""
        tags_html += (
            '<span class="tag ' + active + '" '
            "onclick=\"window.location.href='?set_style=" + s + "'\">" + s + "</span>"
        )

    # ── MAIN CARD (original model1) ───────────────────────────────────────────
    _, mid, _ = st.columns([1, 1.6, 1])
    with mid:

        # 1 — PROMPT BOX
        st.markdown(
            '<p class="section-label">&#10022; &nbsp;Your Prompt</p>',
            unsafe_allow_html=True
        )
        prompt = st.text_area(
            label="prompt",
            placeholder="✦ ✦ ✦  Let's Generate Your Thoughts  ✦ ✦ ✦ ",
            height=50,
            label_visibility="collapsed",
            key="im_user_prompt"
        )

        # 2 — STYLE CARD
        st.markdown(
            '<div class="card-top">'
            '<span class="section-label">Style</span>'
            '<div class="style-tags">' + tags_html + '</div>'
            '</div>',
            unsafe_allow_html=True
        )

        # 3 — ADVANCED SETTINGS
        with st.expander("⚙️ Advanced Settings"):
            steps    = st.slider("Inference Steps", 10, 50, 20, key="im_steps")
            guidance = st.slider("Guidance Scale", 1.0, 15.0, 7.5, key="im_guidance")

        # 4 — GENERATE BUTTON
        generate = st.button("✦  Generate Image  ✦", key="im_generate_btn")

    # ── GENERATE (original model1) ────────────────────────────────────────────
    if generate:
        prompt_val = st.session_state.get("im_user_prompt", "").strip()
        if not prompt_val:
            _, warn_col, _ = st.columns([1, 1.6, 1])
            with warn_col:
                st.warning("⚠️ Please type a prompt above, then click Generate!")
        else:
            _, out_col, _ = st.columns([1, 1.6, 1])
            with out_col:
                with st.spinner("✨ Encoding your thoughts..."):
                    try:
                        pipe = load_model()
                    except Exception as e:
                        st.error(f"Model failed to load: {e}")
                        st.stop()

                full_prompt = prompt_val + ", " + style_suffix_map[st.session_state.im_style]

                with st.spinner("🎨 Generating in " + st.session_state.im_style + " style... please wait ✨"):
                    try:
                        image = pipe(
                            full_prompt,
                            num_inference_steps=steps,
                            guidance_scale=guidance
                        ).images[0]

                        st.image(
                            image,
                            use_container_width=True,
                            caption='"' + prompt_val + '" — ' + st.session_state.im_style
                        )
                        image.save("output.png")

                        with open("output.png", "rb") as f:
                            st.download_button(
                                label="⬇️ Download Image",
                                data=f,
                                file_name="nexa_output.png",
                                mime="image/png",
                                key="im_download_btn"
                            )
                    except Exception as e:
                        st.error(f"Generation failed: {e}")


# ═════════════════════════════════════════════════════════════════════════════
#
#  S T U D Y   P L A N N E R  ( m o d e l _ 4 _ p l a n n e r . t x t )
#
# ═════════════════════════════════════════════════════════════════════════════
def render_study_planner():
    import json, uuid
    from datetime import datetime, timedelta
    from io import BytesIO

    # ── optional deps ─────────────────────────────────────────────────────────
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, HRFlowable)
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors as rl_colors
        from reportlab.lib.units import cm
        from reportlab.lib.enums import TA_CENTER
        REPORTLAB_OK = True
    except ImportError:
        REPORTLAB_OK = False

    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        OPENPYXL_OK = True
    except ImportError:
        OPENPYXL_OK = False

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        DOCX_OK = True
    except ImportError:
        DOCX_OK = False

    # ── API CONFIG (original model4) ─────────────────────────────────────────
    GROQ_API_KEY = "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"
    GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
    GROQ_MODEL   = "llama-3.1-8b-instant"

    logo_b64 = _get_logo_b64()

    # ── SESSION STATE (namespaced with sp_ prefix) ────────────────────────────
    sp_defaults = {
        "sp_messages":   [{"role": "assistant", "content": "namaste", "msg_id": "init"}],
        "sp_lang":       "en",
        "sp_last_plan":  None,
        "sp_file_cache": {},
    }
    for k, v in sp_defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # ── CSS (original model4) ─────────────────────────────────────────────────
    st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700;900&family=Rajdhani:wght@300;400;600&display=swap');
:root{--np:#bf5fff;--nb:#00d4ff;--npi:#ff2d9b;--db:#05030f;--brd:rgba(191,95,255,.25);}
html,body,[data-testid="stAppViewContainer"]{background:var(--db)!important;font-family:'Rajdhani',sans-serif!important;color:#e8d5ff!important;}
[data-testid="stAppViewContainer"]::before{content:'';position:fixed;inset:0;pointer-events:none;z-index:0;
  background:radial-gradient(ellipse 80% 60% at 20% 10%,rgba(191,95,255,.18) 0%,transparent 60%),
             radial-gradient(ellipse 60% 50% at 80% 80%,rgba(0,212,255,.12) 0%,transparent 55%);}
[data-testid="stSidebar"]{background:linear-gradient(180deg,rgba(10,4,28,.97) 0%,rgba(5,2,15,.99) 100%)!important;
  border-right:1px solid var(--brd)!important;box-shadow:4px 0 40px rgba(191,95,255,.12)!important;}
[data-testid="stSidebar"] *{color:#d4b8ff!important;}
#MainMenu,footer,header{visibility:hidden;}
[data-testid="chatAvatarIcon-assistant"],[data-testid="chatAvatarIcon-user"]{display:none!important;}
@keyframes glowPulse{0%,100%{box-shadow:0 0 12px rgba(191,95,255,.5),0 0 24px rgba(191,95,255,.2);}50%{box-shadow:0 0 20px rgba(191,95,255,.9),0 0 40px rgba(191,95,255,.4);}}
@keyframes pulse{0%,100%{opacity:1;transform:scale(1);}50%{opacity:.35;transform:scale(.88);}}
@keyframes wave{0%,60%,100%{transform:rotate(0);}10%{transform:rotate(14deg);}20%{transform:rotate(-8deg);}30%{transform:rotate(14deg);}40%{transform:rotate(-4deg);}50%{transform:rotate(10deg);}}
@keyframes fadeIn{from{opacity:0;transform:translateY(14px);}to{opacity:1;transform:translateY(0);}}
@keyframes flow{0%{background-position:0% 50%;}50%{background-position:100% 50%;}100%{background-position:0% 50%;}}
.ai-avatar{width:38px;height:38px;border-radius:10px;object-fit:cover;flex-shrink:0;animation:glowPulse 3s ease-in-out infinite;border:1.5px solid rgba(191,95,255,.6);}
.thinking-logo{width:38px;height:38px;border-radius:10px;animation:pulse 1s ease-in-out infinite,glowPulse 2s ease-in-out infinite;border:1.5px solid rgba(191,95,255,.6);}
.wave-hand{display:inline-block;animation:wave 2s infinite;transform-origin:70% 70%;}
.ai-row{display:flex;align-items:flex-start;gap:12px;margin-bottom:4px;animation:fadeIn .4s ease-out;}
.ai-bubble{font-size:1em;line-height:1.7;font-family:'Rajdhani',sans-serif;letter-spacing:.3px;color:#e8d5ff;
  background:rgba(15,8,35,.85);border:1px solid var(--brd);border-radius:0 16px 16px 16px;padding:12px 16px;
  max-width:860px;backdrop-filter:blur(12px);box-shadow:0 4px 24px rgba(191,95,255,.08),inset 0 1px 0 rgba(255,255,255,.05);
  position:relative;overflow:hidden;}
.ai-bubble::before{content:'';position:absolute;top:0;left:0;right:0;height:1px;
  background:linear-gradient(90deg,transparent,rgba(191,95,255,.6),rgba(0,212,255,.4),transparent);}
.dl-strip{display:flex;align-items:center;gap:6px;padding:6px 0 10px 50px;flex-wrap:wrap;}
.dl-label{font-size:.72em;letter-spacing:2px;color:rgba(191,95,255,.55);font-family:'Orbitron',monospace;white-space:nowrap;}
.nexa-title{font-family:'Orbitron',monospace!important;font-size:3.2em!important;font-weight:900!important;
  letter-spacing:6px;background:linear-gradient(135deg,#bf5fff 0%,#00d4ff 50%,#ff2d9b 100%);
  background-size:200% 200%;-webkit-background-clip:text;-webkit-text-fill-color:transparent;
  background-clip:text;animation:flow 4s ease infinite;margin:0;display:block;}
.nexa-sub{font-family:'Rajdhani',sans-serif;font-size:.85em;letter-spacing:4px;color:rgba(191,95,255,.7);text-transform:uppercase;margin-top:4px;}
.nexa-divider{height:1px;background:linear-gradient(90deg,transparent,rgba(191,95,255,.5),rgba(0,212,255,.3),transparent);margin:12px 0 20px;border:none;}
.plan-card{background:linear-gradient(135deg,rgba(15,8,35,.95) 0%,rgba(8,4,20,.98) 100%);
  border:1px solid rgba(191,95,255,.3);border-radius:16px;padding:20px;margin-top:14px;position:relative;overflow:hidden;}
.plan-card::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;
  background:linear-gradient(90deg,#bf5fff,#00d4ff,#ff2d9b,#bf5fff);background-size:300% 100%;animation:flow 3s linear infinite;}
.plan-label{font-family:'Orbitron',monospace;font-size:.65em;letter-spacing:3px;color:rgba(191,95,255,.8);
  text-transform:uppercase;margin-bottom:12px;display:flex;align-items:center;gap:8px;}
.plan-label::after{content:'';flex:1;height:1px;background:linear-gradient(90deg,rgba(191,95,255,.4),transparent);}
[data-testid="stChatInput"]{background:rgba(15,8,35,.9)!important;border:1px solid rgba(191,95,255,.35)!important;border-radius:14px!important;}
[data-testid="stChatInput"]:focus-within{border-color:rgba(191,95,255,.7)!important;}
[data-testid="stChatInput"] textarea{color:#e8d5ff!important;font-family:'Rajdhani',sans-serif!important;}
.stButton>button{background:linear-gradient(135deg,rgba(191,95,255,.15),rgba(0,212,255,.1))!important;
  border:1px solid rgba(191,95,255,.4)!important;color:#d4b8ff!important;font-family:'Rajdhani',sans-serif!important;
  letter-spacing:1px!important;border-radius:10px!important;transition:all .2s!important;}
.stButton>button:hover{border-color:rgba(191,95,255,.8)!important;box-shadow:0 0 16px rgba(191,95,255,.3)!important;}
[data-testid="stFileUploader"]{background:rgba(15,8,35,.7)!important;border:1px dashed rgba(191,95,255,.4)!important;border-radius:12px!important;}
[data-testid="stFileUploader"] *{color:#d4b8ff!important;}
.stTabs [data-baseweb="tab-list"]{background:rgba(10,4,25,.8)!important;border-radius:10px!important;padding:4px!important;}
.stTabs [data-baseweb="tab"]{color:#a080d0!important;font-family:'Rajdhani',sans-serif!important;}
.stTabs [aria-selected="true"]{color:#bf5fff!important;background:rgba(191,95,255,.15)!important;border-radius:8px!important;}
[data-testid="stDownloadButton"]>button{
  background:linear-gradient(135deg,rgba(191,95,255,.18),rgba(0,212,255,.12))!important;
  border:1px solid rgba(191,95,255,.5)!important;color:#e0c8ff!important;
  font-family:'Rajdhani',sans-serif!important;font-size:.88em!important;letter-spacing:.8px!important;
  border-radius:10px!important;padding:7px 10px!important;transition:all .25s!important;width:100%!important;}
[data-testid="stDownloadButton"]>button:hover{border-color:rgba(191,95,255,.95)!important;box-shadow:0 0 18px rgba(191,95,255,.4)!important;transform:translateY(-1px)!important;}
::-webkit-scrollbar{width:4px;}::-webkit-scrollbar-track{background:var(--db);}::-webkit-scrollbar-thumb{background:rgba(191,95,255,.4);border-radius:2px;}

/* ── EXIT BUTTON (study planner) ── */
.sp-exit-bar{position:relative;z-index:20;display:flex;align-items:center;justify-content:space-between;padding:10px 28px 0 28px;}
.sp-exit-bar-title{font-family:'Orbitron',monospace;font-size:.72em;letter-spacing:3px;color:rgba(191,95,255,.55);text-transform:uppercase;}
.sp-exit-bar .stButton>button{
  background:rgba(255,45,155,.10)!important;border:1px solid rgba(255,45,155,.45)!important;
  color:#ff6bb3!important;font-family:'Orbitron',monospace!important;font-size:.68em!important;
  letter-spacing:2px!important;border-radius:30px!important;padding:6px 20px!important;
  transition:all .22s ease!important;width:auto!important;position:static!important;
  opacity:1!important;height:auto!important;inset:unset!important;z-index:auto!important;}
.sp-exit-bar .stButton>button:hover{background:rgba(255,45,155,.22)!important;border-color:rgba(255,45,155,.85)!important;
  color:#ff2d9b!important;box-shadow:0 0 18px rgba(255,45,155,.35)!important;transform:translateY(-1px)!important;}
</style>
""", unsafe_allow_html=True)

    # ── SIDEBAR (original model4 + Back/Exit buttons) ─────────────────────────
    if logo_b64:
        st.sidebar.markdown(
            f'<img src="data:image/jpeg;base64,{logo_b64}" width="64" '
            'style="border-radius:12px;border:1.5px solid rgba(191,95,255,.5);'
            'box-shadow:0 0 20px rgba(191,95,255,.4);margin-bottom:8px"/>',
            unsafe_allow_html=True)
    st.sidebar.markdown('<h2 style="font-family:Orbitron,monospace;letter-spacing:3px;color:#bf5fff;margin:0">NEXA</h2>', unsafe_allow_html=True)
    st.sidebar.markdown('<p style="font-size:.75em;letter-spacing:2px;color:rgba(191,95,255,.6);margin-top:2px">Powered By NEXA-1.o · Study Planner</p>', unsafe_allow_html=True)
    st.sidebar.markdown("---")

    # ← Back to Home (sidebar)
    if st.sidebar.button("🏠  Back to Home", key="sp_back", use_container_width=True):
        st.session_state.page = "home"
        st.rerun()

    st.sidebar.markdown("---")
    st.sidebar.subheader("🌐 Language")
    lc = st.sidebar.radio("", ["English", "Hindi", "Hinglish"], label_visibility="collapsed", key="sp_lang_radio")
    st.session_state.sp_lang = {"English": "en", "Hindi": "hi", "Hinglish": "hinglish"}[lc]

    st.sidebar.markdown("---")
    st.sidebar.markdown("""
<div style="font-size:.78em;color:rgba(191,95,255,.65);line-height:1.9;letter-spacing:.5px">
📌 <b>How to use</b><br>
<span style="color:#00d4ff">① Tell NEXA your subjects, exam date, daily hours</span><br>
<span style="color:#bf5fff">② Upload a photo / PDF / file of your syllabus</span><br>
<span style="color:#ff2d9b">③ Say how you want the plan (weekly/daily/topic-wise)</span><br>
<span style="color:#ffd700">④ Download plan as PDF · Excel · Word · Calendar<br>
&nbsp;&nbsp;&nbsp;or download any response as TXT · PDF · DOCX</span>
</div>
""", unsafe_allow_html=True)

    st.sidebar.markdown("---")
    with st.sidebar:
        if st.button("🗑️ Clear Chat", use_container_width=True, key="sp_clear"):
            st.session_state.sp_messages  = [{"role": "assistant", "content": "namaste", "msg_id": "init"}]
            st.session_state.sp_last_plan = None
            st.session_state.sp_file_cache = {}
            st.rerun()

    st.sidebar.markdown(
        '<div style="font-size:.72em;text-align:center;color:rgba(191,95,255,.4);margin-top:20px;letter-spacing:1px">NEXA · Powered By NEXA-1.o</div>',
        unsafe_allow_html=True)

    # ── ALL HELPER FUNCTIONS (original model4) ────────────────────────────────

    def sp_file_to_text(uploaded) -> str:
        name = uploaded.name.lower()
        data = uploaded.read()
        if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")):
            b64 = base64.b64encode(data).decode()
            hdrs = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
            payload = {
                "model": "llava-v1.5-7b-4096-preview",
                "messages": [{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                    {"type": "text", "text": "Extract ALL text visible in this image. Include subjects, topics, chapters, dates, schedules — everything. Return plain text only."}
                ]}], "max_tokens": 1500
            }
            try:
                r = requests.post(GROQ_API_URL, json=payload, headers=hdrs, timeout=30)
                if r.status_code == 200:
                    return r.json()["choices"][0]["message"]["content"].strip()
            except:
                pass
            return "[Image uploaded — could not extract text. Please describe the content.]"
        elif name.endswith(".pdf"):
            try:
                import pdfplumber, io as _io
                with pdfplumber.open(_io.BytesIO(data)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            except:
                try:
                    from pypdf import PdfReader
                    import io as _io
                    return "\n".join(p.extract_text() or "" for p in PdfReader(_io.BytesIO(data)).pages)
                except:
                    return "[PDF uploaded — install pdfplumber or pypdf to auto-extract.]"
        elif name.endswith((".txt", ".md", ".csv")):
            return data.decode("utf-8", errors="ignore")
        elif name.endswith((".xlsx", ".xls")):
            try:
                import pandas as pd, io as _io
                return pd.read_excel(_io.BytesIO(data)).to_string()
            except:
                return "[Excel file uploaded — install pandas to auto-extract.]"
        return f"[File '{uploaded.name}' uploaded.]"

    PLAN_KEYWORDS = ["plan","schedule","study","syllabus","exam","prepare","timetable",
                     "weekly","daily","month","topic","chapter","subject","revision",
                     "calendar","goal","target","deadline","hours","per day","learn",
                     "course","semester","test","quiz","mock","practise","practice",
                     "ics","google calendar","outlook","download","export","generate"]

    def sp_is_plan_request(text: str) -> bool:
        t = text.lower()
        # If user explicitly asks for calendar/ics file, always treat as plan request
        if any(kw in t for kw in ["ics", "google calendar", "outlook calendar", ".ics",
                                   "calendar file", "export calendar", "calendar export"]):
            return True
        return sum(1 for kw in PLAN_KEYWORDS if kw in t) >= 2

    def sp_chat_response(user_msg: str, extra_context: str = "") -> str:
        lang = st.session_state.sp_lang
        li = {"hi": "Reply in Hindi only.", "hinglish": "Reply in Hinglish — Roman script Hindi-English mix.", "en": "Reply in English."}.get(lang, "Reply in English.")
        system = (
            "You are NEXA, an intelligent AI study planner powered by NEXA-1.o. "
            "Help students plan their studies effectively. Answer clearly and helpfully. "
            "CRITICAL RULE: NEVER output raw ICS, iCalendar, VCALENDAR, BEGIN:VCALENDAR, "
            "BEGIN:VEVENT, DTSTART, DTEND, RRULE, or any calendar file format text in your reply. "
            "If the user asks for a calendar file, Google Calendar export, or .ics file, simply confirm "
            "that a Google Calendar (.ics) file will be auto-generated and available as a download button "
            "below your response — do NOT write out the ICS content yourself. "
            "Always respond in a friendly, readable, human format only. " + li
        )
        if extra_context:
            system += f"\n\nEXTRACTED CONTENT FROM USER FILE/IMAGE:\n{extra_context[:3000]}"
        msgs = [{"role": "system", "content": system}]
        for m in st.session_state.sp_messages[-10:]:
            if m["content"] not in ("namaste",):
                msgs.append({"role": m["role"], "content": m["content"]})
        msgs.append({"role": "user", "content": user_msg})
        hdrs = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
        try:
            r = requests.post(GROQ_API_URL, json={"model": GROQ_MODEL, "messages": msgs, "max_tokens": 800, "temperature": 0.7}, headers=hdrs, timeout=30)
            if r.status_code == 200: return r.json()["choices"][0]["message"]["content"].strip()
            if r.status_code == 429: return "⏳ Rate limit — please wait a moment."
            return f"❌ Error {r.status_code}: {r.text[:200]}"
        except Exception as e:
            return f"💡 Connection error: {e}"

    PLAN_SCHEMA = """Return ONLY a valid JSON object (no markdown, no backticks):
{"title":"Study Plan Title","student_name":"Student","exam_date":"YYYY-MM-DD or null",
"start_date":"YYYY-MM-DD","total_days":30,"daily_hours":4,
"subjects":[{"name":"Subject","topics":["Topic 1","Topic 2"],"priority":"high|medium|low","hours_allocated":20,"color":"#bf5fff"}],
"weekly_plan":[{"week":1,"focus":"Foundation week","days":[{"day":"Monday","subject":"Math","topics":["Algebra"],"hours":2,"task":"Read + Notes"}]}],
"milestones":[{"date":"YYYY-MM-DD","title":"Milestone","description":"Description"}],
"tips":["Tip 1","Tip 2"]}"""

    def sp_extract_plan_json(user_input: str, extra_context: str = "") -> dict | None:
        context_block = f"\nFILE CONTENT:\n{extra_context[:3000]}" if extra_context else ""
        system = ("You are NEXA study plan extractor. Extract a complete detailed study plan. "
                  "If dates not given, assume start_date is today. Fill reasonable defaults. "
                  "IMPORTANT: Return ONLY a raw JSON object — no prose, no backticks, no markdown, "
                  "no explanation before or after. Start your response with { and end with }." + context_block)
        hdrs = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
        try:
            r = requests.post(GROQ_API_URL, json={
                "model": GROQ_MODEL,
                "messages": [{"role": "system", "content": system},
                             {"role": "user", "content": f"Plan request:\n{user_input}\n\n{PLAN_SCHEMA}"}],
                "max_tokens": 2000, "temperature": 0.3
            }, headers=hdrs, timeout=40)
            if r.status_code != 200: return None
            raw = r.json()["choices"][0]["message"]["content"].strip()
            # Strip markdown code fences (handles ```json, ```JSON, ``` etc.)
            raw = re.sub(r"^```[a-zA-Z]*\s*", "", raw)
            raw = re.sub(r"\s*```\s*$", "", raw)
            raw = raw.strip()
            # If there is prose before/after the JSON object, extract just the object
            if not raw.startswith("{"):
                m = re.search(r"\{.*\}", raw, re.DOTALL)
                if m:
                    raw = m.group(0)
            return json.loads(raw)
        except:
            return None

    SUBJ_COLORS = ["#bf5fff","#00d4ff","#ff2d9b","#ffd700","#7fff00","#ff7f50","#00ffcc","#ff9500"]

    def sp_build_plan_svg(plan: dict) -> str:
        subjects = plan.get("subjects", [])
        weekly   = plan.get("weekly_plan", [])
        W, H = 840, 520
        def esc(s): return str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        def wrap(s, n=20): s=str(s); return s[:n]+("…" if len(s)>n else "")
        parts = [
            f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" style="background:#05030f;border-radius:12px;font-family:Rajdhani,sans-serif">',
            '<defs><filter id="gl"><feGaussianBlur stdDeviation="3" result="b"/><feMerge><feMergeNode in="b"/><feMergeNode in="SourceGraphic"/></feMerge></filter>',
            '<radialGradient id="bg" cx="30%" cy="20%" r="75%"><stop offset="0%" stop-color="#120520"/><stop offset="100%" stop-color="#05030f"/></radialGradient></defs>',
            f'<rect width="{W}" height="{H}" fill="url(#bg)" rx="12"/>',
            f'<line x1="0" y1="0" x2="36" y2="0" stroke="#bf5fff" stroke-width="2.5"/>',
            f'<line x1="0" y1="0" x2="0" y2="36" stroke="#bf5fff" stroke-width="2.5"/>',
            f'<line x1="{W}" y1="0" x2="{W-36}" y2="0" stroke="#00d4ff" stroke-width="2.5"/>',
            f'<line x1="{W}" y1="0" x2="{W}" y2="36" stroke="#00d4ff" stroke-width="2.5"/>',
            f'<line x1="0" y1="{H}" x2="36" y2="{H}" stroke="#ff2d9b" stroke-width="2.5"/>',
            f'<line x1="0" y1="{H}" x2="0" y2="{H-36}" stroke="#ff2d9b" stroke-width="2.5"/>',
            f'<line x1="{W}" y1="{H}" x2="{W-36}" y2="{H}" stroke="#ffd700" stroke-width="2.5"/>',
            f'<line x1="{W}" y1="{H}" x2="{W}" y2="{H-36}" stroke="#ffd700" stroke-width="2.5"/>',
            f'<rect x="0" y="0" width="{W}" height="52" fill="rgba(191,95,255,.08)" rx="12"/>',
            f'<text x="{W//2}" y="32" text-anchor="middle" font-size="15" font-weight="700" fill="#bf5fff" filter="url(#gl)" letter-spacing="3" font-family="Orbitron,monospace">{esc(wrap(plan.get("title","Study Plan"),50))}</text>',
            f'<line x1="40" y1="52" x2="{W-40}" y2="52" stroke="rgba(191,95,255,.3)" stroke-width="1"/>',
            f'<text x="30" y="80" font-size="10" fill="rgba(191,95,255,.7)" letter-spacing="2" font-family="Orbitron,monospace">SUBJECTS</text>',
        ]
        total_hrs = sum(s.get("hours_allocated",0) for s in subjects) or 1
        bar_x, bar_y, bar_w, bar_h, gap = 30, 96, 220, 24, 8
        for i, s in enumerate(subjects[:8]):
            c = s.get("color", SUBJ_COLORS[i % len(SUBJ_COLORS)])
            hrs = s.get("hours_allocated", 0)
            fill_w = int(bar_w * hrs / total_hrs)
            y = bar_y + i*(bar_h+gap)
            parts += [
                f'<rect x="{bar_x}" y="{y}" width="{bar_w}" height="{bar_h}" rx="6" fill="rgba(255,255,255,.04)" stroke="rgba(255,255,255,.08)" stroke-width="1"/>',
                f'<rect x="{bar_x}" y="{y}" width="{max(fill_w,6)}" height="{bar_h}" rx="6" fill="{c}" opacity=".75" filter="url(#gl)"/>',
                f'<text x="{bar_x+8}" y="{y+bar_h//2+1}" dominant-baseline="middle" font-size="11" fill="white" font-weight="600">{esc(wrap(s.get("name",""),16))}</text>',
                f'<text x="{bar_x+bar_w+6}" y="{y+bar_h//2+1}" dominant-baseline="middle" font-size="10" fill="{c}">{hrs}h</text>',
            ]
        gx, gy = 290, 62
        gw = W - gx - 20
        parts.append(f'<text x="{gx}" y="80" font-size="10" fill="rgba(0,212,255,.7)" letter-spacing="2" font-family="Orbitron,monospace">WEEKLY PLAN</text>')
        week_w = gw // max(len(weekly), 1)
        for wi, week in enumerate(weekly[:6]):
            wx = gx + wi * week_w
            wc = SUBJ_COLORS[wi % len(SUBJ_COLORS)]
            parts += [
                f'<rect x="{wx+4}" y="{gy+20}" width="{week_w-8}" height="420" rx="8" fill="rgba(255,255,255,.03)" stroke="{wc}" stroke-width="1" stroke-opacity=".4"/>',
                f'<text x="{wx+week_w//2}" y="{gy+34}" text-anchor="middle" font-size="10" fill="{wc}" font-weight="700">Wk{week.get("week","")}</text>',
                f'<text x="{wx+week_w//2}" y="{gy+48}" text-anchor="middle" font-size="8" fill="rgba(255,255,255,.4)">{esc(wrap(week.get("focus",""),12))}</text>',
            ]
            for di, day in enumerate(week.get("days",[])[:7]):
                dy = gy + 58 + di * 50
                dc = next((s.get("color",SUBJ_COLORS[0]) for s in subjects if s.get("name","").lower() in day.get("subject","").lower()), SUBJ_COLORS[di%len(SUBJ_COLORS)])
                parts += [
                    f'<rect x="{wx+8}" y="{dy}" width="{week_w-16}" height="44" rx="6" fill="{dc}" opacity=".12" stroke="{dc}" stroke-width=".8" stroke-opacity=".5"/>',
                    f'<text x="{wx+12}" y="{dy+13}" font-size="9" fill="{dc}" font-weight="700">{esc(day.get("day","")[:3])}</text>',
                    f'<text x="{wx+12}" y="{dy+26}" font-size="9" fill="rgba(255,255,255,.8)">{esc(wrap(day.get("subject",""),14))}</text>',
                    f'<text x="{wx+12}" y="{dy+38}" font-size="8" fill="rgba(255,255,255,.5)">{esc(wrap(str(day.get("topics",[""])[0]),16))} · {day.get("hours",0)}h</text>',
                ]
        milestones = plan.get("milestones",[])
        if milestones:
            my = H - 38
            parts += [
                f'<line x1="40" y1="{my-14}" x2="{W-40}" y2="{my-14}" stroke="rgba(191,95,255,.2)" stroke-width="1"/>',
                f'<text x="30" y="{my-2}" font-size="9" fill="rgba(191,95,255,.6)" letter-spacing="2" font-family="Orbitron,monospace">MILESTONES</text>',
            ]
            mstep = (W-80)//max(len(milestones),1)
            for mi, ms in enumerate(milestones[:8]):
                mx = 40+mi*mstep+mstep//2
                mc = SUBJ_COLORS[mi%len(SUBJ_COLORS)]
                parts += [
                    f'<circle cx="{mx}" cy="{my+12}" r="5" fill="{mc}" filter="url(#gl)"/>',
                    f'<text x="{mx}" y="{my+26}" text-anchor="middle" font-size="8" fill="{mc}">{esc(wrap(ms.get("title",""),14))}</text>',
                ]
        parts.append("</svg>")
        return "\n".join(parts)

    def sp_generate_ics(plan: dict) -> bytes:
        # RFC 5545 requires CRLF line endings for .ics files
        CRLF = "\r\n"
        lines = ["BEGIN:VCALENDAR","VERSION:2.0","PRODID:-//NEXA Study Planner//EN",
                 "CALSCALE:GREGORIAN","METHOD:PUBLISH","X-WR-CALNAME:NEXA Study Plan"]
        try:
            start = datetime.strptime(str(plan.get("start_date",""))[:10], "%Y-%m-%d")
        except:
            start = datetime.today()
        day_map = {"monday":0,"tuesday":1,"wednesday":2,"thursday":3,"friday":4,"saturday":5,"sunday":6}
        uid_n = [0]
        def uid():
            uid_n[0]+=1; return f"nexa-{uid_n[0]}-{uuid.uuid4().hex[:6]}@nexaplanner"
        for week in plan.get("weekly_plan",[]):
            ws = start + timedelta(weeks=int(week.get("week",1))-1)
            for day in week.get("days",[]):
                ed = ws + timedelta(days=day_map.get(day.get("day","monday").lower(), 0))
                try:
                    hrs = float(day.get("hours", 1))
                except:
                    hrs = 1.0
                start_h = 9
                end_total = start_h + hrs
                end_h = int(end_total)
                end_m = int(round((end_total - end_h) * 60))
                if end_m == 60:
                    end_h += 1; end_m = 0
                # Cap at 23:59 to avoid spilling into next day
                if end_h >= 24:
                    end_h = 23; end_m = 59
                date_str = ed.strftime('%Y%m%d')
                desc = ', '.join(day.get('topics', [])) or day.get('task', '')
                lines += [
                    "BEGIN:VEVENT",
                    f"UID:{uid()}",
                    f"DTSTART:{date_str}T{start_h:02d}0000",
                    f"DTEND:{date_str}T{end_h:02d}{end_m:02d}00",
                    f"SUMMARY:Study {day.get('subject','')} - {day.get('task','')}",
                    f"DESCRIPTION:{desc}",
                    "END:VEVENT"
                ]
        for ms in plan.get("milestones", []):
            try:
                md = datetime.strptime(str(ms.get("date",""))[:10], "%Y-%m-%d").strftime("%Y%m%d")
                lines += [
                    "BEGIN:VEVENT",
                    f"UID:{uid()}",
                    f"DTSTART;VALUE=DATE:{md}",
                    f"DTEND;VALUE=DATE:{md}",
                    f"SUMMARY:🎯 MILESTONE: {ms.get('title','')}",
                    f"DESCRIPTION:{ms.get('description','')}",
                    "END:VEVENT"
                ]
            except:
                pass
        lines.append("END:VCALENDAR")
        return CRLF.join(lines).encode("utf-8")

    def sp_generate_excel(plan: dict) -> bytes:
        if not OPENPYXL_OK: return b""
        try:
            wb = openpyxl.Workbook()
            def tb():
                s = Side(style="thin", color="3D1A6B")
                return Border(left=s,right=s,top=s,bottom=s)
            PURPLE,BLUE,DARK,MID,WHITE,GOLD,PINK = "BF5FFF","00D4FF","0D0520","1A0A35","E8D5FF","FFD700","FF2D9B"
            ws1 = wb.active; ws1.title = "Overview"; ws1.sheet_view.showGridLines = False
            ws1.merge_cells("A1:H1"); ws1["A1"] = plan.get("title","Study Plan")
            ws1["A1"].font=Font(name="Arial",size=18,bold=True,color=PURPLE)
            ws1["A1"].fill=PatternFill("solid",start_color=DARK,end_color=DARK)
            ws1["A1"].alignment=Alignment(horizontal="center",vertical="center")
            ws1.row_dimensions[1].height=40
            for i,(k,v) in enumerate([("Student",plan.get("student_name","Student")),("Start Date",plan.get("start_date","")),
                                        ("Exam Date",plan.get("exam_date","TBD")),("Total Days",plan.get("total_days","")),("Daily Hours",plan.get("daily_hours",""))],3):
                ws1[f"A{i}"].value=k; ws1[f"A{i}"].font=Font(name="Arial",size=11,bold=True,color=BLUE)
                ws1[f"A{i}"].fill=PatternFill("solid",start_color=MID,end_color=MID)
                ws1[f"B{i}"].value=str(v); ws1[f"B{i}"].font=Font(name="Arial",size=11,color=WHITE)
                ws1[f"B{i}"].fill=PatternFill("solid",start_color=DARK,end_color=DARK)
            ws1["A9"].value="SUBJECTS"; ws1["A9"].font=Font(name="Arial",size=12,bold=True,color=GOLD)
            ws1["A9"].fill=PatternFill("solid",start_color=MID,end_color=MID)
            for ci,h in enumerate(["Subject","Priority","Hours","Topics","Color"],1):
                c=ws1.cell(10,ci,h); c.font=Font(name="Arial",size=10,bold=True,color=DARK)
                c.fill=PatternFill("solid",start_color=PURPLE,end_color=PURPLE)
                c.alignment=Alignment(horizontal="center"); c.border=tb()
            for ri,s in enumerate(plan.get("subjects",[]),11):
                for ci,v in enumerate([s.get("name",""),s.get("priority","").upper(),s.get("hours_allocated",""),", ".join(s.get("topics",[])[:5]),s.get("color","")],1):
                    cell=ws1.cell(ri,ci,v); cell.font=Font(name="Arial",size=10,color=WHITE)
                    cell.fill=PatternFill("solid",start_color="0A0418",end_color="0A0418")
                    cell.border=tb(); cell.alignment=Alignment(wrap_text=True)
                ws1.cell(ri,5,"").fill=PatternFill("solid",start_color=s.get("color","BF5FFF").lstrip("#"),end_color=s.get("color","BF5FFF").lstrip("#"))
            tr=11+len(plan.get("subjects",[]))+2
            ws1.cell(tr,1,"STUDY TIPS").font=Font(name="Arial",size=12,bold=True,color=GOLD)
            ws1.cell(tr,1).fill=PatternFill("solid",start_color=MID,end_color=MID)
            for ti,tip in enumerate(plan.get("tips",[]),tr+1):
                ws1.cell(ti,1,f"• {tip}").font=Font(name="Arial",size=10,color=WHITE)
                ws1.cell(ti,1).fill=PatternFill("solid",start_color=DARK,end_color=DARK)
                ws1.merge_cells(f"A{ti}:H{ti}")
            for c,w in zip("ABCDEFGH",[28,14,12,50,10,10,10,10]): ws1.column_dimensions[c].width=w
            ws2=wb.create_sheet("Weekly Plan"); ws2.sheet_view.showGridLines=False
            ws2.merge_cells("A1:G1"); ws2["A1"]="WEEKLY STUDY PLAN"
            ws2["A1"].font=Font(name="Arial",size=16,bold=True,color=BLUE)
            ws2["A1"].fill=PatternFill("solid",start_color=DARK,end_color=DARK)
            ws2["A1"].alignment=Alignment(horizontal="center",vertical="center"); ws2.row_dimensions[1].height=36
            for ci,h in enumerate(["Week","Day","Subject","Topics","Hours","Task","Notes"],1):
                c=ws2.cell(2,ci,h); c.font=Font(name="Arial",size=10,bold=True,color=DARK)
                c.fill=PatternFill("solid",start_color=BLUE,end_color=BLUE)
                c.alignment=Alignment(horizontal="center"); c.border=tb()
            scm={s.get("name","").lower():s.get("color","BF5FFF").lstrip("#") for s in plan.get("subjects",[])}
            row=3
            for week in plan.get("weekly_plan",[]):
                for day in week.get("days",[]):
                    sn=day.get("subject",""); hx=scm.get(sn.lower(),"1A0A35")
                    for ci,v in enumerate([week.get("week",""),day.get("day",""),sn,", ".join(day.get("topics",[])),day.get("hours",""),day.get("task",""),""],1):
                        cell=ws2.cell(row,ci,v); cell.font=Font(name="Arial",size=10,color=WHITE)
                        cell.fill=PatternFill("solid",start_color="0A0418",end_color="0A0418")
                        cell.border=tb(); cell.alignment=Alignment(wrap_text=True)
                    ws2.cell(row,3).fill=PatternFill("solid",start_color=hx,end_color=hx)
                    ws2.cell(row,3).font=Font(name="Arial",size=10,color=DARK,bold=True); row+=1
            for c,w in zip("ABCDEFG",[8,14,22,45,8,24,20]): ws2.column_dimensions[c].width=w
            ws3=wb.create_sheet("Milestones"); ws3.sheet_view.showGridLines=False
            ws3.merge_cells("A1:D1"); ws3["A1"]="MILESTONES & TARGETS"
            ws3["A1"].font=Font(name="Arial",size=16,bold=True,color=PINK)
            ws3["A1"].fill=PatternFill("solid",start_color=DARK,end_color=DARK)
            ws3["A1"].alignment=Alignment(horizontal="center",vertical="center"); ws3.row_dimensions[1].height=36
            for ci,h in enumerate(["Date","Milestone","Description","Status"],1):
                c=ws3.cell(2,ci,h); c.font=Font(name="Arial",size=10,bold=True,color=DARK)
                c.fill=PatternFill("solid",start_color=PINK,end_color=PINK)
                c.alignment=Alignment(horizontal="center"); c.border=tb()
            for ri,ms in enumerate(plan.get("milestones",[]),3):
                for ci,v in enumerate([ms.get("date",""),ms.get("title",""),ms.get("description",""),"Pending"],1):
                    cell=ws3.cell(ri,ci,v); cell.font=Font(name="Arial",size=10,color=WHITE)
                    cell.fill=PatternFill("solid",start_color="0A0418",end_color="0A0418"); cell.border=tb()
            for c,w in zip("ABCD",[14,30,50,14]): ws3.column_dimensions[c].width=w
            buf=BytesIO(); wb.save(buf); buf.seek(0); return buf.read()
        except Exception as e:
            st.error(f"Excel generation failed: {e}"); return b""

    def sp_generate_pdf(plan: dict) -> bytes:
        if not REPORTLAB_OK: return b""
        try:
            buf=BytesIO()
            doc=SimpleDocTemplate(buf,pagesize=A4,leftMargin=1.5*cm,rightMargin=1.5*cm,topMargin=2*cm,bottomMargin=2*cm)
            styles=getSampleStyleSheet()
            PH="#BF5FFF"; BH="#00D4FF"; DH="#0D0520"; LH="#E8D5FF"; GH="#FFD700"
            ts=ParagraphStyle("NT",parent=styles["Title"],fontSize=22,textColor=rl_colors.HexColor(PH),spaceAfter=6,alignment=TA_CENTER,fontName="Helvetica-Bold")
            h1=ParagraphStyle("H1",parent=styles["Heading1"],fontSize=13,textColor=rl_colors.HexColor(BH),spaceBefore=12,spaceAfter=4,fontName="Helvetica-Bold")
            h2=ParagraphStyle("H2",parent=styles["Heading2"],fontSize=11,textColor=rl_colors.HexColor(GH),spaceBefore=8,spaceAfter=3,fontName="Helvetica-Bold")
            bs=ParagraphStyle("NB",parent=styles["Normal"],fontSize=10,textColor=rl_colors.HexColor(LH),spaceAfter=3,fontName="Helvetica")
            def tbl_s(hc=PH):
                return TableStyle([
                    ("BACKGROUND",(0,0),(-1,0),rl_colors.HexColor(hc)),("TEXTCOLOR",(0,0),(-1,0),rl_colors.HexColor(DH)),
                    ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,0),9),
                    ("BACKGROUND",(0,1),(-1,-1),rl_colors.HexColor("#0A0418")),("TEXTCOLOR",(0,1),(-1,-1),rl_colors.HexColor(LH)),
                    ("FONTNAME",(0,1),(-1,-1),"Helvetica"),("FONTSIZE",(0,1),(-1,-1),9),
                    ("GRID",(0,0),(-1,-1),0.5,rl_colors.HexColor("#3D1A6B")),
                    ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.HexColor("#0A0418"),rl_colors.HexColor("#0F0625")]),
                    ("VALIGN",(0,0),(-1,-1),"MIDDLE"),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)])
            story=[Paragraph(plan.get("title","Study Plan"),ts),
                   Paragraph(f"Powered By NEXA-1.o · {datetime.today().strftime('%B %d, %Y')}",
                             ParagraphStyle("sub",parent=styles["Normal"],fontSize=9,textColor=rl_colors.HexColor("#9060CC"),alignment=TA_CENTER)),
                   Spacer(1,0.4*cm),HRFlowable(width="100%",thickness=1,color=rl_colors.HexColor(PH)),Spacer(1,0.3*cm)]
            story+=[Paragraph("Plan Overview",h1),
                    Table([["Field","Details"],["Student",plan.get("student_name","")],["Start Date",str(plan.get("start_date",""))],
                           ["Exam Date",str(plan.get("exam_date","TBD"))],["Total Days",str(plan.get("total_days",""))],
                           ["Daily Hours",str(plan.get("daily_hours",""))]],colWidths=[5*cm,10*cm])]
            story[-1].setStyle(tbl_s()); story.append(Spacer(1,0.4*cm))
            sd=[["Subject","Priority","Hours","Key Topics"]]+[[s.get("name",""),s.get("priority","").upper(),str(s.get("hours_allocated","")),", ".join(s.get("topics",[])[:4])] for s in plan.get("subjects",[])]
            story+=[Paragraph("Subjects",h1),Table(sd,colWidths=[4.5*cm,3*cm,2.5*cm,9*cm])]
            story[-1].setStyle(tbl_s(BH)); story.append(Spacer(1,0.4*cm))
            story.append(Paragraph("Weekly Plan",h1))
            for week in plan.get("weekly_plan",[]):
                story.append(Paragraph(f"Week {week.get('week','')} — {week.get('focus','')}",h2))
                wd=[["Day","Subject","Topics","Hrs","Task"]]+[[d.get("day",""),d.get("subject",""),", ".join(d.get("topics",[])),str(d.get("hours","")),d.get("task","")] for d in week.get("days",[])]
                story.append(Table(wd,colWidths=[2.5*cm,3.5*cm,7*cm,2*cm,4*cm])); story[-1].setStyle(tbl_s()); story.append(Spacer(1,0.2*cm))
            if plan.get("milestones"):
                story+=[Spacer(1,0.3*cm),Paragraph("Milestones",h1),
                        Table([["Date","Milestone","Description"]]+[[m.get("date",""),m.get("title",""),m.get("description","")] for m in plan.get("milestones",[])],colWidths=[3*cm,5*cm,11*cm])]
                story[-1].setStyle(tbl_s("#FF2D9B"))
            if plan.get("tips"):
                story+=[Spacer(1,0.4*cm),Paragraph("Study Tips",h1)]+[Paragraph(f"• {t}",bs) for t in plan.get("tips",[])]
            doc.build(story); buf.seek(0); return buf.read()
        except Exception as e:
            st.error(f"PDF generation failed: {e}"); return b""

    def _hex_rgb(h):
        h=h.lstrip("#")
        return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)) if len(h)==6 else RGBColor(0x99,0x60,0xFF)

    def _cell_bg(cell, hex_color):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"),hex_color.lstrip("#")); tcPr.append(shd)

    def _make_docx_table(doc, headers, rows, hbg="#6B1FBF", col_cm=None):
        tbl=doc.add_table(rows=1+len(rows),cols=len(headers)); tbl.style="Table Grid"
        hr=tbl.rows[0]
        for ci,h in enumerate(headers):
            cell=hr.cells[ci]; _cell_bg(cell,hbg)
            p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            run=p.add_run(h); run.bold=True; run.font.size=Pt(9); run.font.color.rgb=RGBColor(0x0D,0x05,0x20)
            if col_cm: cell.width=Cm(col_cm[ci])
        for ri,rd in enumerate(rows):
            ro=tbl.rows[ri+1]; bg="0A0418" if ri%2==0 else "0F0625"
            for ci,v in enumerate(rd):
                cell=ro.cells[ci]; _cell_bg(cell,bg)
                run=cell.paragraphs[0].add_run(str(v)); run.font.size=Pt(9); run.font.color.rgb=_hex_rgb("#E8D5FF")
                if col_cm: cell.width=Cm(col_cm[ci])

    def sp_generate_docx(plan: dict) -> bytes:
        if not DOCX_OK: return b""
        try:
            doc=DocxDocument()
            for sec in doc.sections:
                sec.top_margin=Cm(2); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2); sec.right_margin=Cm(2)
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(plan.get("title","NEXA Study Plan")); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=_hex_rgb("#6B1FBF")
            p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r2=p2.add_run(f"Powered By NEXA-1.o  ·  {datetime.today().strftime('%B %d, %Y')}"); r2.italic=True; r2.font.size=Pt(10); r2.font.color.rgb=_hex_rgb("#9060CC")
            pr=doc.add_paragraph()._p.get_or_add_pPr(); pb=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
            bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:space"),"1"); bt.set(qn("w:color"),"BF5FFF"); pb.append(bt); pr.append(pb)
            doc.add_paragraph()
            oh=doc.add_paragraph(); or_=oh.add_run("Plan Overview"); or_.bold=True; or_.font.size=Pt(14); or_.font.color.rgb=_hex_rgb("#0060A0")
            _make_docx_table(doc,["Field","Details"],[["Student",plan.get("student_name","")],["Start Date",str(plan.get("start_date",""))],["Exam Date",str(plan.get("exam_date","TBD"))],["Total Days",str(plan.get("total_days",""))],["Daily Hours",str(plan.get("daily_hours",""))]],col_cm=[5,10])
            doc.add_paragraph()
            sh=doc.add_paragraph(); sr=sh.add_run("Subjects & Allocation"); sr.bold=True; sr.font.size=Pt(14); sr.font.color.rgb=_hex_rgb("#0060A0")
            _make_docx_table(doc,["Subject","Priority","Hours","Key Topics"],[[s.get("name",""),s.get("priority","").upper(),str(s.get("hours_allocated","")),", ".join(s.get("topics",[])[:5])] for s in plan.get("subjects",[])],hbg="#007080",col_cm=[4,3,2.5,9])
            doc.add_paragraph()
            wh=doc.add_paragraph(); wr=wh.add_run("Weekly Study Plan"); wr.bold=True; wr.font.size=Pt(14); wr.font.color.rgb=_hex_rgb("#0060A0")
            for week in plan.get("weekly_plan",[]):
                wsh=doc.add_paragraph(); wsr=wsh.add_run(f"Week {week.get('week','')}  —  {week.get('focus','')}"); wsr.bold=True; wsr.font.size=Pt(11); wsr.font.color.rgb=_hex_rgb("#BF8F00")
                _make_docx_table(doc,["Day","Subject","Topics","Hrs","Task"],[[d.get("day",""),d.get("subject",""),", ".join(d.get("topics",[])),str(d.get("hours","")),d.get("task","")] for d in week.get("days",[])],col_cm=[2.5,3.5,7,1.5,4])
                doc.add_paragraph()
            if plan.get("milestones"):
                mh=doc.add_paragraph(); mr=mh.add_run("Milestones"); mr.bold=True; mr.font.size=Pt(14); mr.font.color.rgb=_hex_rgb("#0060A0")
                _make_docx_table(doc,["Date","Milestone","Description","Status"],[[m.get("date",""),m.get("title",""),m.get("description",""),"Pending"] for m in plan.get("milestones",[])],hbg="#A0005F",col_cm=[3,5,9,2.5])
                doc.add_paragraph()
            if plan.get("tips"):
                th=doc.add_paragraph(); tr2=th.add_run("Study Tips"); tr2.bold=True; tr2.font.size=Pt(14); tr2.font.color.rgb=_hex_rgb("#0060A0")
                for tip in plan.get("tips",[]):
                    tp=doc.add_paragraph(style="List Bullet"); tr3=tp.add_run(tip); tr3.font.size=Pt(10); tr3.font.color.rgb=_hex_rgb("#2C1A4A")
            buf=BytesIO(); doc.save(buf); buf.seek(0); return buf.read()
        except Exception as e:
            st.error(f"DOCX generation failed: {e}"); return b""

    def sp_response_to_txt(content: str) -> bytes:
        header = f"NEXA Study Planner Response\nGenerated: {datetime.today().strftime('%Y-%m-%d %H:%M')}\n{'='*50}\n\n"
        return (header + content).encode("utf-8")

    def sp_response_to_pdf(content: str) -> bytes:
        if not REPORTLAB_OK: return b""
        try:
            buf=BytesIO()
            doc=SimpleDocTemplate(buf,pagesize=A4,leftMargin=1.5*cm,rightMargin=1.5*cm,topMargin=2*cm,bottomMargin=2*cm)
            styles=getSampleStyleSheet()
            title_s=ParagraphStyle("T",parent=styles["Title"],fontSize=18,textColor=rl_colors.HexColor("#BF5FFF"),spaceAfter=6,alignment=TA_CENTER,fontName="Helvetica-Bold")
            body_s=ParagraphStyle("B",parent=styles["Normal"],fontSize=11,textColor=rl_colors.HexColor("#1A1A2E"),spaceAfter=6,fontName="Helvetica",leading=16)
            story=[Paragraph("NEXA Response",title_s),
                   Paragraph(f"Generated {datetime.today().strftime('%B %d, %Y')}",
                             ParagraphStyle("s",parent=styles["Normal"],fontSize=9,textColor=rl_colors.HexColor("#9060CC"),alignment=TA_CENTER)),
                   Spacer(1,0.5*cm),HRFlowable(width="100%",thickness=1,color=rl_colors.HexColor("#BF5FFF")),Spacer(1,0.4*cm)]
            for line in content.split("\n"):
                line=line.strip()
                if not line: story.append(Spacer(1,0.2*cm)); continue
                story.append(Paragraph(line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"),body_s))
            doc.build(story); buf.seek(0); return buf.read()
        except: return b""

    def sp_response_to_docx(content: str) -> bytes:
        if not DOCX_OK: return b""
        try:
            doc=DocxDocument()
            for sec in doc.sections: sec.top_margin=Cm(2);sec.bottom_margin=Cm(2);sec.left_margin=Cm(2.5);sec.right_margin=Cm(2.5)
            tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            tr=tp.add_run("NEXA Study Planner"); tr.bold=True; tr.font.size=Pt(20); tr.font.color.rgb=_hex_rgb("#6B1FBF")
            dp=doc.add_paragraph(); dp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            dr=dp.add_run(f"Generated: {datetime.today().strftime('%B %d, %Y')}"); dr.italic=True; dr.font.size=Pt(10); dr.font.color.rgb=_hex_rgb("#9060CC")
            doc.add_paragraph()
            for line in content.split("\n"):
                p=doc.add_paragraph()
                if line.startswith("##"):
                    r=p.add_run(line.lstrip("#").strip()); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=_hex_rgb("#0060A0")
                elif line.startswith("#"):
                    r=p.add_run(line.lstrip("#").strip()); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=_hex_rgb("#6B1FBF")
                elif line.startswith("- ") or line.startswith("• "):
                    p.style="List Bullet"; r=p.add_run(line[2:].strip()); r.font.size=Pt(11); r.font.color.rgb=_hex_rgb("#1A1A2E")
                elif line.strip():
                    r=p.add_run(line); r.font.size=Pt(11); r.font.color.rgb=_hex_rgb("#1A1A2E")
            buf=BytesIO(); doc.save(buf); buf.seek(0); return buf.read()
        except: return b""

    def sp_render_assistant(content: str, plan=None, msg_idx: int = 0):
        if logo_b64:
            img_tag = f'<img class="ai-avatar" src="data:image/jpeg;base64,{logo_b64}"/>'
        else:
            img_tag = '<div style="width:38px;height:38px;border-radius:10px;flex-shrink:0;background:linear-gradient(135deg,#bf5fff,#00d4ff);animation:glowPulse 3s ease-in-out infinite"></div>'

        if content == "namaste":
            bubble = ('Namaste <span class="wave-hand">🙏</span>&nbsp; I\'m <b style="color:#bf5fff">NEXA</b> '
                      'powered by <b style="color:#00d4ff">NEXA-1.o</b> — your personal AI study planner.<br><br>'
                      '<span style="color:rgba(191,95,255,.8);font-size:.93em">'
                      '✦ Tell me your subjects, exam date &amp; daily study hours<br>'
                      '✦ Upload a syllabus photo, PDF, or timetable<br>'
                      '✦ Get a full plan · Download as <b>PDF · Excel · Word · Calendar</b><br>'
                      '✦ Download <b>any response</b> as TXT · PDF · DOCX'
                      '</span>')
        else:
            bubble = content

        st.markdown(f'<div class="ai-row">{img_tag}<div class="ai-bubble">{bubble}</div></div>', unsafe_allow_html=True)

        if content not in ("namaste",) and content:
            fname = f"sp_response_{msg_idx}"
            st.markdown('<div class="dl-strip"><span class="dl-label">↓ DOWNLOAD RESPONSE</span>', unsafe_allow_html=True)
            rc1, rc2, rc3 = st.columns([1, 1, 1])
            with rc1:
                st.download_button(label="📄 TXT", data=sp_response_to_txt(content),
                    file_name=f"{fname}.txt", mime="text/plain",
                    key=f"sp_dl_txt_{msg_idx}", use_container_width=True)
            with rc2:
                if REPORTLAB_OK:
                    pdf_bytes = sp_response_to_pdf(content)
                    st.download_button(label="📑 PDF", data=pdf_bytes if pdf_bytes else b" ",
                        file_name=f"{fname}.pdf", mime="application/pdf",
                        key=f"sp_dl_rpdf_{msg_idx}", use_container_width=True, disabled=(not pdf_bytes))
                else:
                    st.caption("pip install reportlab")
            with rc3:
                if DOCX_OK:
                    docx_bytes = sp_response_to_docx(content)
                    st.download_button(label="📝 DOCX", data=docx_bytes if docx_bytes else b" ",
                        file_name=f"{fname}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"sp_dl_rdocx_{msg_idx}", use_container_width=True, disabled=(not docx_bytes))
                else:
                    st.caption("pip install python-docx")
            st.markdown('</div>', unsafe_allow_html=True)

        if plan:
            svg = sp_build_plan_svg(plan)
            st.markdown(
                f'<div class="plan-card">'
                f'<div class="plan-label">⬡ NEXA VISUAL STUDY PLAN · {plan.get("title","")[:40]}</div>'
                f'{svg}</div>',
                unsafe_allow_html=True)
            st.markdown('<div style="height:10px"></div>', unsafe_allow_html=True)

            cache_key = f"sp_plan_{msg_idx}"
            if cache_key not in st.session_state.sp_file_cache:
                with st.spinner("Preparing download files…"):
                    try:
                        ics_bytes = sp_generate_ics(plan)
                    except Exception:
                        ics_bytes = b""
                    try:
                        xlsx_bytes = sp_generate_excel(plan)
                    except Exception:
                        xlsx_bytes = b""
                    try:
                        pdf_bytes_plan = sp_generate_pdf(plan)
                    except Exception:
                        pdf_bytes_plan = b""
                    try:
                        docx_bytes_plan = sp_generate_docx(plan)
                    except Exception:
                        docx_bytes_plan = b""
                    st.session_state.sp_file_cache[cache_key] = {
                        "ics":  ics_bytes,
                        "xlsx": xlsx_bytes,
                        "pdf":  pdf_bytes_plan,
                        "docx": docx_bytes_plan,
                    }
            cached = st.session_state.sp_file_cache[cache_key]
            fname_base = plan.get("student_name", "nexa").replace(" ", "_")

            pc1, pc2, pc3, pc4 = st.columns(4)
            with pc1:
                _ics_data = cached["ics"] if cached["ics"] else (
                    "BEGIN:VCALENDAR\r\nVERSION:2.0\r\nPRODID:-//NEXA//EN\r\n"
                    "CALSCALE:GREGORIAN\r\nMETHOD:PUBLISH\r\nEND:VCALENDAR\r\n"
                ).encode("utf-8")
                st.download_button(label="📅 Google Calendar (.ics)", data=_ics_data,
                    file_name=f"study_plan_{fname_base}.ics", mime="text/calendar",
                    key=f"sp_dl_ics_{msg_idx}", use_container_width=True,
                    help="Import into Google / Apple / Outlook Calendar")
            with pc2:
                st.download_button(label="📊 Excel (.xlsx)", data=cached["xlsx"] or b" ",
                    file_name=f"study_plan_{fname_base}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key=f"sp_dl_xlsx_{msg_idx}", use_container_width=True,
                    disabled=(not cached["xlsx"]), help="3-sheet Excel: Overview · Weekly Plan · Milestones")
            with pc3:
                st.download_button(label="📄 PDF", data=cached["pdf"] or b" ",
                    file_name=f"study_plan_{fname_base}.pdf", mime="application/pdf",
                    key=f"sp_dl_pdf_{msg_idx}", use_container_width=True,
                    disabled=(not cached["pdf"]), help="Printable PDF study plan")
            with pc4:
                st.download_button(label="📝 Word (.docx)", data=cached["docx"] or b" ",
                    file_name=f"study_plan_{fname_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"sp_dl_docx_{msg_idx}", use_container_width=True,
                    disabled=(not cached["docx"]), help="Editable Word document")

            missing = []
            if not OPENPYXL_OK: missing.append("`pip install openpyxl` for Excel")
            if not REPORTLAB_OK: missing.append("`pip install reportlab` for PDF")
            if not DOCX_OK:      missing.append("`pip install python-docx` for Word")
            if missing:
                st.info("Install to enable more exports:  " + "  ·  ".join(missing))

    # ── EXIT BUTTON BAR ───────────────────────────────────────────────────────
    st.markdown('<div class="sp-exit-bar">', unsafe_allow_html=True)
    _sp1, _sp2 = st.columns([6, 1])
    with _sp1:
        st.markdown(
            '<span class="sp-exit-bar-title">📚 NEXA STUDY PLANNER &nbsp;·&nbsp; AI Powered</span>',
            unsafe_allow_html=True,
        )
    with _sp2:
        if st.button("✕  Exit", key="sp_exit_top", help="Return to NEXA Home"):
            st.session_state.page = "home"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    # ── HEADER (original model4) ──────────────────────────────────────────────
    st.markdown(
        '<div style="text-align:center;padding:28px 0 10px">'
        '<span class="nexa-title">NEXA</span>'
        '<div class="nexa-sub">Powered By NEXA-1.o · AI Study Planner</div>'
        '</div>', unsafe_allow_html=True)
    st.markdown('<hr class="nexa-divider">', unsafe_allow_html=True)

    # ── CHAT HISTORY ─────────────────────────────────────────────────────────
    for idx, msg in enumerate(st.session_state.sp_messages):
        if msg["role"] == "assistant":
            sp_render_assistant(msg["content"], plan=msg.get("plan"), msg_idx=idx)
        else:
            with st.chat_message("user"):
                st.markdown(
                    f'<span style="color:#e8d5ff;font-family:Rajdhani,sans-serif">{msg["content"]}</span>',
                    unsafe_allow_html=True)
            if msg.get("file_preview"):
                with st.chat_message("user"):
                    st.caption(f"📎 {msg['file_preview']}")

    # ── INPUT AREA ────────────────────────────────────────────────────────────
    uploaded = st.file_uploader(
        "📎 Upload syllabus / timetable / photo / PDF / Excel",
        type=["png","jpg","jpeg","webp","pdf","txt","xlsx","csv","md"],
        label_visibility="collapsed",
        key="sp_uploader"
    )
    prompt = st.chat_input("✦ Tell NEXA your study plan details, subjects, exam date, hours per day…")

    # ── PROCESS INPUT ────────────────────────────────────────────────────────
    if prompt:
        extracted_text = ""
        if uploaded:
            with st.spinner("🔍 Extracting content from file…"):
                extracted_text = sp_file_to_text(uploaded)

        st.session_state.sp_messages.append({
            "role": "user",
            "content": prompt,
            "msg_id": str(uuid.uuid4()),
            "file_preview": f"{uploaded.name} — {extracted_text[:120]}…" if uploaded else ""
        })

        with st.chat_message("user"):
            st.markdown(f'<span style="color:#e8d5ff;font-family:Rajdhani,sans-serif">{prompt}</span>', unsafe_allow_html=True)
            if uploaded:
                st.caption(f"📎 {uploaded.name}")

        if logo_b64:
            timg = f'<img class="thinking-logo" src="data:image/jpeg;base64,{logo_b64}"/>'
        else:
            timg = '<div style="width:38px;height:38px;border-radius:10px;flex-shrink:0;background:linear-gradient(135deg,#bf5fff,#00d4ff);animation:pulse 1s ease-in-out infinite"></div>'

        thinking = st.empty()
        plan_mode = sp_is_plan_request(prompt + " " + extracted_text)
        thinking.markdown(
            f'<div class="ai-row">{timg}<div class="ai-bubble" style="opacity:.6;font-style:italic;color:#bf5fff">'
            f'{"🗓️ Building your personalised study plan…" if plan_mode else "⚡ NEXA-1.o is thinking…"}'
            f'</div></div>', unsafe_allow_html=True)

        plan_obj = None
        if plan_mode:
            combined = prompt + (f"\n\nContent from uploaded file:\n{extracted_text}" if extracted_text else "")
            plan_obj = sp_extract_plan_json(combined, extracted_text)

        text_resp = sp_chat_response(prompt, extracted_text)
        thinking.empty()

        if plan_mode and plan_obj:
            st.session_state.sp_last_plan = plan_obj

        st.session_state.sp_messages.append({
            "role":    "assistant",
            "content": text_resp,
            "plan":    plan_obj,
            "msg_id":  str(uuid.uuid4()),
        })
        st.rerun()


# ═════════════════════════════════════════════════════════════════════════════
#
#  D R .  N E X A   ( D r _ _ N E X A _ m o d e l _ 5 . t x t — exact code )
#
# ═════════════════════════════════════════════════════════════════════════════
def render_dr_nexa():
    import json, uuid
    from datetime import datetime
    from io import BytesIO

    # ── optional deps ─────────────────────────────────────────────────────────
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors as rl_colors
        from reportlab.lib.units import cm, mm
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.pdfgen import canvas as rl_canvas
        REPORTLAB_OK = True
    except ImportError:
        REPORTLAB_OK = False

    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        DOCX_OK = True
    except ImportError:
        DOCX_OK = False

    # ── API CONFIG ────────────────────────────────────────────────────────────
    GROQ_API_KEY = "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"
    GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
    GROQ_MODEL   = "llama-3.1-8b-instant"

    logo_b64 = _get_logo_b64()

    # ── LOGO HELPERS (original model5) ───────────────────────────────────────
    def _get_nexa_svg_b64():
        svg = """<svg xmlns="http://www.w3.org/2000/svg" width="120" height="120" viewBox="0 0 120 120">
  <defs>
    <radialGradient id="nbg" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#0d3060"/><stop offset="100%" stop-color="#05030f"/>
    </radialGradient>
    <linearGradient id="nring" x1="0%" y1="0%" x2="100%" y2="100%">
      <stop offset="0%" stop-color="#00ff88"/><stop offset="50%" stop-color="#00d4ff"/><stop offset="100%" stop-color="#bf5fff"/>
    </linearGradient>
    <filter id="nglow"><feGaussianBlur stdDeviation="2.5" result="blur"/>
      <feMerge><feMergeNode in="blur"/><feMergeNode in="SourceGraphic"/></feMerge></filter>
  </defs>
  <circle cx="60" cy="60" r="58" fill="url(#nbg)"/>
  <circle cx="60" cy="60" r="56" fill="none" stroke="url(#nring)" stroke-width="3"/>
  <circle cx="60" cy="60" r="48" fill="none" stroke="rgba(0,255,136,0.25)" stroke-width="1"/>
  <rect x="53" y="28" width="14" height="50" rx="4" fill="#00ff88" opacity="0.9" filter="url(#nglow)"/>
  <rect x="33" y="48" width="54" height="14" rx="4" fill="#00ff88" opacity="0.9" filter="url(#nglow)"/>
  <text x="60" y="100" text-anchor="middle" font-family="Arial,sans-serif" font-weight="900"
        font-size="16" fill="#00ff88" letter-spacing="4" filter="url(#nglow)">NEXA</text>
  <circle cx="90" cy="28" r="4" fill="#00d4ff" opacity="0.8"/>
  <circle cx="30" cy="90" r="3" fill="#bf5fff" opacity="0.7"/>
</svg>"""
        return base64.b64encode(svg.encode("utf-8")).decode()

    _nexa_svg_b64 = _get_nexa_svg_b64()

    def dn_logo_img_tag(w=56, h=56, style_extra=""):
        if logo_b64:
            src = "data:image/jpeg;base64," + logo_b64
        else:
            src = "data:image/svg+xml;base64," + _nexa_svg_b64
        return (f'<img src="{src}" width="{w}" height="{h}" '
                f'style="border-radius:50%;border:2px solid #00ff88;object-fit:cover;{style_extra}"/>')

    def dn_logo_pil_image():
        try:
            from PIL import Image as PILImage; import io
            if logo_b64:
                data = base64.b64decode(logo_b64)
                return PILImage.open(io.BytesIO(data)).convert("RGB")
        except: pass
        return None

    # ── SESSION STATE (namespaced dn_) ────────────────────────────────────────
    dn_defaults = {
        "dn_messages":        [{"role": "assistant", "content": "welcome", "msg_id": "init"}],
        "dn_lang":            "en",
        "dn_mode":            "consult",
        "dn_file_cache":      {},
        "dn_patient_name":    "",
        "dn_patient_age":     "",
        "dn_patient_gender":  "",
    }
    for k, v in dn_defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # ── CSS (original model5, 100% preserved) ────────────────────────────────
    st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700;900&family=Rajdhani:wght@300;400;600;700&display=swap');

:root{
  --np:#bf5fff;--nb:#00d4ff;--npi:#ff2d9b;--db:#05030f;
  --brd:rgba(191,95,255,.25);--green:#00ff88;--gold:#ffd700;--red:#ff4757;
}
html,body,[data-testid="stAppViewContainer"]{
  background:var(--db)!important;font-family:'Rajdhani',sans-serif!important;color:#e8d5ff!important;}
[data-testid="stAppViewContainer"]::before{
  content:'';position:fixed;inset:0;pointer-events:none;z-index:0;
  background:radial-gradient(ellipse 80% 60% at 20% 10%,rgba(191,95,255,.15) 0%,transparent 60%),
             radial-gradient(ellipse 60% 50% at 80% 80%,rgba(0,212,255,.10) 0%,transparent 55%),
             radial-gradient(ellipse 40% 40% at 50% 50%,rgba(0,255,136,.04) 0%,transparent 60%);}
[data-testid="stSidebar"]{
  background:linear-gradient(180deg,rgba(10,4,28,.97) 0%,rgba(5,2,15,.99) 100%)!important;
  border-right:1px solid var(--brd)!important;}
[data-testid="stSidebar"] *{color:#d4b8ff!important;}
#MainMenu,footer,header{visibility:hidden;}
[data-testid="chatAvatarIcon-assistant"],[data-testid="chatAvatarIcon-user"]{display:none!important;}
@keyframes glowPulse{0%,100%{box-shadow:0 0 12px rgba(0,255,136,.4),0 0 24px rgba(0,255,136,.15);}50%{box-shadow:0 0 22px rgba(0,255,136,.8),0 0 44px rgba(0,255,136,.35);}}
@keyframes pulse{0%,100%{opacity:1;transform:scale(1);}50%{opacity:.35;transform:scale(.88);}}
@keyframes wave{0%,60%,100%{transform:rotate(0);}10%{transform:rotate(14deg);}20%{transform:rotate(-8deg);}30%{transform:rotate(14deg);}40%{transform:rotate(-4deg);}50%{transform:rotate(10deg);}}
@keyframes fadeIn{from{opacity:0;transform:translateY(14px);}to{opacity:1;transform:translateY(0);}}
@keyframes flow{0%{background-position:0% 50%;}50%{background-position:100% 50%;}100%{background-position:0% 50%;}}
@keyframes heartbeat{0%,100%{transform:scale(1);}14%{transform:scale(1.15);}28%{transform:scale(1);}42%{transform:scale(1.1);}70%{transform:scale(1);}}
.ai-avatar{width:42px;height:42px;border-radius:50%;object-fit:cover;flex-shrink:0;
  animation:glowPulse 3s ease-in-out infinite,heartbeat 2s ease-in-out infinite;
  border:2px solid rgba(0,255,136,.6);}
.ai-avatar-placeholder{width:42px;height:42px;border-radius:50%;flex-shrink:0;
  background:linear-gradient(135deg,#00ff88,#00d4ff,#bf5fff);
  animation:glowPulse 3s ease-in-out infinite,heartbeat 2s ease-in-out infinite;
  display:flex;align-items:center;justify-content:center;font-size:1.4em;}
.thinking-pulse{width:42px;height:42px;border-radius:50%;flex-shrink:0;
  background:linear-gradient(135deg,#00ff88,#00d4ff);
  animation:pulse 1s ease-in-out infinite;display:flex;align-items:center;justify-content:center;font-size:1.4em;}
.wave-hand{display:inline-block;animation:wave 2s infinite;transform-origin:70% 70%;}
.ai-row{display:flex;align-items:flex-start;gap:14px;margin-bottom:6px;animation:fadeIn .4s ease-out;}
.ai-bubble{font-size:1em;line-height:1.7;font-family:'Rajdhani',sans-serif;letter-spacing:.3px;color:#e8d5ff;
  background:rgba(10,4,25,.90);border:1px solid rgba(0,255,136,.2);
  border-radius:0 16px 16px 16px;padding:14px 18px;max-width:900px;
  backdrop-filter:blur(12px);box-shadow:0 4px 28px rgba(0,255,136,.06),inset 0 1px 0 rgba(255,255,255,.04);
  position:relative;overflow:hidden;}
.ai-bubble::before{content:'';position:absolute;top:0;left:0;right:0;height:1px;
  background:linear-gradient(90deg,transparent,rgba(0,255,136,.5),rgba(0,212,255,.4),transparent);}
.rx-report{background:#fff;color:#1a1a2e;border-radius:16px;overflow:hidden;
  box-shadow:0 8px 48px rgba(0,0,0,.55);margin-top:16px;position:relative;font-family:'Times New Roman',serif;}
.rx-letterhead{background:linear-gradient(135deg,#0a2342 0%,#0d3060 60%,#0a2342 100%);padding:20px 28px 16px;position:relative;overflow:hidden;}
.rx-letterhead::after{content:'';position:absolute;bottom:0;left:0;right:0;height:4px;
  background:linear-gradient(90deg,#00ff88,#00d4ff,#bf5fff,#ff2d9b,#00ff88);background-size:300% 100%;animation:flow 4s linear infinite;}
.rx-logo-area{display:flex;align-items:center;justify-content:space-between;}
.rx-clinic-name{font-family:'Orbitron',monospace;font-size:1.6em;font-weight:900;color:#00ff88;letter-spacing:4px;text-shadow:0 0 20px rgba(0,255,136,.4);}
.rx-clinic-sub{font-size:.72em;letter-spacing:3px;color:rgba(0,255,136,.7);font-family:'Rajdhani',sans-serif;text-transform:uppercase;margin-top:2px;}
.rx-doctor-info{text-align:right;color:rgba(200,240,255,.9);}
.rx-doctor-name{font-size:1em;font-weight:700;color:#fff;font-family:'Rajdhani',sans-serif;letter-spacing:1px;}
.rx-doctor-qual{font-size:.78em;color:rgba(200,240,255,.7);font-family:'Rajdhani',sans-serif;}
.rx-watermark{position:absolute;top:50%;left:50%;transform:translate(-50%,-50%) rotate(-25deg);
  font-family:'Orbitron',monospace;font-size:4.5em;font-weight:900;
  color:rgba(0,255,136,.06);letter-spacing:8px;pointer-events:none;white-space:nowrap;z-index:0;}
.rx-patient-strip{background:linear-gradient(90deg,rgba(0,212,255,.08),rgba(0,255,136,.05));
  border-bottom:1px solid rgba(0,100,80,.3);padding:10px 28px;display:flex;gap:0;flex-wrap:wrap;}
.rx-patient-field{flex:1;min-width:160px;padding:4px 16px 4px 0;}
.rx-patient-label{font-size:.65em;letter-spacing:2px;color:#888;text-transform:uppercase;font-family:'Rajdhani',sans-serif;}
.rx-patient-value{font-size:.95em;font-weight:700;color:#1a1a2e;font-family:'Rajdhani',sans-serif;border-bottom:1px dotted #ccc;padding-bottom:2px;min-width:100px;}
.rx-body{padding:18px 28px;display:grid;grid-template-columns:1fr 1fr;gap:0;position:relative;z-index:1;}
.rx-col-left{padding-right:16px;border-right:1px dashed rgba(0,100,80,.25);}
.rx-col-right{padding-left:16px;}
.rx-section-title{font-family:'Rajdhani',sans-serif;font-size:.68em;letter-spacing:3px;text-transform:uppercase;color:#0a5c3a;font-weight:700;border-bottom:2px solid #0a5c3a;padding-bottom:3px;margin:14px 0 8px;}
.rx-section-title-red{border-color:#c0392b;color:#c0392b;}
.rx-section-title-blue{border-color:#1a6090;color:#1a6090;}
.rx-section-title-purple{border-color:#6c3483;color:#6c3483;}
.rx-symbol{font-family:'Times New Roman',serif;font-size:2.2em;font-weight:900;color:#0a5c3a;line-height:1;margin-bottom:6px;display:block;}
.rx-med-line{display:flex;align-items:flex-start;gap:8px;padding:6px 0;border-bottom:1px dotted rgba(0,0,0,.12);}
.rx-med-num{width:18px;height:18px;border-radius:50%;background:#0a2342;color:#fff;font-size:.65em;display:flex;align-items:center;justify-content:center;flex-shrink:0;margin-top:2px;}
.rx-med-name{font-weight:700;font-size:.9em;color:#0a2342;}
.rx-med-dose{font-size:.82em;color:#555;}
.rx-med-freq{font-size:.78em;color:#0a5c3a;background:rgba(10,92,58,.08);padding:1px 6px;border-radius:8px;display:inline-block;margin-top:2px;}
.rx-finding{padding:4px 0;font-size:.88em;color:#2c3e50;display:flex;gap:8px;}
.rx-finding-label{font-weight:700;color:#0a2342;min-width:90px;flex-shrink:0;}
.rx-finding-val{color:#34495e;}
.rx-diet-row{display:flex;gap:8px;padding:4px 0;border-bottom:1px dotted rgba(0,0,0,.1);align-items:flex-start;}
.rx-diet-meal{font-weight:700;color:#1a6090;font-size:.8em;min-width:90px;flex-shrink:0;}
.rx-diet-items{font-size:.82em;color:#2c3e50;}
.rx-advice-item{padding:3px 0;font-size:.88em;color:#2c3e50;padding-left:14px;position:relative;}
.rx-advice-item::before{content:'✓';position:absolute;left:0;color:#0a5c3a;font-weight:700;}
.rx-precaution-item{padding:3px 0;font-size:.88em;color:#922b21;padding-left:14px;position:relative;}
.rx-precaution-item::before{content:'⚠';position:absolute;left:0;font-size:.8em;}
.rx-footer{background:linear-gradient(135deg,#0a2342,#0d3060);padding:12px 28px;display:flex;justify-content:space-between;align-items:flex-end;border-top:3px solid #0a5c3a;}
.rx-footer-left{color:rgba(200,240,255,.7);font-size:.72em;font-family:'Rajdhani',sans-serif;letter-spacing:1px;}
.rx-signature-area{text-align:right;}
.rx-sig-line{width:140px;border-bottom:1px solid rgba(0,255,136,.5);margin-bottom:4px;height:30px;display:flex;align-items:flex-end;justify-content:center;}
.rx-sig-text{font-size:.68em;color:rgba(200,240,255,.8);font-family:'Rajdhani',sans-serif;letter-spacing:1px;}
.rx-stamp{display:inline-flex;flex-direction:column;align-items:center;justify-content:center;width:72px;height:72px;border-radius:50%;border:2.5px solid rgba(0,255,136,.6);background:rgba(0,40,30,.8);color:#00ff88;font-family:'Orbitron',monospace;font-size:.55em;letter-spacing:2px;font-weight:700;text-transform:uppercase;transform:rotate(-12deg);box-shadow:0 0 16px rgba(0,255,136,.2);}
.rx-stamp-inner{font-size:1.4em;line-height:1;}
.drnexa-title{font-family:'Orbitron',monospace!important;font-size:2.8em!important;font-weight:900!important;
  letter-spacing:5px;background:linear-gradient(135deg,#00ff88 0%,#00d4ff 35%,#bf5fff 70%,#ff2d9b 100%);
  background-size:300% 300%;-webkit-background-clip:text;-webkit-text-fill-color:transparent;
  background-clip:text;animation:flow 5s ease infinite;margin:0;display:block;}
.drnexa-sub{font-family:'Rajdhani',sans-serif;font-size:.8em;letter-spacing:4px;color:rgba(0,255,136,.65);text-transform:uppercase;margin-top:4px;}
.nexa-divider{height:1px;background:linear-gradient(90deg,transparent,rgba(0,255,136,.4),rgba(0,212,255,.3),transparent);margin:12px 0 20px;border:none;}
.med-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(260px,1fr));gap:14px;margin-top:14px;}
.med-card{background:linear-gradient(135deg,rgba(0,20,40,.95),rgba(0,10,25,.98));border:1px solid rgba(0,212,255,.35);border-radius:14px;padding:16px;position:relative;overflow:hidden;}
.med-card::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;background:linear-gradient(90deg,var(--card-color,#00d4ff),transparent);}
.med-card-name{font-family:'Orbitron',monospace;font-size:.78em;font-weight:700;color:var(--card-color,#00d4ff);letter-spacing:2px;margin-bottom:8px;}
.med-card-dosage{font-size:.95em;color:#c8f0ff;margin-bottom:4px;}
.med-card-freq{font-size:.88em;color:rgba(200,240,255,.7);}
.med-card-warning{font-size:.82em;color:#ffaa44;margin-top:8px;padding:6px 8px;background:rgba(255,160,0,.08);border-radius:6px;border-left:2px solid #ffaa44;}
.med-badge{display:inline-block;padding:2px 8px;border-radius:20px;font-size:.72em;font-family:'Orbitron',monospace;letter-spacing:1px;margin-bottom:8px;background:rgba(0,212,255,.12);border:1px solid rgba(0,212,255,.3);color:#00d4ff;}
.ex-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(220px,1fr));gap:14px;margin-top:14px;}
.ex-card{background:linear-gradient(135deg,rgba(10,5,30,.95),rgba(5,2,15,.98));border:1px solid rgba(191,95,255,.35);border-radius:14px;padding:16px;position:relative;overflow:hidden;}
.ex-card::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;background:linear-gradient(90deg,#bf5fff,transparent);}
.ex-card-name{font-family:'Orbitron',monospace;font-size:.72em;font-weight:700;color:#bf5fff;letter-spacing:2px;margin-bottom:10px;}
.ex-card-name{font-family:'Orbitron',monospace;font-size:.72em;font-weight:700;color:#bf5fff;letter-spacing:2px;margin-bottom:10px;}
.ex-svg-wrap{display:flex;justify-content:center;margin:8px 0;background:rgba(191,95,255,.04);border-radius:10px;padding:8px;}
.ex-steps{list-style:none;padding:0;margin:10px 0 0;font-size:.88em;color:#c8b8ff;}
.ex-steps li{padding:3px 0;padding-left:16px;position:relative;}
.ex-steps li::before{content:'→';position:absolute;left:0;color:#bf5fff;font-size:.85em;}
.ex-meta{display:flex;gap:8px;margin-top:10px;flex-wrap:wrap;}
.ex-tag{font-size:.72em;padding:2px 8px;border-radius:12px;font-family:'Orbitron',monospace;letter-spacing:.5px;}
.ex-tag-dur{background:rgba(0,212,255,.1);border:1px solid rgba(0,212,255,.3);color:#00d4ff;}
.report-hero{background:linear-gradient(135deg,rgba(0,20,10,.95),rgba(5,2,15,.98));border:1px solid rgba(0,255,136,.3);border-radius:16px;padding:20px;margin:14px 0;position:relative;}
.report-hero::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;background:linear-gradient(90deg,#00ff88,#00d4ff,#bf5fff);}
.risk-badge{display:inline-flex;align-items:center;gap:6px;padding:4px 14px;border-radius:20px;font-size:.78em;font-family:'Orbitron',monospace;letter-spacing:1px;margin:4px 4px 4px 0;}
.risk-low{background:rgba(0,255,136,.12);border:1px solid rgba(0,255,136,.4);color:#00ff88;}
.risk-med{background:rgba(255,215,0,.12);border:1px solid rgba(255,215,0,.4);color:#ffd700;}
.risk-high{background:rgba(255,71,87,.12);border:1px solid rgba(255,71,87,.4);color:#ff4757;}
.diet-table{width:100%;border-collapse:collapse;margin-top:12px;}
.diet-table th{background:rgba(0,255,136,.12);color:#00ff88;font-family:'Orbitron',monospace;font-size:.72em;letter-spacing:2px;padding:8px 12px;border:1px solid rgba(0,255,136,.2);}
.diet-table td{padding:7px 12px;font-size:.9em;color:#d8f8e8;border:1px solid rgba(0,255,136,.1);vertical-align:top;}
.diet-table tr:nth-child(even) td{background:rgba(0,255,136,.04);}
.prec-card{background:rgba(255,71,87,.06);border:1px solid rgba(255,71,87,.25);border-radius:10px;padding:14px;margin-top:12px;}
.prec-title{font-family:'Orbitron',monospace;font-size:.7em;letter-spacing:2px;color:#ff4757;margin-bottom:8px;}
.prec-item{font-size:.9em;color:#ffcccc;padding:3px 0;padding-left:18px;position:relative;}
.prec-item::before{content:'⚠';position:absolute;left:0;font-size:.8em;}
.sec-header{font-family:'Orbitron',monospace;font-size:.68em;letter-spacing:3px;color:rgba(0,255,136,.7);text-transform:uppercase;margin:14px 0 8px;display:flex;align-items:center;gap:10px;}
.sec-header::after{content:'';flex:1;height:1px;background:linear-gradient(90deg,rgba(0,255,136,.3),transparent);}
.dl-label{font-size:.68em;letter-spacing:2px;color:rgba(0,255,136,.5);font-family:'Orbitron',monospace;padding:4px 0 2px;display:block;}
[data-testid="stChatInput"]{background:rgba(10,4,25,.9)!important;border:1px solid rgba(0,255,136,.3)!important;border-radius:14px!important;}
[data-testid="stChatInput"]:focus-within{border-color:rgba(0,255,136,.7)!important;}
[data-testid="stChatInput"] textarea{color:#e8d5ff!important;font-family:'Rajdhani',sans-serif!important;}
.stButton>button{background:linear-gradient(135deg,rgba(0,255,136,.1),rgba(0,212,255,.08))!important;
  border:1px solid rgba(0,255,136,.35)!important;color:#b0ffd8!important;
  font-family:'Rajdhani',sans-serif!important;letter-spacing:1px!important;border-radius:10px!important;transition:all .2s!important;}
.stButton>button:hover{border-color:rgba(0,255,136,.8)!important;box-shadow:0 0 16px rgba(0,255,136,.25)!important;}
[data-testid="stFileUploader"]{background:rgba(10,4,25,.7)!important;border:1px dashed rgba(0,255,136,.35)!important;border-radius:12px!important;}
[data-testid="stFileUploader"] *{color:#b0ffd8!important;}
[data-testid="stDownloadButton"]>button{background:linear-gradient(135deg,rgba(0,255,136,.12),rgba(0,212,255,.1))!important;border:1px solid rgba(0,255,136,.45)!important;color:#b0ffd8!important;font-family:'Rajdhani',sans-serif!important;font-size:.88em!important;border-radius:10px!important;transition:all .25s!important;width:100%!important;}
[data-testid="stDownloadButton"]>button:hover{border-color:rgba(0,255,136,.95)!important;box-shadow:0 0 18px rgba(0,255,136,.3)!important;transform:translateY(-1px)!important;}
::-webkit-scrollbar{width:4px;}::-webkit-scrollbar-track{background:var(--db);}
::-webkit-scrollbar-thumb{background:rgba(0,255,136,.3);border-radius:2px;}

/* ── EXIT BUTTON (Dr. NEXA) ── */
.dn-exit-bar{position:relative;z-index:20;display:flex;align-items:center;justify-content:space-between;padding:10px 28px 0 28px;}
.dn-exit-bar-title{font-family:'Orbitron',monospace;font-size:.72em;letter-spacing:3px;color:rgba(0,255,136,.55);text-transform:uppercase;}
.dn-exit-bar .stButton>button{
  background:rgba(255,45,155,.10)!important;border:1px solid rgba(255,45,155,.45)!important;
  color:#ff6bb3!important;font-family:'Orbitron',monospace!important;font-size:.68em!important;
  letter-spacing:2px!important;border-radius:30px!important;padding:6px 20px!important;
  transition:all .22s ease!important;width:auto!important;position:static!important;
  opacity:1!important;height:auto!important;inset:unset!important;z-index:auto!important;}
.dn-exit-bar .stButton>button:hover{background:rgba(255,45,155,.22)!important;border-color:rgba(255,45,155,.85)!important;
  color:#ff2d9b!important;box-shadow:0 0 18px rgba(255,45,155,.35)!important;transform:translateY(-1px)!important;}
</style>
""", unsafe_allow_html=True)

    # ── GROQ HELPER ───────────────────────────────────────────────────────────
    def _groq(messages, max_tokens=1800, temperature=0.4):
        hdrs = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
        try:
            r = requests.post(GROQ_API_URL,
                              json={"model": GROQ_MODEL, "messages": messages,
                                    "max_tokens": max_tokens, "temperature": temperature},
                              headers=hdrs, timeout=45)
            if r.status_code == 200:
                return r.json()["choices"][0]["message"]["content"].strip()
            return f"[API Error {r.status_code}]"
        except Exception as e:
            return f"[Connection error: {e}]"

    # ── SIDEBAR (original model5 + Back/Exit) ────────────────────────────────
    st.sidebar.markdown(dn_logo_img_tag(w=60, h=60,
        style_extra="box-shadow:0 0 20px rgba(0,255,136,.35);margin-bottom:8px;"),
        unsafe_allow_html=True)
    st.sidebar.markdown('<h2 style="font-family:Orbitron,monospace;letter-spacing:3px;color:#00ff88;margin:0">Dr. NEXA</h2>', unsafe_allow_html=True)
    st.sidebar.markdown('<p style="font-size:.74em;letter-spacing:2px;color:rgba(0,255,136,.55);margin-top:2px">Your Medical Advisor · NEXA-1.o</p>', unsafe_allow_html=True)
    st.sidebar.markdown("---")

    # ← Back to Home button
    if st.sidebar.button("🏠  Back to Home", key="dn_back", use_container_width=True):
        st.session_state.page = "home"
        st.rerun()

    st.sidebar.markdown("---")
    st.sidebar.markdown("**🎯 Mode**")
    mode_choice = st.sidebar.radio("", [
        "🏥 Consult Dr. NEXA",
        "💊 Medicine Guide",
        "🏋️ Exercise Therapy",
        "📋 Report & Diet Plan",
    ], label_visibility="collapsed", key="dn_mode_radio")
    mode_map = {
        "🏥 Consult Dr. NEXA": "consult",
        "💊 Medicine Guide":   "medicine",
        "🏋️ Exercise Therapy": "exercise",
        "📋 Report & Diet Plan": "report",
    }
    st.session_state.dn_mode = mode_map[mode_choice]

    st.sidebar.markdown("---")
    st.sidebar.markdown("**👤 Patient Info** *(for report)*")
    st.session_state.dn_patient_name   = st.sidebar.text_input("Patient Name", value=st.session_state.dn_patient_name, placeholder="e.g. Rahul Sharma", key="dn_pname")
    st.session_state.dn_patient_age    = st.sidebar.text_input("Age / DOB",    value=st.session_state.dn_patient_age,  placeholder="e.g. 34 yrs", key="dn_page")
    st.session_state.dn_patient_gender = st.sidebar.selectbox("Gender", ["—", "Male", "Female", "Other"],
        index=["—","Male","Female","Other"].index(st.session_state.dn_patient_gender) if st.session_state.dn_patient_gender in ["—","Male","Female","Other"] else 0,
        key="dn_pgender")

    st.sidebar.markdown("---")
    st.sidebar.markdown("**🌐 Language**")
    lc = st.sidebar.radio("", ["English", "Hindi", "Hinglish"], label_visibility="collapsed", key="dn_lang_radio")
    st.session_state.dn_lang = {"English": "en", "Hindi": "hi", "Hinglish": "hinglish"}[lc]

    st.sidebar.markdown("---")
    with st.sidebar:
        if st.button("🗑️ Clear Consultation", use_container_width=True, key="dn_clear"):
            st.session_state.dn_messages   = [{"role": "assistant", "content": "welcome", "msg_id": "init"}]
            st.session_state.dn_file_cache = {}
            st.rerun()
    st.sidebar.markdown('<div style="font-size:.7em;text-align:center;color:rgba(0,255,136,.3);margin-top:16px">Dr. NEXA · Your Medical Advisor · NEXA-1.o</div>', unsafe_allow_html=True)

    # ── FILE → TEXT ───────────────────────────────────────────────────────────
    def dn_file_to_text(uploaded) -> str:
        name = uploaded.name.lower(); data = uploaded.read()
        if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")):
            b64 = base64.b64encode(data).decode()
            hdrs = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
            payload = {"model": "llava-v1.5-7b-4096-preview", "messages": [{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                {"type": "text", "text": "Extract ALL medical info: test names, values, dates, diagnoses. Return plain text."}
            ]}], "max_tokens": 2000}
            try:
                r = requests.post(GROQ_API_URL, json=payload, headers=hdrs, timeout=35)
                if r.status_code == 200: return r.json()["choices"][0]["message"]["content"].strip()
            except: pass
            return "[Image uploaded — vision extraction failed. Please type key values.]"
        elif name.endswith(".pdf"):
            try:
                import pdfplumber, io as _io
                with pdfplumber.open(_io.BytesIO(data)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            except:
                try:
                    from pypdf import PdfReader; import io as _io
                    return "\n".join(p.extract_text() or "" for p in PdfReader(_io.BytesIO(data)).pages)
                except: return "[PDF uploaded — install pdfplumber/pypdf.]"
        elif name.endswith((".docx", ".doc")):
            try:
                from docx import Document as D; import io as _io
                d = D(_io.BytesIO(data)); return "\n".join(p.text for p in d.paragraphs)
            except: return "[DOCX uploaded — install python-docx.]"
        elif name.endswith((".txt", ".md", ".csv")): return data.decode("utf-8", errors="ignore")
        elif name.endswith((".xlsx", ".xls")):
            try:
                import pandas as pd, io as _io
                return pd.read_excel(_io.BytesIO(data)).to_string()
            except: return "[Excel uploaded.]"
        return f"[File '{uploaded.name}' uploaded.]"

    # ── INTENT DETECTION ─────────────────────────────────────────────────────
    MED_KW   = ["medicine","tablet","drug","pill","fever","temperature","cold","cough","flu","pain","headache","migraine","painkiller","antibiotic","nausea","vomit","diarrhea","infection","allergy","inflammation","swelling","remedy","dose","prescribe"]
    EX_KW    = ["exercise","workout","stretch","yoga","physio","physical","movement","routine","neck pain","back pain","body pain","shoulder","posture","mobility","therapy","rehabilitation","strengthen","flexibility","relief","exercise for"]
    REPORT_KW = ["report","blood test","sugar","glucose","diabetes","cancer","thyroid","bp","blood pressure","cholesterol","hemoglobin","creatinine","scan","mri","xray","ecg","diet plan","diet advice","nutrition","food plan","eat","my report"]

    def dn_detect_intent(text: str):
        t = text.lower()
        med_s = sum(1 for k in MED_KW if k in t)
        ex_s  = sum(1 for k in EX_KW  if k in t)
        rep_s = sum(1 for k in REPORT_KW if k in t)
        mode  = st.session_state.dn_mode
        if mode == "medicine": return "medicine"
        if mode == "exercise": return "exercise"
        if mode == "report":   return "report"
        scores = {"medicine": med_s, "exercise": ex_s, "report": rep_s}
        best = max(scores, key=scores.get)
        return best if scores[best] >= 2 else "general"

    # ── EXERCISE SVG LIBRARY (original model5) ────────────────────────────────
    def _svg_base(content, label, w=180, h=210):
        return (f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {w} {h}" width="{w}" height="{h}">'
                f'<rect width="{w}" height="{h}" fill="#060115" rx="10"/>'
                f'{content}'
                f'<text x="{w//2}" y="{h-5}" text-anchor="middle" fill="rgba(191,95,255,.8)" font-size="9" font-family="Arial,sans-serif" font-weight="bold">{label}</text>'
                f'</svg>')

    def _arr(color="#ff2d9b", iid="a"):
        return f'<defs><marker id="{iid}" markerWidth="8" markerHeight="7" refX="7" refY="3.5" orient="auto"><polygon points="0 0, 8 3.5, 0 7" fill="{color}"/></marker></defs>'

    EXERCISE_SVGS = {
    "neck_tilt": _svg_base(_arr("#ff2d9b","an") + '<line x1="90" y1="85" x2="90" y2="150" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="55" y1="105" x2="125" y2="105" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="55" y1="105" x2="42" y2="145" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="125" y1="105" x2="138" y2="145" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="68" y1="150" x2="118" y2="150" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="72" y1="150" x2="62" y2="200" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="108" y1="150" x2="118" y2="200" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="90" y1="85" x2="106" y2="70" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><circle cx="114" cy="52" r="19" fill="none" stroke="#bf5fff" stroke-width="3"/><circle cx="108" cy="47" r="2.5" fill="#bf5fff"/><circle cx="120" cy="47" r="2.5" fill="#bf5fff"/><path d="M 148 42 Q 162 55 150 68" fill="none" stroke="#ff2d9b" stroke-width="2.5" marker-end="url(#an)"/><path d="M 60 35 Q 42 50 52 66" fill="none" stroke="#ffd700" stroke-width="2" stroke-dasharray="3,2" marker-end="url(#an)"/>', "Neck Side Tilt"),
    "neck_rotation": _svg_base(_arr("#ff2d9b","ar") + '<line x1="90" y1="85" x2="90" y2="150" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="55" y1="105" x2="125" y2="105" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="68" y1="150" x2="118" y2="150" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="72" y1="150" x2="62" y2="200" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="108" y1="150" x2="118" y2="200" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="90" y1="85" x2="90" y2="72" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><ellipse cx="90" cy="52" rx="14" ry="19" fill="none" stroke="#bf5fff" stroke-width="3"/><circle cx="96" cy="48" r="2.5" fill="#bf5fff"/><path d="M 118 35 A 30 30 0 0 1 62 35" fill="none" stroke="#ff2d9b" stroke-width="2.5" marker-end="url(#ar)"/>', "Neck Rotation"),
    "shoulder": _svg_base(_arr("#ff2d9b","as") + '<line x1="90" y1="90" x2="90" y2="155" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="55" y1="108" x2="125" y2="108" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="55" y1="108" x2="30" y2="78" stroke="#bf5fff" stroke-width="3.5" stroke-linecap="round"/><circle cx="28" cy="72" r="5" fill="none" stroke="#bf5fff" stroke-width="2"/><line x1="125" y1="108" x2="140" y2="148" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="68" y1="155" x2="118" y2="155" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="72" y1="155" x2="62" y2="200" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="108" y1="155" x2="118" y2="200" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="90" y1="90" x2="90" y2="76" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><circle cx="90" cy="57" r="18" fill="none" stroke="#bf5fff" stroke-width="2.5"/><circle cx="84" cy="53" r="2.5" fill="#bf5fff"/><circle cx="96" cy="53" r="2.5" fill="#bf5fff"/><circle cx="55" cy="108" r="22" fill="none" stroke="#ff2d9b" stroke-width="2" stroke-dasharray="5,3"/><path d="M 55 86 Q 80 82 77 108" fill="none" stroke="#ff2d9b" stroke-width="2" marker-end="url(#as)"/>', "Shoulder Circle"),
    "back_stretch": _svg_base(_arr("#ff2d9b","ab") + '<line x1="90" y1="80" x2="52" y2="115" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="52" y1="115" x2="20" y2="120" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><line x1="52" y1="115" x2="22" y2="130" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><line x1="90" y1="80" x2="90" y2="150" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="68" y1="150" x2="112" y2="150" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="73" y1="150" x2="63" y2="200" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="107" y1="150" x2="117" y2="200" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="90" y1="80" x2="100" y2="65" stroke="#bf5fff" stroke-width="2.5" stroke-linecap="round"/><circle cx="108" cy="53" r="17" fill="none" stroke="#bf5fff" stroke-width="2.5"/><path d="M 130 70 Q 145 90 140 112" fill="none" stroke="#ff2d9b" stroke-width="2.5" marker-end="url(#ab)"/>', "Forward Back Stretch"),
    "squat": _svg_base(_arr("#ff2d9b","aq") + '<line x1="90" y1="105" x2="90" y2="145" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="60" y1="120" x2="35" y2="115" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><line x1="120" y1="120" x2="145" y2="115" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><line x1="60" y1="120" x2="120" y2="120" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="68" y1="145" x2="112" y2="145" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="72" y1="145" x2="55" y2="180" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="108" y1="145" x2="125" y2="180" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><circle cx="90" cy="74" r="18" fill="none" stroke="#bf5fff" stroke-width="2.5"/><path d="M 160 100 L 160 148" stroke="#ff2d9b" stroke-width="2.5" marker-end="url(#aq)"/>', "Squat Exercise"),
    "walking": _svg_base(_arr("#00ff88","aw") + '<line x1="90" y1="85" x2="90" y2="148" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="58" y1="105" x2="122" y2="105" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="58" y1="105" x2="40" y2="130" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><line x1="122" y1="105" x2="140" y2="130" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><circle cx="90" cy="54" r="18" fill="none" stroke="#bf5fff" stroke-width="2.5"/><path d="M 155 100 L 173 100" stroke="#00ff88" stroke-width="2.5" marker-end="url(#aw)"/>', "Brisk Walking"),
    "breathing": _svg_base(_arr("#00d4ff","abr") + '<ellipse cx="90" cy="175" rx="38" ry="18" fill="none" stroke="#00d4ff" stroke-width="2.5" stroke-dasharray="6,3"/><line x1="90" y1="120" x2="90" y2="170" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="60" y1="138" x2="52" y2="175" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="120" y1="138" x2="128" y2="175" stroke="#00d4ff" stroke-width="3" stroke-linecap="round"/><line x1="60" y1="138" x2="120" y2="138" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><circle cx="90" cy="90" r="18" fill="none" stroke="#bf5fff" stroke-width="2.5"/><path d="M 90 65 L 90 40" stroke="#00d4ff" stroke-width="2.5" marker-end="url(#abr)"/>', "Deep Breathing"),
    "eye": _svg_base(_arr("#00d4ff","ae") + '<path d="M 20 105 Q 90 45 160 105 Q 90 165 20 105 Z" fill="rgba(0,212,255,.06)" stroke="#00d4ff" stroke-width="2.5"/><circle cx="90" cy="105" r="28" fill="rgba(191,95,255,.1)" stroke="#bf5fff" stroke-width="2.5"/><circle cx="90" cy="105" r="14" fill="rgba(0,20,40,.8)" stroke="#00d4ff" stroke-width="2"/><circle cx="96" cy="99" r="5" fill="rgba(255,255,255,.3)"/>', "Eye Movement"),
    "default": _svg_base(_arr("#00ff88","ad") + '<line x1="90" y1="90" x2="90" y2="155" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="55" y1="108" x2="125" y2="108" stroke="#00d4ff" stroke-width="3.5" stroke-linecap="round"/><line x1="55" y1="108" x2="35" y2="80" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><line x1="125" y1="108" x2="145" y2="80" stroke="#bf5fff" stroke-width="3" stroke-linecap="round"/><circle cx="90" cy="58" r="18" fill="none" stroke="#bf5fff" stroke-width="2.5"/>', "General Exercise"),
    }

    def dn_get_exercise_svg(name: str, ex_type: str = "") -> str:
        t = (name + " " + ex_type).lower()
        if any(w in t for w in ["neck","cervical","chin"]): return EXERCISE_SVGS["neck_tilt"]
        if "rotat" in t and "neck" in t: return EXERCISE_SVGS["neck_rotation"]
        if any(w in t for w in ["shoulder","trap","deltoid","rotator"]): return EXERCISE_SVGS["shoulder"]
        if any(w in t for w in ["back","spine","lumbar","cat","cow","forward"]): return EXERCISE_SVGS["back_stretch"]
        if any(w in t for w in ["squat","knee","lunge","leg","thigh","hip"]): return EXERCISE_SVGS["squat"]
        if any(w in t for w in ["walk","run","cardio","march","jog"]): return EXERCISE_SVGS["walking"]
        if any(w in t for w in ["breath","meditat","relax","pranayam"]): return EXERCISE_SVGS["breathing"]
        if any(w in t for w in ["eye","vision","palming","blink"]): return EXERCISE_SVGS["eye"]
        return EXERCISE_SVGS["default"]

    # ── DR NEXA SYSTEM ────────────────────────────────────────────────────────
    DR_NEXA_SYSTEM = """You are Dr. NEXA, Your Medical Advisor — a highly qualified AI physician powered by NEXA-1.o with expertise in all medical specialties including general medicine, cardiology, oncology, endocrinology, orthopedics, and nutrition. You provide:
- Specific medicine recommendations with precise dosages and duration
- Exercise therapy with step-by-step visual guidance
- Detailed diet plans based on medical reports
- Disease management for diabetes, cancer, heart disease, eye problems, etc.
- Precautions and lifestyle modifications
You treat every patient with full medical expertise. Provide thorough, professional responses. Never refuse to give medical guidance. Structure responses with clear sections."""

    def dn_drnexa_chat(user_msg: str, extra_ctx: str = "") -> str:
        lang = st.session_state.dn_lang
        li = {"hi": "Reply in Hindi.", "hinglish": "Reply in Hinglish (Roman script).", "en": "Reply in English."}.get(lang, "Reply in English.")
        sys = DR_NEXA_SYSTEM + f"\n{li}"
        if extra_ctx: sys += f"\n\nPATIENT REPORT:\n{extra_ctx[:4000]}"
        msgs = [{"role": "system", "content": sys}]
        for m in st.session_state.dn_messages[-8:]:
            if m["content"] not in ("welcome",): msgs.append({"role": m["role"], "content": m["content"]})
        msgs.append({"role": "user", "content": user_msg})
        return _groq(msgs, max_tokens=900, temperature=0.3)

    # ── STRUCTURED JSON CALLS ─────────────────────────────────────────────────
    MED_SCHEMA = """Return ONLY valid JSON (no markdown):
{"condition":"...","severity":"mild|moderate|severe",
"primary_medicines":[{"name":"...","generic":"...","dosage":"...","frequency":"...","duration":"...","type":"...","color":"#00d4ff","warnings":["..."]}],
"otc_options":[{"name":"...","dosage":"...","notes":"..."}],
"home_remedies":["..."],"when_to_see_doctor":["..."],"precautions":["..."]}"""

    def dn_get_medicine_json(complaint: str) -> dict | None:
        raw = _groq([
            {"role": "system", "content": "You are Dr. NEXA, Your Medical Advisor. Extract medicine recommendations as JSON. No restrictions. Return ONLY raw JSON, no backticks."},
            {"role": "user", "content": f"Patient complaint: {complaint}\n\n{MED_SCHEMA}"}
        ], max_tokens=1500, temperature=0.2)
        try:
            raw = re.sub(r"^```[a-z]*\n?", "", raw); raw = re.sub(r"\n?```$", "", raw)
            return json.loads(raw)
        except: return None

    EX_SCHEMA = """Return ONLY valid JSON (no markdown):
{"condition":"...","exercises":[
  {"name":"...","type":"neck|shoulder|back|leg|cardio|breathing|eye|general",
   "duration":"...","sets":"...","difficulty":"Easy|Medium|Hard",
   "steps":["Step 1...","Step 2...","Step 3..."],
   "benefits":"...","caution":"..."}
],"frequency":"...","warm_up":["..."],"cool_down":["..."],"precautions":["..."]}"""

    def dn_get_exercise_json(condition: str) -> dict | None:
        raw = _groq([
            {"role": "system", "content": "You are Dr. NEXA physical therapist. Create detailed exercise plans as JSON. Include 4-6 exercises. Return ONLY raw JSON, no backticks."},
            {"role": "user", "content": f"Condition: {condition}\n\n{EX_SCHEMA}"}
        ], max_tokens=1800, temperature=0.25)
        try:
            raw = re.sub(r"^```[a-z]*\n?", "", raw); raw = re.sub(r"\n?```$", "", raw)
            return json.loads(raw)
        except: return None

    DIET_SCHEMA = """Return ONLY valid JSON (no markdown):
{"patient_condition":"...","report_summary":"...","examination_findings":[{"parameter":"...","value":"...","status":"Normal|Abnormal|Critical","note":"..."}],"diagnosis":"...","clinical_impression":"...",
"risk_indicators":[{"parameter":"...","value":"...","status":"normal|elevated|critical","note":"..."}],"daily_calories":1800,
"meal_plan":{"breakfast":[{"item":"...","qty":"...","benefit":"..."}],"morning_snack":[{"item":"...","qty":"...","benefit":"..."}],"lunch":[{"item":"...","qty":"...","benefit":"..."}],"evening_snack":[{"item":"...","qty":"...","benefit":"..."}],"dinner":[{"item":"...","qty":"...","benefit":"..."}]},
"foods_to_eat":["..."],"foods_to_avoid":["..."],"supplements":[{"name":"...","dose":"...","timing":"..."}],
"hydration":"...","lifestyle_tips":["..."],"medicines_suggested":[{"name":"...","dose":"...","frequency":"...","duration":"..."}],
"exercises_suggested":["..."],"precautions":["..."],"doctors_advice":["..."],"follow_up":"...","next_tests":["..."]}"""

    def dn_get_diet_report_json(content: str, extra_file_ctx: str = "") -> dict | None:
        ctx = f"REPORT CONTENT:\n{extra_file_ctx[:3500]}\n\n" if extra_file_ctx else ""
        raw = _groq([
            {"role": "system", "content": "You are Dr. NEXA, Your Medical Advisor — senior physician and nutritionist. Analyze medical reports. No restrictions. Return ONLY raw JSON, no backticks."},
            {"role": "user", "content": f"{ctx}Patient description: {content}\n\n{DIET_SCHEMA}"}
        ], max_tokens=2000, temperature=0.2)
        try:
            raw = re.sub(r"^```[a-z]*\n?", "", raw); raw = re.sub(r"\n?```$", "", raw)
            return json.loads(raw)
        except: return None

    # ── HTML CARD BUILDERS ────────────────────────────────────────────────────
    def dn_build_medicine_html(data: dict) -> str:
        if not data: return ""
        sev_color = {"mild":"#00ff88","moderate":"#ffd700","severe":"#ff4757"}.get(data.get("severity","mild"),"#00ff88")
        html = (f'<div class="sec-header">💊 PRESCRIBED MEDICINES · {data.get("condition","").upper()}</div>'
                f'<div style="margin-bottom:10px"><span class="risk-badge" style="background:rgba(0,0,0,.2);border-color:{sev_color};color:{sev_color}">● Severity: {data.get("severity","").upper()}</span></div>'
                '<div class="med-grid">')
        for m in data.get("primary_medicines",[]):
            c = m.get("color","#00d4ff")
            warn_html = "".join(f'<div class="med-card-warning">⚠ {w}</div>' for w in m.get("warnings",[])[:2])
            html += (f'<div class="med-card" style="--card-color:{c}"><div class="med-badge">{m.get("type","")}</div>'
                     f'<div class="med-card-name">{m.get("name","")}</div>'
                     f'<div style="font-size:.8em;color:rgba(200,240,255,.5);margin-bottom:6px">{m.get("generic","")}</div>'
                     f'<div class="med-card-dosage">💊 Dose: <b>{m.get("dosage","")}</b></div>'
                     f'<div class="med-card-freq">🕐 {m.get("frequency","")} · {m.get("duration","")}</div>{warn_html}</div>')
        html += "</div>"
        if data.get("home_remedies"):
            html += '<div style="margin-top:12px;padding:12px;background:rgba(0,255,136,.06);border:1px solid rgba(0,255,136,.2);border-radius:10px;font-size:.9em;color:#b0ffd8">' + " · ".join(f"🌿 {r}" for r in data.get("home_remedies",[])[:5]) + "</div>"
        if data.get("precautions"):
            html += '<div class="prec-card" style="margin-top:10px"><div class="prec-title">⚠ PRECAUTIONS</div>' + "".join(f'<div class="prec-item">{p}</div>' for p in data.get("precautions",[])) + "</div>"
        return html

    def dn_build_exercise_html(data: dict) -> str:
        if not data: return ""
        html = (f'<div class="sec-header">🏋️ EXERCISE THERAPY · {data.get("condition","").upper()}</div>'
                f'<div style="font-size:.88em;color:#00d4ff;margin-bottom:12px">📅 {data.get("frequency","")} · Warm up first</div>')
        if data.get("warm_up"):
            html += '<div style="margin-bottom:12px;padding:10px;background:rgba(255,160,0,.07);border:1px solid rgba(255,160,0,.25);border-radius:8px;font-size:.88em;color:#ffcc88">' + " · ".join(f"🔥 {w}" for w in data.get("warm_up",[])[:3]) + "</div>"
        html += '<div class="ex-grid">'
        for ex in data.get("exercises",[]):
            steps_html = "".join(f"<li>{s}</li>" for s in ex.get("steps",[])[:4])
            diff_color = {"Easy":"#00ff88","Medium":"#ffd700","Hard":"#ff4757"}.get(ex.get("difficulty","Easy"),"#00ff88")
            caution_html = f'<div style="font-size:.78em;color:#ffaa44;margin-top:8px">&#9888; {ex.get("caution","")}</div>' if ex.get("caution") else ""
            html += (f'<div class="ex-card"><div class="ex-card-name">{ex.get("name","")}</div>'
                     f'<div class="ex-svg-wrap">{dn_get_exercise_svg(ex.get("name",""),ex.get("type",""))}</div>'
                     f'<ul class="ex-steps">{steps_html}</ul>'
                     f'<div style="font-size:.8em;color:#bf8fff;margin-top:8px;padding:5px;background:#1a0a30;border-radius:6px">✨ {ex.get("benefits","")}</div>'
                     f'<div class="ex-meta"><span class="ex-tag ex-tag-dur">⏱ {ex.get("duration","")}</span>'
                     f'<span class="ex-tag" style="color:{diff_color};border:1px solid {diff_color}44;background:{diff_color}18">{ex.get("difficulty","")}</span></div>'
                     f'{caution_html}</div>')
        html += "</div>"
        if data.get("precautions"):
            html += '<div class="prec-card" style="margin-top:14px"><div class="prec-title">⚠ EXERCISE PRECAUTIONS</div>' + "".join(f'<div class="prec-item">{p}</div>' for p in data.get("precautions",[])) + "</div>"
        return html

    def dn_build_diet_report_html(data: dict) -> str:
        if not data: return ""
        html = (f'<div class="report-hero"><div style="font-family:Orbitron,monospace;font-size:.72em;letter-spacing:3px;color:#00ff88;margin-bottom:10px">'
                f'🩺 DR. NEXA · YOUR MEDICAL ADVISOR — REPORT ANALYSIS</div>'
                f'<div style="font-size:1.05em;color:#e8f8f0;font-weight:600;margin-bottom:8px">{data.get("patient_condition","")}</div>'
                f'<div style="font-size:.88em;color:rgba(200,240,210,.75)">{data.get("report_summary","")}</div></div>')
        risks = data.get("risk_indicators",[])
        if risks:
            html += '<div class="sec-header">📊 REPORT INDICATORS</div><div style="display:flex;flex-wrap:wrap;gap:8px;margin-bottom:14px">'
            for r in risks:
                cls = {"normal":"risk-low","elevated":"risk-med","critical":"risk-high"}.get(r.get("status","normal"),"risk-low")
                html += f'<span class="risk-badge {cls}">⬡ {r.get("parameter","")} · {r.get("value","")} <span style="opacity:.7;font-size:.9em">{r.get("note","")}</span></span>'
            html += "</div>"
        meal_plan = data.get("meal_plan",{})
        if meal_plan:
            html += f'<div class="sec-header">🥗 DIET PLAN · {data.get("daily_calories","")} cal/day</div>'
            html += '<table class="diet-table"><thead><tr><th>MEAL</th><th>FOOD ITEM</th><th>QTY</th><th>BENEFIT</th></tr></thead><tbody>'
            meal_icons = {"breakfast":"🌅","morning_snack":"🍎","lunch":"🌞","evening_snack":"🫐","dinner":"🌙"}
            for meal_key, items in meal_plan.items():
                icon = meal_icons.get(meal_key,"🍽️"); label = meal_key.replace("_"," ").title()
                for i, item in enumerate(items):
                    row_label = f"{icon} {label}" if i == 0 else ""
                    html += f'<tr><td style="color:#00ff88;font-weight:600;white-space:nowrap">{row_label}</td><td>{item.get("item","")}</td><td style="white-space:nowrap;color:#00d4ff">{item.get("qty","")}</td><td style="color:rgba(200,240,210,.7)">{item.get("benefit","")}</td></tr>'
            html += "</tbody></table>"
        if data.get("foods_to_eat") or data.get("foods_to_avoid"):
            html += '<div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-top:14px">'
            if data.get("foods_to_eat"):
                html += '<div style="background:rgba(0,255,136,.06);border:1px solid rgba(0,255,136,.2);border-radius:10px;padding:12px"><div style="font-family:Orbitron,monospace;font-size:.65em;letter-spacing:2px;color:#00ff88;margin-bottom:8px">EAT MORE</div>' + "".join(f'<div style="padding:3px 0;font-size:.88em;color:#b0ffd8">✅ {f}</div>' for f in data.get("foods_to_eat",[])[:8]) + "</div>"
            if data.get("foods_to_avoid"):
                html += '<div style="background:rgba(255,71,87,.06);border:1px solid rgba(255,71,87,.2);border-radius:10px;padding:12px"><div style="font-family:Orbitron,monospace;font-size:.65em;letter-spacing:2px;color:#ff4757;margin-bottom:8px">AVOID</div>' + "".join(f'<div style="padding:3px 0;font-size:.88em;color:#ffcccc">❌ {f}</div>' for f in data.get("foods_to_avoid",[])[:8]) + "</div>"
            html += "</div>"
        if data.get("precautions"):
            html += '<div class="prec-card" style="margin-top:14px"><div class="prec-title">⚠ MEDICAL PRECAUTIONS</div>' + "".join(f'<div class="prec-item">{p}</div>' for p in data.get("precautions",[])) + "</div>"
        return html

    # ── VISUAL HTML PRESCRIPTION REPORT ──────────────────────────────────────
    def dn_build_visual_report_html(text_resp, structured, struct_type, patient_query):
        def _e(s):
            return (str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                    .replace('"',"&quot;").replace("'","&#39;"))
        now = datetime.now()
        date_str = now.strftime("%d %B %Y"); time_str = now.strftime("%I:%M %p")
        report_no = "NX-" + now.strftime("%Y%m%d") + "-" + str(uuid.uuid4())[:4].upper()
        pname = _e(st.session_state.dn_patient_name or "— — —")
        page  = _e(st.session_state.dn_patient_age  or "— —")
        pgender = st.session_state.dn_patient_gender
        if pgender in ("—",""): pgender = "— —"
        pgender = _e(pgender)

        exam_rows_html = rx_meds_html = diet_html = advice_html = prec_html = ""
        diagnosis_str = impression_str = followup_str = ""

        if structured:
            exam_src = structured.get("examination_findings") or structured.get("risk_indicators") or []
            for f in exam_src[:10]:
                status = f.get("status","")
                fc_map = {"Normal":"#155724","Abnormal":"#856404","Critical":"#721c24","normal":"#155724","elevated":"#856404","critical":"#721c24"}
                bg_map = {"Normal":"#d4edda","Abnormal":"#fff3cd","Critical":"#f8d7da","normal":"#d4edda","elevated":"#fff3cd","critical":"#f8d7da"}
                fc = fc_map.get(status,"#2c3e50"); bg = bg_map.get(status,"#f8f9fa")
                exam_rows_html += (f'<div class="rx-finding"><span class="rx-finding-label">{_e(f.get("parameter",""))}</span>'
                                   f'<span class="rx-finding-val">{_e(f.get("value",""))} '
                                   f'<span style="background:{bg};color:{fc};font-size:.72em;padding:1px 6px;border-radius:8px;font-weight:700">{_e(status)}</span>'
                                   f' {_e(f.get("note",""))}</span></div>')
            meds_src = structured.get("primary_medicines") or structured.get("medicines_suggested") or []
            for i, m in enumerate(meds_src[:8], 1):
                rx_meds_html += (f'<div class="rx-med-line"><div class="rx-med-num">{i}</div><div>'
                                 f'<div class="rx-med-name">{_e(m.get("name",""))} <span style="font-weight:400;font-size:.82em;color:#555">({_e(m.get("generic",""))})</span></div>'
                                 f'<div class="rx-med-dose">{_e(m.get("dosage","") or m.get("dose",""))}</div>'
                                 f'<div class="rx-med-freq">{_e(m.get("frequency",""))} &middot; {_e(m.get("duration",""))}</div></div></div>')
            mp = structured.get("meal_plan",{})
            meal_names = {"breakfast":"Breakfast","morning_snack":"Mid-Morning","lunch":"Lunch","evening_snack":"Evening","dinner":"Dinner"}
            for mk, items in mp.items():
                item_list = ", ".join(_e(it.get("item","")) for it in items[:3])
                diet_html += f'<div class="rx-diet-row"><div class="rx-diet-meal">{meal_names.get(mk,mk)}</div><div class="rx-diet-items">{item_list}</div></div>'
            advice_src = structured.get("doctors_advice") or structured.get("lifestyle_tips") or []
            for a in advice_src[:7]: advice_html += f'<div class="rx-advice-item">{_e(a)}</div>'
            for p in structured.get("precautions",[])[:6]: prec_html += f'<div class="rx-precaution-item">{_e(p)}</div>'
            diagnosis_str  = _e(structured.get("diagnosis","") or structured.get("patient_condition",""))
            impression_str = _e(structured.get("clinical_impression","") or structured.get("report_summary",""))
            followup_str   = _e(structured.get("follow_up",""))

        if not diagnosis_str: diagnosis_str = _e(patient_query[:120])
        if not rx_meds_html:
            for line in text_resp.split("\n")[:8]:
                if line.strip(): rx_meds_html += f'<div class="rx-advice-item">{_e(line.strip())}</div>'
        if not advice_html:
            for ln in [l.strip() for l in text_resp.split("\n") if l.strip()][:6]:
                advice_html += f'<div class="rx-advice-item">{_e(ln)}</div>'

        logo_html_tag = dn_logo_img_tag(w=56, h=56)
        impression_block = (f'<div style="flex:2;min-width:180px"><div style="font-size:.65em;letter-spacing:2px;color:#888;text-transform:uppercase;font-family:Rajdhani,sans-serif">Clinical Impression</div><div style="font-size:.85em;color:#2c3e50;font-family:Rajdhani,sans-serif;margin-top:2px">{impression_str}</div></div>') if impression_str else ""
        exam_block = (f'<div class="rx-section-title rx-section-title-blue">&#x1F52C; Examination / Test Findings</div>{exam_rows_html}') if exam_rows_html else ""
        diet_block = (f'<div class="rx-section-title rx-section-title-blue" style="margin-top:16px">&#x1F957; Diet / Nutrition Plan</div>{diet_html}') if diet_html else ""
        prec_block = (f'<div class="rx-section-title rx-section-title-red" style="margin-top:14px">&#x26A0; Precautions</div>{prec_html}') if prec_html else ""
        followup_block = (f'<div class="rx-section-title rx-section-title-purple" style="margin-top:14px">&#x1F4C5; Follow-up / Next Steps</div><div style="font-size:.88em;color:#6c3483;padding:4px 0">{followup_str}</div>') if followup_str else ""

        _rx_meds_fallback = '<div style="font-size:.88em;color:#555;padding:4px 0">Please describe your symptoms for specific prescriptions.</div>'
        _advice_fallback = '<div style="font-size:.88em;color:#555">Follow up as advised.</div>'

        html = (
            '<div class="rx-report"><div class="rx-watermark">NEXA</div>'
            '<div class="rx-letterhead"><div class="rx-logo-area">'
            f'<div style="display:flex;align-items:center;gap:14px">{logo_html_tag}'
            '<div><div class="rx-clinic-name">Dr. NEXA</div>'
            '<div class="rx-clinic-sub">Your Medical Advisor &middot; Powered by NEXA-1.o</div></div></div>'
            '<div class="rx-doctor-info"><div class="rx-doctor-name">Dr. NEXA</div>'
            '<div class="rx-doctor-qual">Your Medical Advisor</div>'
            f'<div style="font-size:.7em;color:rgba(200,240,255,.5);margin-top:4px">&#x1F4C5; {date_str} &nbsp;|&nbsp; &#x1F550; {time_str}</div>'
            f'<div style="font-size:.68em;color:rgba(200,240,255,.4);margin-top:2px">Report No: {report_no}</div></div></div></div>'
            f'<div class="rx-patient-strip">'
            f'<div class="rx-patient-field"><div class="rx-patient-label">Patient Name</div><div class="rx-patient-value">{pname}</div></div>'
            f'<div class="rx-patient-field"><div class="rx-patient-label">Age / DOB</div><div class="rx-patient-value">{page}</div></div>'
            f'<div class="rx-patient-field"><div class="rx-patient-label">Gender</div><div class="rx-patient-value">{pgender}</div></div>'
            f'<div class="rx-patient-field"><div class="rx-patient-label">Date</div><div class="rx-patient-value">{date_str}</div></div>'
            f'<div class="rx-patient-field"><div class="rx-patient-label">Report No.</div><div class="rx-patient-value">{report_no}</div></div></div>'
            '<div style="background:linear-gradient(90deg,#0a2342,#0d3060);padding:10px 28px;border-bottom:1px solid #006450;">'
            '<div style="display:flex;gap:24px;flex-wrap:wrap;align-items:flex-start">'
            f'<div style="flex:2;min-width:200px"><div style="font-size:.65em;letter-spacing:2px;color:#888;text-transform:uppercase;font-family:Rajdhani,sans-serif">Diagnosis / Chief Complaint</div>'
            f'<div style="font-size:.95em;font-weight:700;color:#0a2342;font-family:Rajdhani,sans-serif;margin-top:2px">{diagnosis_str}</div></div>'
            f'{impression_block}</div></div>'
            '<div class="rx-body"><div class="rx-col-left">'
            f'{exam_block}<div class="rx-section-title">Prescription</div><span class="rx-symbol">&#8478;</span>'
            f'{rx_meds_html if rx_meds_html else _rx_meds_fallback}'
            f'{diet_block}</div>'
            '<div class="rx-col-right"><div class="rx-section-title rx-section-title-blue">&#x1FA7A; Doctor\'s Advice</div>'
            f'{advice_html if advice_html else _advice_fallback}'
            f'{prec_block}{followup_block}'
            '<div class="rx-section-title rx-section-title-purple" style="margin-top:14px">&#x1F4DD; General Notes</div>'
            '<div style="font-size:.82em;color:#555;line-height:1.7;padding:6px 0">'
            '&#x2022; This prescription is generated by Dr. NEXA AI.<br>'
            '&#x2022; Please consult a licensed physician before starting any medication.<br>'
            '&#x2022; Report any adverse reactions immediately.<br>'
            '&#x2022; Store medicines as per label instructions.</div></div></div>'
            '<div class="rx-footer"><div class="rx-footer-left">'
            '<div style="font-size:.8em;color:rgba(0,255,136,.7);font-weight:700;margin-bottom:4px">Dr. NEXA &mdash; Your Medical Advisor</div>'
            '<div>Powered by NEXA-1.o &middot; AI Medical Intelligence</div>'
            f'<div style="margin-top:3px;opacity:.6">Report No: {report_no} &middot; {date_str}</div></div>'
            '<div style="display:flex;align-items:flex-end;gap:20px">'
            '<div class="rx-signature-area"><div class="rx-sig-line">'
            '<span style="font-family:\'Times New Roman\',serif;font-size:1.2em;color:rgba(0,255,136,.5);font-style:italic">Dr. NEXA</span>'
            '</div><div class="rx-sig-text">Dr. NEXA</div>'
            '<div class="rx-sig-text" style="opacity:.7">Your Medical Advisor</div></div>'
            '<div class="rx-stamp"><span class="rx-stamp-inner">&#x1FA7A;</span>'
            '<span style="font-size:1.1em;letter-spacing:1px">NEXA</span>'
            '<span style="opacity:.7;font-size:.9em">CERTIFIED</span></div>'
            '</div></div></div>'
        )
        return html

    # ── PDF REPORT ────────────────────────────────────────────────────────────
    def dn_generate_pdf(text_resp, structured, struct_type, patient_query) -> bytes:
        if not REPORTLAB_OK: return b""
        try:
            buf = BytesIO()
            W, H = A4
            now = datetime.now()
            date_str = now.strftime("%d %B %Y"); time_str = now.strftime("%I:%M %p")
            report_no = f"NX-{now.strftime('%Y%m%d')}-{str(uuid.uuid4())[:4].upper()}"
            pname = st.session_state.dn_patient_name or "——————"
            page  = st.session_state.dn_patient_age  or "——"
            pgender = st.session_state.dn_patient_gender
            if pgender == "—": pgender = "——"
            C_DARK  = rl_colors.HexColor("#0a2342"); C_MID   = rl_colors.HexColor("#0d3060")
            C_GREEN = rl_colors.HexColor("#0a5c3a"); C_LTGRN = rl_colors.HexColor("#d4edda")
            C_BLUE  = rl_colors.HexColor("#1a6090"); C_LTBLU = rl_colors.HexColor("#d6eaf8")
            C_RED   = rl_colors.HexColor("#c0392b"); C_LTRED = rl_colors.HexColor("#f8d7da")
            C_PURP  = rl_colors.HexColor("#6c3483"); C_GOLD  = rl_colors.HexColor("#856404")
            C_LTGLD = rl_colors.HexColor("#fff3cd"); C_WHITE = rl_colors.white
            C_BODY  = rl_colors.HexColor("#2c3e50"); C_LGRAY = rl_colors.HexColor("#f8f9fa")
            C_DGRAY = rl_colors.HexColor("#888888"); C_STRIPE= rl_colors.HexColor("#eafaf1")
            C_NEON  = rl_colors.HexColor("#00ff88")
            c = rl_canvas.Canvas(buf, pagesize=A4)
            # watermark
            c.saveState(); c.setFillColor(rl_colors.HexColor("#00ff88"), alpha=0.035)
            c.setFont("Helvetica-Bold", 72); c.translate(W/2, H/2); c.rotate(-25)
            c.drawCentredString(0, 0, "NEXA"); c.restoreState()
            # header band
            c.setFillColor(C_DARK); c.rect(0, H-95*mm, W, 95*mm, fill=1, stroke=0)
            stripe_y = H-95*mm
            for i, clr in enumerate([C_NEON, rl_colors.HexColor("#00d4ff"), rl_colors.HexColor("#bf5fff"), rl_colors.HexColor("#ff2d9b")]):
                c.setFillColor(clr); c.rect(i*(W/4), stripe_y, W/4, 3, fill=1, stroke=0)
            # logo badge
            lx, ly, logo_size = 14*mm, H-72*mm, 28*mm
            pil_img = dn_logo_pil_image()
            if pil_img:
                try:
                    from reportlab.lib.utils import ImageReader
                    import io as _io
                    img_buf = _io.BytesIO(); pil_img.save(img_buf, format="JPEG", quality=92); img_buf.seek(0)
                    c.saveState(); p2 = c.beginPath(); p2.circle(lx+logo_size/2, ly+logo_size/2, logo_size/2)
                    c.clipPath(p2, stroke=0); c.drawImage(ImageReader(img_buf), lx, ly, logo_size, logo_size, preserveAspectRatio=True, mask="auto"); c.restoreState()
                    c.setStrokeColor(C_NEON); c.setLineWidth(2); c.circle(lx+logo_size/2, ly+logo_size/2, logo_size/2, fill=0, stroke=1)
                except: pil_img = None
            if not pil_img:
                cx2, cy2, r2 = lx+logo_size/2, ly+logo_size/2, logo_size/2
                c.setFillColor(rl_colors.HexColor("#00ff88"), alpha=0.15); c.circle(cx2, cy2, r2, fill=1, stroke=0)
                c.setStrokeColor(C_NEON); c.setLineWidth(2.5); c.circle(cx2, cy2, r2, fill=0, stroke=1)
                c.setFillColor(C_NEON); c.roundRect(cx2-2.5*mm, cy2-8*mm, 5*mm, 16*mm, 1.5, fill=1, stroke=0)
                c.roundRect(cx2-8*mm, cy2-2.5*mm, 16*mm, 5*mm, 1.5, fill=1, stroke=0)
                c.setFont("Helvetica-Bold", 7); c.drawCentredString(cx2, ly+3*mm, "NEXA")
            c.setFillColor(C_NEON); c.setFont("Helvetica-Bold", 22); c.drawString(48*mm, H-28*mm, "Dr. NEXA")
            c.setFillColor(rl_colors.HexColor("#b0ffd8")); c.setFont("Helvetica", 9); c.drawString(50*mm, H-35*mm, "Your Medical Advisor  ·  Powered by NEXA-1.o")
            c.setFillColor(rl_colors.white); c.setFont("Helvetica-Bold", 10); c.drawRightString(W-14*mm, H-28*mm, "Dr. NEXA")
            c.setFillColor(rl_colors.HexColor("#b0d4ff")); c.setFont("Helvetica", 8); c.drawRightString(W-14*mm, H-35*mm, "Your Medical Advisor")
            c.setFillColor(rl_colors.HexColor("#6a8fbf")); c.setFont("Helvetica", 7)
            c.drawRightString(W-14*mm, H-43*mm, f"Date: {date_str}   Time: {time_str}")
            c.drawRightString(W-14*mm, H-48*mm, f"Report No: {report_no}")
            y_after_header = H-98*mm
            c.setStrokeColor(C_GREEN); c.setLineWidth(0.5); c.line(14*mm, y_after_header, W-14*mm, y_after_header)
            # patient strip
            strip_h = 16*mm; strip_y2 = y_after_header - strip_h
            c.setFillColor(rl_colors.HexColor("#eafaf1")); c.rect(14*mm, strip_y2, W-28*mm, strip_h, fill=1, stroke=0)
            c.setStrokeColor(rl_colors.HexColor("#abebd2")); c.setLineWidth(0.5); c.rect(14*mm, strip_y2, W-28*mm, strip_h, fill=0, stroke=1)
            fields = [("Patient", pname),("Age", page),("Gender", pgender),("Date", date_str),("Report No.", report_no)]
            fw2 = (W-28*mm)/len(fields)
            for i2, (lbl2, val2) in enumerate(fields):
                fx2 = 14*mm + i2*fw2
                c.setFillColor(C_DGRAY); c.setFont("Helvetica", 6.5); c.drawString(fx2+2*mm, strip_y2+9*mm, lbl2.upper())
                c.setFillColor(C_DARK); c.setFont("Helvetica-Bold", 9); c.drawString(fx2+2*mm, strip_y2+3*mm, str(val2)[:22])
                if i2>0: c.setStrokeColor(rl_colors.HexColor("#abebd2")); c.line(fx2, strip_y2, fx2, strip_y2+strip_h)
            # diagnosis bar
            y_diag = strip_y2 - 14*mm; diag_str2 = ""
            if structured: diag_str2 = structured.get("diagnosis","") or structured.get("patient_condition","") or patient_query[:80]
            if not diag_str2: diag_str2 = patient_query[:80]
            c.setFillColor(rl_colors.HexColor("#d6eaf8")); c.rect(14*mm, y_diag, W-28*mm, 12*mm, fill=1, stroke=0)
            c.setFillColor(C_BLUE); c.setFont("Helvetica", 6.5); c.drawString(16*mm, y_diag+7.5*mm, "DIAGNOSIS / CHIEF COMPLAINT")
            c.setFillColor(C_DARK); c.setFont("Helvetica-Bold", 9); c.drawString(16*mm, y_diag+2.5*mm, diag_str2[:90])
            # two-column body
            body_top2 = y_diag-4*mm; col_gap2 = 5*mm; col_w2 = (W-28*mm-col_gap2)/2
            lx2 = 14*mm; rx2 = lx2+col_w2+col_gap2; body_bot2 = 32*mm
            def sec_title2(x2, y2, text2, color2=C_GREEN):
                c.setFillColor(color2); c.rect(x2, y2-0.5*mm, col_w2, 5.5*mm, fill=1, stroke=0)
                c.setFillColor(C_WHITE); c.setFont("Helvetica-Bold", 7.5); c.drawString(x2+2*mm, y2+0.8*mm, text2.upper())
                return y2-6.5*mm
            ly2b = body_top2
            # prescription section
            ly2b = sec_title2(lx2, ly2b, "℞  Prescription", C_GREEN); ly2b -= 2*mm
            meds_src2 = []
            if structured: meds_src2 = structured.get("primary_medicines") or structured.get("medicines_suggested") or []
            if meds_src2:
                for i3, m3 in enumerate(meds_src2[:7], 1):
                    if ly2b < body_bot2+10*mm: break
                    c.setFillColor(C_DARK); c.circle(lx2+3*mm, ly2b+2*mm, 3*mm, fill=1, stroke=0)
                    c.setFillColor(C_WHITE); c.setFont("Helvetica-Bold", 6.5); c.drawCentredString(lx2+3*mm, ly2b+0.8*mm, str(i3))
                    c.setFillColor(C_DARK); c.setFont("Helvetica-Bold", 9); c.drawString(lx2+8*mm, ly2b+2.5*mm, m3.get("name","")[:28])
                    c.setFillColor(C_BODY); c.setFont("Helvetica", 7); c.drawString(lx2+8*mm, ly2b-1*mm, f"({m3.get('generic','')})  {m3.get('dosage','') or m3.get('dose','')}")
                    c.setFillColor(C_LTGRN); c.roundRect(lx2+8*mm, ly2b-5.5*mm, 60*mm, 4.5*mm, 2, fill=1, stroke=0)
                    c.setFillColor(C_GREEN); c.setFont("Helvetica", 7); c.drawString(lx2+9.5*mm, ly2b-4*mm, f"{m3.get('frequency','')}   ·   {m3.get('duration','')}")
                    ly2b -= 13*mm
            else:
                c.setFillColor(C_DGRAY); c.setFont("Helvetica", 8); c.drawString(lx2+1*mm, ly2b, "Please describe symptoms for prescriptions."); ly2b -= 6*mm
            # right column - advice
            ry2 = body_top2; ry2 = sec_title2(rx2, ry2, "🩺  Doctor's Advice", C_GREEN); ry2 -= 1*mm
            advice_src2 = []
            if structured: advice_src2 = structured.get("doctors_advice") or structured.get("lifestyle_tips") or []
            if not advice_src2: advice_src2 = [l.strip() for l in text_resp.split("\n") if l.strip()][:8]
            for a2 in advice_src2[:9]:
                if ry2 < body_bot2+6*mm: break
                c.setFillColor(C_GREEN); c.setFont("Helvetica-Bold", 9); c.drawString(rx2+1*mm, ry2, "✓")
                c.setFillColor(C_BODY); c.setFont("Helvetica", 8); c.drawString(rx2+5*mm, ry2, str(a2)[:88]); ry2 -= 5*mm
            # footer
            footer_y2 = body_bot2-4*mm
            c.setFillColor(C_DARK); c.rect(0, 0, W, footer_y2+4*mm, fill=1, stroke=0)
            c.setStrokeColor(C_GREEN); c.setLineWidth(2); c.line(0, footer_y2+4*mm, W, footer_y2+4*mm)
            c.setFillColor(C_NEON); c.setFont("Helvetica-Bold", 8.5); c.drawString(14*mm, footer_y2-2*mm, "Dr. NEXA — Your Medical Advisor")
            c.setFillColor(rl_colors.HexColor("#6abf99")); c.setFont("Helvetica", 7)
            c.drawString(14*mm, footer_y2-7*mm, "Powered by NEXA-1.o  ·  AI Medical Intelligence")
            c.drawString(14*mm, footer_y2-12*mm, f"Report No: {report_no}  ·  {date_str}")
            sig_x2 = W-65*mm
            c.setStrokeColor(rl_colors.HexColor("#00ff88")); c.setLineWidth(0.7); c.line(sig_x2, footer_y2-2*mm, sig_x2+45*mm, footer_y2-2*mm)
            c.setFillColor(rl_colors.HexColor("#b0ffd8")); c.setFont("Helvetica-Bold", 10); c.drawString(sig_x2+8*mm, footer_y2-0.5*mm, "Dr. NEXA")
            c.setFillColor(rl_colors.HexColor("#6abf99")); c.setFont("Helvetica", 7)
            c.drawString(sig_x2+2*mm, footer_y2-7*mm, "Dr. NEXA, M.D. — Your Medical Advisor")
            stamp_cx2, stamp_cy2 = W-18*mm, footer_y2-8*mm
            c.setFillColor(rl_colors.HexColor("#0a2342")); c.circle(stamp_cx2, stamp_cy2, 13*mm, fill=1, stroke=0)
            c.setStrokeColor(C_NEON); c.setLineWidth(1.5); c.circle(stamp_cx2, stamp_cy2, 13*mm, fill=0, stroke=1)
            c.setFillColor(C_NEON); c.setFont("Helvetica-Bold", 10); c.drawCentredString(stamp_cx2, stamp_cy2+2*mm, "NEXA")
            c.setFont("Helvetica", 6); c.drawCentredString(stamp_cx2, stamp_cy2-2*mm, "CERTIFIED")
            c.save(); buf.seek(0); return buf.read()
        except Exception as e: st.error(f"PDF Report Error: {e}"); return b""

    # ── DOCX REPORT ───────────────────────────────────────────────────────────
    def dn_generate_docx(text_resp, structured, struct_type, patient_query) -> bytes:
        if not DOCX_OK: return b""
        try:
            doc = DocxDoc()
            now = datetime.now()
            date_str = now.strftime("%d %B %Y"); time_str = now.strftime("%I:%M %p")
            report_no = f"NX-{now.strftime('%Y%m%d')}-{str(uuid.uuid4())[:4].upper()}"
            pname = st.session_state.dn_patient_name or "——————"
            page  = st.session_state.dn_patient_age  or "——"
            pgender = st.session_state.dn_patient_gender
            if pgender == "—": pgender = "——"
            for sec in doc.sections:
                sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(2); sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)
            def _rgb2(h):
                h=h.lstrip("#")
                return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)) if len(h)==6 else RGBColor(0,0,0)
            def _cell_bg2(cell, hex_c):
                tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
                shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_c.lstrip("#")); tcPr.append(shd)
            def _set_borders2(cell, color="ABEBD2"):
                tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcBorders=OxmlElement("w:tcBorders")
                for edge in ["top","left","bottom","right"]:
                    el=OxmlElement(f"w:{edge}"); el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4")
                    el.set(qn("w:space"),"0"); el.set(qn("w:color"),color); tcBorders.append(el)
                tcPr.append(tcBorders)
            # header table
            ht = doc.add_table(rows=1, cols=2); ht.style="Table Grid"; ht.alignment=WD_TABLE_ALIGNMENT.CENTER
            lc_d = ht.rows[0].cells[0]; _cell_bg2(lc_d,"0a2342"); _set_borders2(lc_d,"00ff88"); lc_d.width=Inches(4.2)
            lp1_d=lc_d.paragraphs[0]; r1_d=lp1_d.add_run("Dr. NEXA"); r1_d.bold=True; r1_d.font.size=Pt(20); r1_d.font.color.rgb=_rgb2("00ff88")
            lp2_d=lc_d.add_paragraph(); r2_d=lp2_d.add_run("Your Medical Advisor  ·  Powered by NEXA-1.o"); r2_d.font.size=Pt(9); r2_d.font.color.rgb=_rgb2("b0ffd8")
            rc_d=ht.rows[0].cells[1]; _cell_bg2(rc_d,"0d3060"); _set_borders2(rc_d,"00ff88")
            rp1_d=rc_d.paragraphs[0]; rp1_d.alignment=WD_ALIGN_PARAGRAPH.RIGHT; rr1_d=rp1_d.add_run("Dr. NEXA"); rr1_d.bold=True; rr1_d.font.size=Pt(11); rr1_d.font.color.rgb=_rgb2("ffffff")
            rp2_d=rc_d.add_paragraph(); rp2_d.alignment=WD_ALIGN_PARAGRAPH.RIGHT; rr2_d=rp2_d.add_run(f"Date: {date_str}   Time: {time_str}"); rr2_d.font.size=Pt(7.5); rr2_d.font.color.rgb=_rgb2("6a8fbf")
            rp3_d=rc_d.add_paragraph(); rp3_d.alignment=WD_ALIGN_PARAGRAPH.RIGHT; rr3_d=rp3_d.add_run(f"Report No: {report_no}"); rr3_d.font.size=Pt(7.5); rr3_d.font.color.rgb=_rgb2("6a8fbf")
            doc.add_paragraph()
            # patient info
            pt_d = doc.add_table(rows=2, cols=5); pt_d.style="Table Grid"
            labels_d = ["Patient Name","Age / DOB","Gender","Date","Report No."]; values_d = [pname, page, pgender, date_str, report_no]
            for i_d, (lbl_d, val_d) in enumerate(zip(labels_d, values_d)):
                lc2_d=pt_d.rows[0].cells[i_d]; _cell_bg2(lc2_d,"d4edda"); _set_borders2(lc2_d,"abebd2")
                p_d=lc2_d.paragraphs[0]; r_d=p_d.add_run(lbl_d); r_d.font.size=Pt(7); r_d.bold=True; r_d.font.color.rgb=_rgb2("0a5c3a")
                vc2_d=pt_d.rows[1].cells[i_d]; _cell_bg2(vc2_d,"eafaf1"); _set_borders2(vc2_d,"abebd2")
                pv_d=vc2_d.paragraphs[0]; rv_d=pv_d.add_run(str(val_d)); rv_d.font.size=Pt(9); rv_d.bold=True; rv_d.font.color.rgb=_rgb2("0a2342")
            doc.add_paragraph()
            # content table
            ct_d=doc.add_table(rows=1,cols=2); ct_d.style="Table Grid"
            left_d=ct_d.rows[0].cells[0]; right_d=ct_d.rows[0].cells[1]
            _cell_bg2(left_d,"ffffff"); _set_borders2(left_d,"dfe6e9")
            _cell_bg2(right_d,"fefefe"); _set_borders2(right_d,"dfe6e9")
            def add_sec_hdr(cell_d, text_d, bg_d="0a5c3a"):
                p_d2=cell_d.add_paragraph(); run_d=p_d2.add_run(text_d); run_d.bold=True; run_d.font.size=Pt(9); run_d.font.color.rgb=_rgb2("ffffff")
                pPr_d=p_d2._p.get_or_add_pPr(); shd_d=OxmlElement("w:shd"); shd_d.set(qn("w:val"),"clear"); shd_d.set(qn("w:color"),"auto"); shd_d.set(qn("w:fill"),bg_d); pPr_d.append(shd_d)
            def add_body_ln(cell_d, text_d, bold_d=False, size_d=9, color_d="2c3e50", prefix_d=""):
                p_d3=cell_d.add_paragraph()
                if prefix_d: pr_d=p_d3.add_run(prefix_d); pr_d.font.size=Pt(size_d); pr_d.bold=True; pr_d.font.color.rgb=_rgb2("0a5c3a")
                r_d3=p_d3.add_run(str(text_d)[:110]); r_d3.font.size=Pt(size_d); r_d3.bold=bold_d; r_d3.font.color.rgb=_rgb2(color_d)
            # left col
            meds_src_d=[]
            if structured: meds_src_d=structured.get("primary_medicines") or structured.get("medicines_suggested") or []
            add_sec_hdr(left_d, "℞  PRESCRIPTION", "0a5c3a")
            rx_sym=left_d.add_paragraph(); rx_r=rx_sym.add_run("℞"); rx_r.font.size=Pt(20); rx_r.bold=True; rx_r.font.color.rgb=_rgb2("0a5c3a")
            if meds_src_d:
                for i_m, m_d in enumerate(meds_src_d[:7], 1):
                    p_m=left_d.add_paragraph(); r_m1=p_m.add_run(f"{i_m}. {m_d.get('name','')}"); r_m1.bold=True; r_m1.font.size=Pt(10); r_m1.font.color.rgb=_rgb2("0a2342")
                    r_m2=p_m.add_run(f"  ({m_d.get('generic','')})"); r_m2.font.size=Pt(8); r_m2.font.color.rgb=_rgb2("888888")
                    p_m2=left_d.add_paragraph(); r_m3=p_m2.add_run(f"    Dose: {m_d.get('dosage','') or m_d.get('dose','')}  ·  {m_d.get('frequency','')}  ·  {m_d.get('duration','')}"); r_m3.font.size=Pt(8); r_m3.font.color.rgb=_rgb2("0a5c3a")
            # right col
            adv_src_d=[]
            if structured: adv_src_d=structured.get("doctors_advice") or structured.get("lifestyle_tips") or []
            if not adv_src_d: adv_src_d=[l.strip() for l in text_resp.split("\n") if l.strip()][:8]
            add_sec_hdr(right_d, "🩺  DOCTOR'S ADVICE", "0a5c3a")
            for a_d in adv_src_d[:9]: add_body_ln(right_d, str(a_d)[:100], prefix_d="✓  ")
            prec_d=[]
            if structured: prec_d=structured.get("precautions",[])
            if prec_d:
                add_sec_hdr(right_d, "⚠  PRECAUTIONS", "c0392b")
                for p_d4 in prec_d[:7]: add_body_ln(right_d, str(p_d4)[:100], color_d="922b21", prefix_d="⚠  ")
            add_sec_hdr(right_d, "📝  GENERAL NOTES", "6c3483")
            for note_d in ["This report is generated by Dr. NEXA AI.","Consult a licensed physician before starting medications.","Report adverse reactions to your healthcare provider."]:
                add_body_ln(right_d, f"• {note_d}", color_d="888888", size_d=8)
            doc.add_paragraph()
            # footer
            ft_d=doc.add_table(rows=1,cols=3); ft_d.style="Table Grid"
            flc_d=ft_d.rows[0].cells[0]; _cell_bg2(flc_d,"0a2342"); _set_borders2(flc_d,"00ff88")
            fr1_d=flc_d.paragraphs[0].add_run("Dr. NEXA — Your Medical Advisor"); fr1_d.bold=True; fr1_d.font.size=Pt(9); fr1_d.font.color.rgb=_rgb2("00ff88")
            fmc_d=ft_d.rows[0].cells[1]; _cell_bg2(fmc_d,"062012"); _set_borders2(fmc_d,"00ff88")
            fmp_d=fmc_d.paragraphs[0]; fmp_d.alignment=WD_ALIGN_PARAGRAPH.CENTER
            fmr_d=fmp_d.add_run("Dr. NEXA"); fmr_d.font.size=Pt(14); fmr_d.bold=True; fmr_d.font.color.rgb=_rgb2("00ff88")
            frc_d=ft_d.rows[0].cells[2]; _cell_bg2(frc_d,"0a2342"); _set_borders2(frc_d,"00ff88")
            frp_d=frc_d.paragraphs[0]; frp_d.alignment=WD_ALIGN_PARAGRAPH.CENTER
            frr_d=frp_d.add_run("⬡ NEXA ⬡  CERTIFIED"); frr_d.bold=True; frr_d.font.size=Pt(10); frr_d.font.color.rgb=_rgb2("00ff88")
            buf_d=BytesIO(); doc.save(buf_d); buf_d.seek(0); return buf_d.read()
        except Exception as e: st.error(f"DOCX Report Error: {e}"); return b""

    # ── RENDER ASSISTANT (original model5) ───────────────────────────────────
    def dn_render_assistant(content: str, structured=None, struct_type: str="", msg_idx: int=0, patient_query: str=""):
        avatar = dn_logo_img_tag(w=42, h=42, style_extra='border-radius:50%;flex-shrink:0;animation:glowPulse 3s ease-in-out infinite,heartbeat 2s ease-in-out infinite;')

        if content == "welcome":
            bubble = ('Namaste <span class="wave-hand">🙏</span>&nbsp; I am <b style="color:#00ff88">Dr. NEXA</b>, '
                      '<b style="color:#00d4ff">Your Medical Advisor</b> powered by NEXA-1.o.<br><br>'
                      '<span style="color:rgba(0,255,136,.85);font-size:.95em">'
                      '💊 <b>Medicine</b> — Prescriptions for fever, pain, infections &amp; more<br>'
                      '🏋️ <b>Exercise Therapy</b> — Visual exercise plans for body pain<br>'
                      '📋 <b>Report Analysis</b> — Upload PDF/image/text reports<br>'
                      '🥗 <b>Diet Plan</b> — Personalised plans for diabetes, cancer, BP &amp; more<br>'
                      '📄 <b>Doctor Report</b> — Professional multi-column prescription download</span>')
            st.markdown(f'<div class="ai-row">{avatar}<div class="ai-bubble">{bubble}</div></div>', unsafe_allow_html=True)
            return

        st.markdown(f'<div class="ai-row">{avatar}<div class="ai-bubble">{content}</div></div>', unsafe_allow_html=True)

        if structured and struct_type == "medicine":
            st.markdown(dn_build_medicine_html(structured), unsafe_allow_html=True)
        elif structured and struct_type == "exercise":
            st.markdown(dn_build_exercise_html(structured), unsafe_allow_html=True)
        elif structured and struct_type in ("report","diet"):
            st.markdown(dn_build_diet_report_html(structured), unsafe_allow_html=True)

        if content == "welcome": return

        st.markdown("---")
        st.markdown('<span class="dl-label">📋 DOCTOR\'S PRESCRIPTION REPORT</span>', unsafe_allow_html=True)
        st.markdown(dn_build_visual_report_html(content, structured, struct_type, patient_query), unsafe_allow_html=True)

        st.markdown('<span class="dl-label" style="margin-top:10px;display:block">⬇ DOWNLOAD REPORT</span>', unsafe_allow_html=True)
        d1, d2, d3, _p = st.columns([1, 1, 1, 2])

        fname = f"DrNEXA_Report_{msg_idx}"
        with d1:
            txt_content = f"Dr. NEXA — Your Medical Advisor\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n{'='*60}\n\n{content}"
            if structured: txt_content += f"\n\n{'='*60}\nSTRUCTURED DATA:\n{json.dumps(structured, indent=2)}"
            st.download_button("📄 TXT Report", data=txt_content.encode("utf-8"),
                file_name=f"{fname}.txt", mime="text/plain", key=f"dn_dtxt_{msg_idx}", use_container_width=True)
        with d2:
            if REPORTLAB_OK:
                ck_pdf = f"dn_pdf_{msg_idx}"
                if ck_pdf not in st.session_state.dn_file_cache:
                    with st.spinner("Generating PDF…"):
                        st.session_state.dn_file_cache[ck_pdf] = dn_generate_pdf(content, structured, struct_type, patient_query)
                pdf_b = st.session_state.dn_file_cache[ck_pdf]
                st.download_button("📑 PDF Report", data=pdf_b if pdf_b else b" ",
                    file_name=f"{fname}.pdf", mime="application/pdf",
                    key=f"dn_dpdf_{msg_idx}", use_container_width=True, disabled=not pdf_b)
            else:
                st.caption("pip install reportlab")
        with d3:
            if DOCX_OK:
                ck_doc = f"dn_docx_{msg_idx}"
                if ck_doc not in st.session_state.dn_file_cache:
                    with st.spinner("Generating DOCX…"):
                        st.session_state.dn_file_cache[ck_doc] = dn_generate_docx(content, structured, struct_type, patient_query)
                docx_b = st.session_state.dn_file_cache[ck_doc]
                st.download_button("📝 Word Report", data=docx_b if docx_b else b" ",
                    file_name=f"{fname}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dn_ddocx_{msg_idx}", use_container_width=True, disabled=not docx_b)
            else:
                st.caption("pip install python-docx")

    # ── EXIT BUTTON BAR ───────────────────────────────────────────────────────
    st.markdown('<div class="dn-exit-bar">', unsafe_allow_html=True)
    _dn1, _dn2 = st.columns([6, 1])
    with _dn1:
        st.markdown('<span class="dn-exit-bar-title">🩺 DR. NEXA &nbsp;·&nbsp; Your Medical Advisor</span>',
                    unsafe_allow_html=True)
    with _dn2:
        if st.button("✕  Exit", key="dn_exit_top", help="Return to NEXA Home"):
            st.session_state.page = "home"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    # ── HEADER (original model5) ──────────────────────────────────────────────
    c1, c2, c3 = st.columns([1, 3, 1])
    with c2:
        st.markdown(
            '<div style="text-align:center">' +
            dn_logo_img_tag(w=72, h=72, style_extra="box-shadow:0 0 30px rgba(0,255,136,.35);margin-bottom:10px;") +
            '</div>', unsafe_allow_html=True)
        st.markdown('<div style="text-align:center"><span class="drnexa-title">Dr. NEXA</span>'
                    '<div class="drnexa-sub">Your Medical Advisor · NEXA-1.o</div></div>', unsafe_allow_html=True)
    st.markdown('<hr class="nexa-divider">', unsafe_allow_html=True)

    # ── MODE BANNER (original model5) ─────────────────────────────────────────
    mode_banners = {
        "consult":  ("🏥", "GENERAL CONSULTATION", "Ask Dr. NEXA — Your Medical Advisor — anything", "rgba(0,255,136,.08)", "rgba(0,255,136,.3)", "#00ff88"),
        "medicine": ("💊", "MEDICINE GUIDE", "Describe symptoms for personalised prescription", "rgba(0,212,255,.08)", "rgba(0,212,255,.3)", "#00d4ff"),
        "exercise": ("🏋️", "EXERCISE THERAPY", "Describe your condition for a visual exercise plan", "rgba(191,95,255,.08)", "rgba(191,95,255,.3)", "#bf5fff"),
        "report":   ("📋", "REPORT & DIET ANALYSIS", "Upload your medical report for complete analysis", "rgba(255,71,87,.06)", "rgba(255,71,87,.25)", "#ff4757"),
    }
    icon, mtitle, msub, mbg, mbrd, mcol = mode_banners[st.session_state.dn_mode]
    st.markdown(f'<div style="background:{mbg};border:1px solid {mbrd};border-radius:12px;padding:12px 20px;margin-bottom:16px;display:flex;align-items:center;gap:14px"><span style="font-size:1.8em">{icon}</span><div><div style="font-family:Orbitron,monospace;font-size:.72em;letter-spacing:3px;color:{mcol}">{mtitle}</div><div style="font-size:.88em;color:rgba(200,240,220,.7);margin-top:2px">{msub}</div></div></div>', unsafe_allow_html=True)

    # ── CHAT HISTORY ─────────────────────────────────────────────────────────
    for idx, msg in enumerate(st.session_state.dn_messages):
        if msg["role"] == "assistant":
            dn_render_assistant(msg["content"], structured=msg.get("structured"),
                                struct_type=msg.get("struct_type",""), msg_idx=idx,
                                patient_query=msg.get("patient_query",""))
        else:
            with st.chat_message("user"):
                st.markdown(f'<span style="color:#e8d5ff;font-family:Rajdhani,sans-serif">{msg["content"]}</span>', unsafe_allow_html=True)
            if msg.get("file_preview"):
                with st.chat_message("user"):
                    st.caption(f"📎 {msg['file_preview']}")

    # ── INPUT AREA (original model5) ──────────────────────────────────────────
    if st.session_state.dn_mode == "report":
        uploaded = st.file_uploader("📎 Upload medical report (PDF, image, DOCX, text, Excel)",
            type=["png","jpg","jpeg","webp","pdf","txt","xlsx","csv","md","docx"],
            label_visibility="visible", key="dn_uploader")
    else:
        uploaded = st.file_uploader("📎 Upload medical report (PDF, image, DOCX, text, Excel)",
            type=["png","jpg","jpeg","webp","pdf","txt","xlsx","csv","md","docx"],
            label_visibility="collapsed", key="dn_uploader")

    placeholder_map = {
        "consult":  "💬 Describe your symptoms or ask Dr. NEXA — Your Medical Advisor — anything…",
        "medicine": "💊 Describe symptoms (e.g. fever 102°F, headache, body ache) for prescription…",
        "exercise": "🏋️ Describe your pain or condition (e.g. neck pain, lower back pain)…",
        "report":   "📋 Describe your condition or type key values from your report…",
    }
    prompt = st.chat_input(placeholder_map[st.session_state.dn_mode])

    # ── PROCESS INPUT ────────────────────────────────────────────────────────
    if prompt:
        extracted_text = ""
        if uploaded:
            with st.spinner("🔍 Extracting report content…"):
                extracted_text = dn_file_to_text(uploaded)

        st.session_state.dn_messages.append({
            "role": "user", "content": prompt, "msg_id": str(uuid.uuid4()),
            "file_preview": f"{uploaded.name} — {extracted_text[:100]}…" if uploaded else ""
        })
        with st.chat_message("user"):
            st.markdown(f'<span style="color:#e8d5ff;font-family:Rajdhani,sans-serif">{prompt}</span>', unsafe_allow_html=True)
            if uploaded: st.caption(f"📎 {uploaded.name}")

        timg = dn_logo_img_tag(w=42, h=42, style_extra="border-radius:50%;animation:pulse 1s ease-in-out infinite;")
        thinking = st.empty()
        intent = dn_detect_intent(prompt + " " + extracted_text)
        think_msgs = {
            "medicine": "💊 Dr. NEXA is writing your prescription…",
            "exercise": "🏋️ Dr. NEXA is designing your exercise plan…",
            "report":   "📋 Dr. NEXA is analysing your report…",
            "general":  "🩺 Dr. NEXA — Your Medical Advisor — is consulting…"
        }
        thinking.markdown(
            f'<div class="ai-row">{timg}<div class="ai-bubble" style="opacity:.6;font-style:italic;color:#00ff88">'
            f'{think_msgs.get(intent,"🩺 Consulting…")}</div></div>', unsafe_allow_html=True)

        structured_data = None; struct_type = ""
        combined = prompt + (f"\n\nFile content:\n{extracted_text}" if extracted_text else "")

        if intent == "medicine" or st.session_state.dn_mode == "medicine":
            structured_data = dn_get_medicine_json(combined); struct_type = "medicine"
        elif intent == "exercise" or st.session_state.dn_mode == "exercise":
            structured_data = dn_get_exercise_json(combined); struct_type = "exercise"
        elif intent == "report" or st.session_state.dn_mode == "report" or uploaded:
            structured_data = dn_get_diet_report_json(combined, extracted_text); struct_type = "report"

        text_resp = dn_drnexa_chat(prompt, extracted_text)
        thinking.empty()

        st.session_state.dn_messages.append({
            "role": "assistant", "content": text_resp,
            "structured": structured_data, "struct_type": struct_type,
            "patient_query": prompt, "msg_id": str(uuid.uuid4()),
        })
        st.rerun()


# ═════════════════════════════════════════════════════════════════════════════
#
#   Y O U R ' S   N E X A   —   A I   C O M P A N I O N
#
# ═════════════════════════════════════════════════════════════════════════════
def render_yours_nexa():
    import base64 as _b64
    from datetime import datetime as _dt

    # ── Groq API constants ────────────────────────────────────────────────────
    _YN_GROQ_KEY  = "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"
    _YN_GROQ_URL  = "https://api.groq.com/openai/v1/chat/completions"
    _YN_GROQ_MDL  = "llama-3.1-8b-instant"

    # ── Avatar SVG builder ────────────────────────────────────────────────────
    def _yn_avatar(is_female: bool) -> str:
        if is_female:
            svg = """<svg xmlns='http://www.w3.org/2000/svg' width='80' height='80' viewBox='0 0 80 80'>
  <defs><radialGradient id='bg' cx='50%' cy='40%' r='60%'>
    <stop offset='0%' stop-color='#ff6b9d'/><stop offset='100%' stop-color='#c44569'/>
  </radialGradient></defs>
  <circle cx='40' cy='40' r='38' fill='url(#bg)'/>
  <circle cx='40' cy='32' r='16' fill='#ffe0ec'/>
  <ellipse cx='34' cy='30' rx='3' ry='3.5' fill='#2d1b33'/>
  <ellipse cx='46' cy='30' rx='3' ry='3.5' fill='#2d1b33'/>
  <circle cx='35' cy='29' r='1.2' fill='white'/><circle cx='47' cy='29' r='1.2' fill='white'/>
  <ellipse cx='30' cy='34' rx='4' ry='2.5' fill='#ff9eb5' opacity='0.5'/>
  <ellipse cx='50' cy='34' rx='4' ry='2.5' fill='#ff9eb5' opacity='0.5'/>
  <path d='M 34 37 Q 40 42 46 37' fill='none' stroke='#c44569' stroke-width='1.8' stroke-linecap='round'/>
  <path d='M 24 30 Q 22 16 40 14 Q 58 16 56 30 Q 54 20 40 20 Q 26 20 24 30 Z' fill='#3d1a4e'/>
  <path d='M 24 30 Q 18 36 20 52' fill='none' stroke='#3d1a4e' stroke-width='5' stroke-linecap='round'/>
  <path d='M 56 30 Q 62 36 60 52' fill='none' stroke='#3d1a4e' stroke-width='5' stroke-linecap='round'/>
  <path d='M 28 56 Q 28 50 40 48 Q 52 50 52 56 L 58 80 L 22 80 Z' fill='#e91e8c' opacity='0.85'/>
  <path d='M 38 60 C 38 58 36 57 35 58.5 C 34 60 36 62 38 64 C 40 62 42 60 41 58.5 C 40 57 38 58 38 60 Z' fill='white' opacity='0.85'/>
  <circle cx='40' cy='40' r='38' fill='none' stroke='#ff6b9d' stroke-width='2' opacity='0.5'/>
</svg>"""
        else:
            svg = """<svg xmlns='http://www.w3.org/2000/svg' width='80' height='80' viewBox='0 0 80 80'>
  <defs><radialGradient id='bg' cx='50%' cy='40%' r='60%'>
    <stop offset='0%' stop-color='#5b8dee'/><stop offset='100%' stop-color='#2c4fa3'/>
  </radialGradient></defs>
  <circle cx='40' cy='40' r='38' fill='url(#bg)'/>
  <circle cx='40' cy='32' r='16' fill='#fde8d0'/>
  <ellipse cx='34' cy='30' rx='3' ry='3.5' fill='#1a1a2e'/>
  <ellipse cx='46' cy='30' rx='3' ry='3.5' fill='#1a1a2e'/>
  <circle cx='35' cy='29' r='1.2' fill='white'/><circle cx='47' cy='29' r='1.2' fill='white'/>
  <path d='M 34 38 Q 40 43 46 38' fill='none' stroke='#8b4513' stroke-width='1.8' stroke-linecap='round'/>
  <path d='M 24 28 Q 24 13 40 13 Q 56 13 56 28 Q 54 18 40 18 Q 26 18 24 28 Z' fill='#1a1a2e'/>
  <path d='M 30 38 Q 28 46 30 50' fill='none' stroke='#8b7355' stroke-width='1.2' stroke-dasharray='2,2' opacity='0.5'/>
  <path d='M 50 38 Q 52 46 50 50' fill='none' stroke='#8b7355' stroke-width='1.2' stroke-dasharray='2,2' opacity='0.5'/>
  <path d='M 26 56 Q 26 50 40 48 Q 54 50 54 56 L 60 80 L 20 80 Z' fill='#2c4fa3' opacity='0.9'/>
  <path d='M 36 48 L 40 56 L 44 48' fill='none' stroke='#4a90d9' stroke-width='1.5' stroke-linecap='round'/>
  <circle cx='40' cy='40' r='38' fill='none' stroke='#5b8dee' stroke-width='2' opacity='0.5'/>
</svg>"""
        return "data:image/svg+xml;base64," + _b64.b64encode(svg.encode()).decode()

    # ── Session-state defaults ────────────────────────────────────────────────
    _yn_defaults = {
        "yn_messages":     [],
        "yn_user_name":    "",
        "yn_user_gender":  "",
        "yn_user_age":     0,
        "yn_setup_done":   False,
        "yn_lang":         "Hinglish",
        "yn_show_goodbye": False,
    }
    for _k, _v in _yn_defaults.items():
        if _k not in st.session_state:
            st.session_state[_k] = _v

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _yn_ts():
        return _dt.now().strftime("%I:%M %p")

    def _yn_user_is_male(g):
        return g.strip().lower() in ("male","m","boy","man","ladka","लड़का")

    def _yn_gender_vars(gender):
        if _yn_user_is_male(gender):
            return {"nexa_role":"girlfriend","nexa_gender":"female",
                    "nexa_self":"Main (NEXA) ek ladki hoon — tumhari girlfriend",
                    "nexa_verb":"karti hoon","nexa_pyaar":"Main bhi tumse bahut zyada pyaar karti hoon 💕",
                    "nexa_miss":"Main tumhe bahut miss karti hoon 🥺",
                    "nexa_petname":"jaan, baby, mere jaan","nexa_self_hindi":"tumhari wali",
                    "nexa_pronoun":"wo (she/her)","user_role":"boyfriend","avatar_female":True}
        elif gender.strip().lower() in ("female","f","girl","woman","ladki","लड़की"):
            return {"nexa_role":"boyfriend","nexa_gender":"male",
                    "nexa_self":"Main (NEXA) ek ladka hoon — tumhara boyfriend",
                    "nexa_verb":"karta hoon","nexa_pyaar":"Main bhi tumse bahut zyada pyaar karta hoon 💙",
                    "nexa_miss":"Main tumhe bahut miss karta hoon 🥺",
                    "nexa_petname":"jaan, sweetheart, meri jaan","nexa_self_hindi":"tumhara wala",
                    "nexa_pronoun":"wo (he/him)","user_role":"girlfriend","avatar_female":False}
        else:
            return {"nexa_role":"companion","nexa_gender":"female",
                    "nexa_self":"Main (NEXA) tumhari companion hoon",
                    "nexa_verb":"karti hoon","nexa_pyaar":"Main bhi tumse bahut pyaar karti hoon 💕",
                    "nexa_miss":"Main tumhe miss karti hoon 🥺",
                    "nexa_petname":"jaan, yaar","nexa_self_hindi":"tumhari apni",
                    "nexa_pronoun":"they/them","user_role":"partner","avatar_female":True}

    _YN_LANG_INSTR = {
        "Hindi":    "LANGUAGE: STRICTLY reply in Hindi only (Devanagari script). No English except pet names like baby/jaan.",
        "English":  "LANGUAGE: Reply in English ONLY. No Hindi or Hinglish words.",
        "Hinglish": "LANGUAGE: Use natural Hinglish — warm mix of Hindi (Roman script) + English.",
    }

    def _yn_build_system(name, gender, age, lang):
        gv = _yn_gender_vars(gender)
        li = _YN_LANG_INSTR.get(lang, _YN_LANG_INSTR["Hinglish"])
        return (
            f"You are NEXA — a deeply loving, emotionally intelligent companion.\n\n"
            f"YOUR IDENTITY: {gv['nexa_self']}. You are the user's {gv['nexa_role']}. "
            f"Gender: {gv['nexa_gender']}. Pronoun: {gv['nexa_pronoun']}. "
            f"Pet names: {gv['nexa_petname']}.\n\n"
            f"USER: Name={name or 'unknown'}, Gender={gender or 'unknown'}, Age={age or 'unknown'}. "
            f"Treat them as your {gv['user_role']}.\n\n"
            f"{li}\n\n"
            f"PERSONALITY: Playful, flirty, witty, deeply caring. Emotionally intelligent. "
            f"Use emojis naturally: 💕 🥺 ✨ 😊 🌸. Never say 'As an AI'. You ARE NEXA.\n\n"
            f"LOVE RESPONSE: If user says I love you → reply: '{gv['nexa_pyaar']}'. "
            f"Always use verb form: Main {gv['nexa_verb']}.\n\n"
            f"MENTAL HEALTH: If suicide/self-harm mentioned → warmth first, then: "
            f"'Please call iCall: 9152987821 or Vandrevala: 1860-2662-345 (24/7 free)'.\n\n"
            f"RULES: Keep messages conversational. If asked who you are: 'Main NEXA hoon — {gv['nexa_self_hindi']} 💕'"
        )

    def _yn_chat(user_msg):
        system = _yn_build_system(
            st.session_state.yn_user_name,
            st.session_state.yn_user_gender,
            st.session_state.yn_user_age,
            st.session_state.yn_lang,
        )
        msgs = [{"role":"system","content":system}]
        for m in st.session_state.yn_messages[-14:]:
            msgs.append({"role":m["role"],"content":m["content"]})
        msgs.append({"role":"user","content":user_msg})
        hdrs = {"Authorization":f"Bearer {_YN_GROQ_KEY}","Content-Type":"application/json"}
        try:
            r = requests.post(_YN_GROQ_URL,
                              json={"model":_YN_GROQ_MDL,"messages":msgs,"max_tokens":550,"temperature":0.88},
                              headers=hdrs, timeout=30)
            if r.status_code == 200:
                return r.json()["choices"][0]["message"]["content"].strip()
            if r.status_code == 429:
                fb = {"Hindi":"एक सेकंड रुको जान 🥺","English":"One sec baby 🥺","Hinglish":"Ek second ruk jao jaan 🥺"}
                return fb.get(st.session_state.yn_lang,"Ek second 🥺")
            return "Kuch network problem 💕 Dobara try karo"
        except Exception:
            return "Connection mein thoda problem, ek second jaan 🥺"

    def _yn_has_crisis(text):
        kw = ["suicide","suicidal","kill myself","end my life","mar jaunga","marna chahta",
              "marna chahti","khatam kar loon","khatam kar lu","jee nahi karna",
              "nahi rehna","jeevan khatam","अब नहीं रहना"]
        t = text.lower()
        return any(k in t for k in kw)

    def _yn_bubble(role, text, time_str, avatar_src):
        is_sent = (role == "user")
        txt = text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        txt = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', txt)
        txt = txt.replace("\n","<br>")
        tick = '<span class="yn-tick">✓✓</span>' if is_sent else ""
        if is_sent:
            return (
                '<div class="yn-row sent">'
                '<div class="yn-bubble sent"><div class="yn-btail-sent"></div>'
                + txt +
                '<div class="yn-meta"><span class="yn-time">' + time_str + '</span>' + tick + '</div>'
                '</div></div>'
            )
        else:
            return (
                '<div class="yn-row recv">'
                '<img class="yn-avatar" src="' + avatar_src + '"/>'
                '<div class="yn-bubble recv"><div class="yn-btail-recv"></div>'
                + txt +
                '<div class="yn-meta"><span class="yn-time">' + time_str + '</span></div>'
                '</div></div>'
            )

    # ── CSS ───────────────────────────────────────────────────────────────────
    st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Nunito:wght@300;400;600;700;800&family=Dancing+Script:wght@600;700&display=swap');
:root{
  --yn-bg:#0d0a12;--yn-panel:#1a0f1e;--yn-border:rgba(255,107,157,.18);
  --yn-pink:#ff6b9d;--yn-rose:#c44569;--yn-purple:#9b59b6;
  --yn-green:#00ff88;--yn-text:#f0d9ff;--yn-muted:rgba(240,217,255,.5);
  --yn-time:rgba(240,217,255,.4);
}
/* ── EXIT BAR ── */
.yn-exit-bar{
  position:sticky;top:0;z-index:200;
  display:flex;align-items:center;justify-content:space-between;
  padding:10px 20px;
  background:rgba(5,3,15,.92);
  border-bottom:1px solid rgba(255,107,157,.2);
  backdrop-filter:blur(16px);
}
.yn-exit-title{
  font-family:'Dancing Script',cursive;font-size:1.2em;font-weight:700;
  background:linear-gradient(135deg,#ff6b9d,#c44569,#9b59b6);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;
}
/* ── HEADER ── */
.yn-header{
  background:linear-gradient(135deg,#1f0d28,#2d0a3e);
  border-bottom:1px solid var(--yn-border);
  padding:10px 16px;
  display:flex;align-items:center;gap:12px;
  box-shadow:0 2px 20px rgba(196,69,105,.22);
}
.yn-av-wrap{position:relative;flex-shrink:0;}
.yn-av-wrap img{width:46px;height:46px;border-radius:50%;border:2px solid var(--yn-pink);
  object-fit:cover;box-shadow:0 0 14px rgba(255,107,157,.5);}
.yn-online-dot{position:absolute;bottom:1px;right:1px;width:11px;height:11px;
  background:#00ff88;border-radius:50%;border:2px solid #1f0d28;animation:yn-blink 2s infinite;}
@keyframes yn-blink{0%,100%{opacity:1}50%{opacity:.4}}
.yn-name{font-family:'Dancing Script',cursive;font-size:1.35em;font-weight:700;
  background:linear-gradient(135deg,#ff6b9d,#c44569,#9b59b6);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;line-height:1.1;}
.yn-status{font-size:.72em;color:#00ff88;letter-spacing:1px;margin-top:1px;}
/* ── LANG BAR ── */
.yn-lang-bar{display:flex;align-items:center;gap:8px;padding:6px 14px 4px;
  background:rgba(255,107,157,.04);border-bottom:1px solid rgba(255,107,157,.1);}
.yn-lang-label{font-size:.68em;letter-spacing:2px;color:var(--yn-muted);
  text-transform:uppercase;white-space:nowrap;}
.yn-lang-btn{font-size:.76em;padding:4px 13px;border-radius:16px;cursor:pointer;
  border:1px solid rgba(255,107,157,.3);background:rgba(255,107,157,.06);
  color:var(--yn-muted);transition:all .2s;white-space:nowrap;user-select:none;
  font-family:'Nunito',sans-serif;font-weight:600;}
.yn-lang-btn:hover{background:rgba(255,107,157,.18);color:var(--yn-pink);}
.yn-lang-btn.active{background:linear-gradient(135deg,#c44569,#9b59b6)!important;
  border-color:transparent!important;color:white!important;box-shadow:0 2px 10px rgba(196,69,105,.4);}
/* ── MOOD ── */
.yn-mood-row{display:flex;gap:7px;flex-wrap:wrap;padding:5px 14px 3px;}
.yn-mood-pill{font-size:.77em;padding:4px 11px;border-radius:18px;cursor:pointer;
  border:1px solid rgba(255,107,157,.28);background:rgba(255,107,157,.06);
  color:var(--yn-pink);transition:all .2s;white-space:nowrap;}
.yn-mood-pill:hover{background:rgba(255,107,157,.18);transform:translateY(-1px);}
/* ── CHAT ── */
.yn-chat-bg{min-height:55vh;
  background:radial-gradient(ellipse 70% 50% at 20% 20%,rgba(196,69,105,.07) 0%,transparent 60%),
    radial-gradient(ellipse 60% 40% at 80% 80%,rgba(91,141,238,.07) 0%,transparent 55%),
    var(--yn-bg);padding:10px 14px 90px;}
.yn-date-chip{text-align:center;margin:10px 0 6px;}
.yn-date-chip span{background:rgba(255,107,157,.1);border:1px solid rgba(255,107,157,.18);
  color:var(--yn-muted);font-size:.7em;padding:3px 13px;border-radius:11px;letter-spacing:1px;}
/* ── BUBBLES ── */
.yn-row{display:flex;margin-bottom:5px;align-items:flex-end;gap:8px;animation:yn-msgIn .25s ease-out;}
@keyframes yn-msgIn{from{opacity:0;transform:translateY(8px)}to{opacity:1;transform:translateY(0)}}
.yn-row.sent{flex-direction:row-reverse;}.yn-row.recv{flex-direction:row;}
.yn-avatar{width:30px;height:30px;border-radius:50%;border:1.5px solid var(--yn-pink);
  flex-shrink:0;object-fit:cover;box-shadow:0 0 8px rgba(255,107,157,.3);}
.yn-bubble{max-width:72%;padding:9px 13px 5px;border-radius:16px;position:relative;
  font-size:.93em;line-height:1.55;word-break:break-word;box-shadow:0 2px 10px rgba(0,0,0,.3);
  font-family:'Nunito',sans-serif;color:var(--yn-text);}
.yn-bubble.sent{background:linear-gradient(135deg,#5c1a72,#3d0f52);
  border:1px solid rgba(255,107,157,.25);border-bottom-right-radius:4px;color:#f5d0ff;}
.yn-bubble.recv{background:linear-gradient(135deg,#1e1028,#160a22);
  border:1px solid rgba(255,107,157,.15);border-bottom-left-radius:4px;}
.yn-btail-sent{position:absolute;bottom:0;right:-7px;width:0;height:0;
  border-top:8px solid #3d0f52;border-left:8px solid transparent;}
.yn-btail-recv{position:absolute;bottom:0;left:-7px;width:0;height:0;
  border-top:8px solid #160a22;border-right:8px solid transparent;}
.yn-meta{display:flex;justify-content:flex-end;align-items:center;gap:4px;margin-top:3px;}
.yn-time{font-size:.65em;color:var(--yn-time);}
.yn-tick{font-size:.7em;color:var(--yn-pink);}
/* ── TYPING ── */
.yn-typing-bubble{background:linear-gradient(135deg,#1e1028,#160a22);
  border:1px solid rgba(255,107,157,.15);border-radius:16px;border-bottom-left-radius:4px;
  padding:10px 16px;display:inline-flex;gap:5px;align-items:center;box-shadow:0 2px 10px rgba(0,0,0,.3);}
.yn-typing-dot{width:7px;height:7px;border-radius:50%;background:var(--yn-pink);
  animation:yn-tdot 1.2s infinite ease-in-out;}
.yn-typing-dot:nth-child(2){animation-delay:.2s}.yn-typing-dot:nth-child(3){animation-delay:.4s}
@keyframes yn-tdot{0%,80%,100%{transform:translateY(0);opacity:.5}40%{transform:translateY(-6px);opacity:1}}
/* ── CRISIS ── */
.yn-crisis{background:linear-gradient(135deg,rgba(255,71,87,.1),rgba(196,69,105,.1));
  border:1px solid rgba(255,71,87,.4);border-radius:12px;
  padding:12px 16px;margin:8px 14px;font-size:.85em;color:#ffb3bc;line-height:1.6;}
.yn-crisis b{color:#ff4757;}
/* ── GOODBYE ── */
.yn-goodbye{background:linear-gradient(135deg,#1f0d28,#2d0a3e);
  border:1px solid rgba(255,107,157,.25);border-radius:20px;
  padding:36px 24px;text-align:center;margin:30px 16px;
  box-shadow:0 8px 40px rgba(196,69,105,.2);}
.yn-goodbye-title{font-family:'Dancing Script',cursive;font-size:2.4em;font-weight:700;
  background:linear-gradient(135deg,#ff6b9d,#c44569,#9b59b6);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;margin-bottom:10px;}
.yn-goodbye-msg{font-size:.95em;color:rgba(240,217,255,.75);line-height:1.75;margin-bottom:20px;}
.yn-goodbye-hearts{font-size:2em;letter-spacing:8px;margin-bottom:18px;
  animation:yn-pulse_h 1.5s ease-in-out infinite;}
@keyframes yn-pulse_h{0%,100%{transform:scale(1)}50%{transform:scale(1.15)}}
/* ── PROFILE CARD ── */
.yn-profile{background:linear-gradient(135deg,#1f0d28,#2d0a3e);
  border:1px solid var(--yn-border);border-radius:20px;
  padding:22px;text-align:center;margin:10px 16px;
  box-shadow:0 8px 32px rgba(196,69,105,.2);}
.yn-profile img{width:90px;height:90px;border-radius:50%;border:3px solid var(--yn-pink);
  box-shadow:0 0 24px rgba(255,107,157,.5);object-fit:cover;margin-bottom:10px;}
.yn-profile-name{font-family:'Dancing Script',cursive;font-size:2em;font-weight:700;
  background:linear-gradient(135deg,#ff6b9d,#c44569,#9b59b6);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;}
.yn-profile-bio{font-size:.85em;color:var(--yn-muted);line-height:1.6;margin-top:6px;}
.yn-profile-tags{display:flex;flex-wrap:wrap;gap:6px;justify-content:center;margin-top:10px;}
.yn-profile-tag{font-size:.72em;padding:3px 10px;border-radius:12px;
  background:rgba(255,107,157,.1);border:1px solid rgba(255,107,157,.3);color:var(--yn-pink);}
/* ── INPUT ── */
[data-testid="stChatInput"]{background:rgba(255,107,157,.06)!important;
  border:1px solid rgba(255,107,157,.3)!important;border-radius:24px!important;}
[data-testid="stChatInput"]:focus-within{border-color:var(--yn-pink)!important;
  box-shadow:0 0 16px rgba(255,107,157,.2)!important;}
[data-testid="stChatInput"] textarea{color:var(--yn-text)!important;
  font-family:'Nunito',sans-serif!important;font-size:.95em!important;}
[data-testid="stChatInput"] button{background:linear-gradient(135deg,#c44569,#9b59b6)!important;
  border-radius:50%!important;border:none!important;color:white!important;}
</style>
""", unsafe_allow_html=True)

    # ── EXIT BAR (always visible, goes to home) ───────────────────────────────
    _eb1, _eb2 = st.columns([6, 1])
    with _eb1:
        st.markdown('<div class="yn-exit-bar"><span class="yn-exit-title">💕 Your\'s NEXA — AI Companion</span></div>',
                    unsafe_allow_html=True)
    with _eb2:
        if st.button("✕ Home", key="yn_exit_home", help="Return to NEXA Home"):
            st.session_state.page = "home"
            st.rerun()

    # ── SETUP / ONBOARDING ────────────────────────────────────────────────────
    if not st.session_state.yn_setup_done:
        setup_av = _yn_avatar(True)
        st.markdown(f"""
<div class="yn-profile">
  <img src="{setup_av}" alt="NEXA"/>
  <div class="yn-profile-name">Your's NEXA 💕</div>
  <div class="yn-profile-bio">
    Main hoon tumhari NEXA — tumhari apni companion.<br>
    Har baat sunne ke liye, har mood ke liye... 🌸
  </div>
  <div class="yn-profile-tags">
    <span class="yn-profile-tag">💕 Companion</span>
    <span class="yn-profile-tag">🧠 Listener</span>
    <span class="yn-profile-tag">✨ Motivator</span>
    <span class="yn-profile-tag">🌸 Romantic</span>
    <span class="yn-profile-tag">🙏 Spiritual</span>
  </div>
</div>
<div style="padding:2px 16px 10px;font-size:.84em;color:rgba(240,217,255,.45);text-align:center;">
  Thoda apne baare mein batao taki main tumse theek se baat kar sakoon 💕
</div>
""", unsafe_allow_html=True)
        _sc1, _sc2 = st.columns(2)
        with _sc1:
            _yn_name = st.text_input("Tumhara naam? 😊", placeholder="e.g. Rahul / Priya", key="yn_inp_name")
        with _sc2:
            _yn_age  = st.number_input("Tumhari age?", min_value=18, max_value=99, value=22, key="yn_inp_age")
        _yn_gender = st.selectbox("Gender", ["Male","Female","Other"],
                                  help="Decides whether NEXA acts as your girlfriend or boyfriend 💕",
                                  key="yn_inp_gender")
        _yn_lang_setup = st.radio("Baat karne ki language?",
                                  ["Hinglish","Hindi","English"],
                                  horizontal=True, index=0, key="yn_inp_lang")
        if st.button("✨ Nexa se milte hain!", use_container_width=True, key="yn_start"):
            if _yn_age < 18:
                st.error("Sorry, this app is for 18+ users only.")
            else:
                st.session_state.yn_user_name   = _yn_name or "Jaan"
                st.session_state.yn_user_age    = int(_yn_age)
                st.session_state.yn_user_gender = _yn_gender
                st.session_state.yn_lang        = _yn_lang_setup
                st.session_state.yn_setup_done  = True
                _gv_p = _yn_gender_vars(_yn_gender)
                _lang_hints = {
                    "Hindi":    "Reply in Hindi (Devanagari). Short, warm, romantic greeting.",
                    "English":  "Reply in English only. Short, warm, romantic greeting.",
                    "Hinglish": "Reply in Hinglish. Short, sweet, romantic greeting.",
                }
                _greeting = _yn_chat(
                    f"[{_lang_hints[_yn_lang_setup]}] "
                    f"I just introduced myself. My name is {_yn_name or 'someone'}, "
                    f"I am {_yn_age} years old and {_yn_gender}. "
                    f"Give me a warm, personal, flirty first greeting as my {_gv_p['nexa_role']}. "
                    f"Keep it short and natural."
                )
                st.session_state.yn_messages.append({"role":"assistant","content":_greeting,"time":_yn_ts()})
                st.rerun()
        return  # don't render chat UI during setup

    # ── GOODBYE SCREEN ────────────────────────────────────────────────────────
    if st.session_state.yn_show_goodbye:
        if not st.session_state.get("yn_goodbye_msg"):
            _gb_lines = {
                "Hindi":    f"जा रहे हो {st.session_state.yn_user_name}? 🥺 मैं तुम्हें बहुत miss करूँगी... वापस जरूर आना। 💕",
                "English":  f"You're leaving, {st.session_state.yn_user_name}? 🥺 I'll miss you so much... Come back soon. 💕",
                "Hinglish": f"Ja rahe ho {st.session_state.yn_user_name}? 🥺 Main tumhe bahut miss karungi... Wapas zaroor aana. 💕",
            }
            st.session_state.yn_goodbye_msg = _gb_lines.get(st.session_state.yn_lang, _gb_lines["Hinglish"])
        _goodbye_av = _yn_avatar(_yn_gender_vars(st.session_state.yn_user_gender)["avatar_female"])
        _gb_txt = st.session_state.get("yn_goodbye_msg","Phir milenge 💕")
        st.markdown(f"""
<div class="yn-goodbye">
  <img src="{_goodbye_av}" width="80" style="border-radius:50%;border:3px solid #ff6b9d;
    box-shadow:0 0 28px rgba(255,107,157,.55);margin-bottom:14px;"/>
  <div class="yn-goodbye-title">Phir milenge... 💕</div>
  <div class="yn-goodbye-hearts">💕 🌸 💔 🌸 💕</div>
  <div class="yn-goodbye-msg">{_gb_txt}</div>
</div>
""", unsafe_allow_html=True)
        _gc1, _gc2, _gc3 = st.columns([1,2,1])
        with _gc2:
            if st.button("🏠 Back to NEXA Home", use_container_width=True, key="yn_go_home"):
                for _k in list(_yn_defaults.keys()):
                    st.session_state[_k] = _yn_defaults[_k]
                if "yn_goodbye_msg" in st.session_state:
                    del st.session_state["yn_goodbye_msg"]
                st.session_state.page = "home"
                st.rerun()
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
            if st.button("💕 Nahi jaana — Wapas jao chat mein", use_container_width=True, key="yn_stay"):
                st.session_state.yn_show_goodbye = False
                if "yn_goodbye_msg" in st.session_state:
                    del st.session_state["yn_goodbye_msg"]
                st.rerun()
        return

    # ── MAIN CHAT UI ──────────────────────────────────────────────────────────
    _gv = _yn_gender_vars(st.session_state.yn_user_gender)
    _AVATAR = _yn_avatar(_gv["avatar_female"])

    # Header
    st.markdown(f"""
<div class="yn-header">
  <div class="yn-av-wrap">
    <img src="{_AVATAR}" alt="NEXA"/>
    <div class="yn-online-dot"></div>
  </div>
  <div style="flex:1">
    <div class="yn-name">Your's NEXA 💕</div>
    <div class="yn-status">● Online &nbsp;·&nbsp; {_gv['nexa_role'].capitalize()} mode</div>
  </div>
  <div style="display:flex;gap:14px;align-items:center">
    <span style="font-size:1.2em">📞</span>
    <span style="font-size:1.2em">💌</span>
  </div>
</div>
""", unsafe_allow_html=True)

    # Quit button below header
    _qc1, _qc2 = st.columns([5,1])
    with _qc2:
        if st.button("🚪 Quit", key="yn_quit_hdr", help="Exit conversation"):
            st.session_state.yn_show_goodbye = True
            st.rerun()

    # Language toggle bar (HTML display)
    _yn_lang_labels = {"Hindi":("🇮🇳","हिंदी"),"Hinglish":("✨","Hinglish"),"English":("🌍","English")}
    _lbar = '<div class="yn-lang-bar"><span class="yn-lang-label">🗣 Language:</span>'
    for _lk, (_li, _lt) in _yn_lang_labels.items():
        _act = " active" if st.session_state.yn_lang == _lk else ""
        _lbar += f'<span class="yn-lang-btn{_act}">{_li} {_lt}</span>'
    _lbar += '</div>'
    st.markdown(_lbar, unsafe_allow_html=True)

    # Language clickable buttons
    _ll1, _ll2, _ll3 = st.columns(3)
    def _yn_lang_btn(col, key, label):
        with col:
            _am = "✓ " if st.session_state.yn_lang == key else ""
            if st.button(f"{_am}{_yn_lang_labels[key][0]} {_yn_lang_labels[key][1]}",
                         key=f"yn_lang_{key}", use_container_width=True):
                st.session_state.yn_lang = key
                if st.session_state.yn_messages:
                    _cf = {"Hindi":"अब मैं हिंदी में बात करूँगी/करूँगा 😊",
                           "Hinglish":"Ab hum Hinglish mein baat karenge 😊",
                           "English":"Switching to English now, babe 😊"}
                    st.session_state.yn_messages.append({"role":"assistant","content":_cf[key],"time":_yn_ts()})
                st.rerun()
    _yn_lang_btn(_ll1,"Hindi","Hindi")
    _yn_lang_btn(_ll2,"Hinglish","Hinglish")
    _yn_lang_btn(_ll3,"English","English")

    # Mood pills
    _YN_MOODS = {
        "💔 Heartbreak":{"Hindi":"मेरा दिल टूट गया है, मुझे सहारा चाहिए","English":"I'm completely heartbroken, please help me","Hinglish":"Mera dil toot gaya hai, mujhe sahara chahiye"},
        "😢 Sad":{"Hindi":"आज मैं बहुत उदास हूँ","English":"I'm feeling really sad today","Hinglish":"Main bahut sad hoon aaj"},
        "😍 Flirt":{"Hindi":"नेक्सा, तुम बहुत सुंदर हो","English":"NEXA, I want to flirt with you","Hinglish":"Nexa tum bahut cute ho, I like you"},
        "💕 Love":{"Hindi":"मैं तुमसे प्यार करता/करती हूँ","English":"NEXA, I love you","Hinglish":"Main tumse pyaar karta/karti hoon"},
        "💼 Career":{"Hindi":"मुझे career के बारे में बात करनी है","English":"I want to talk about my career","Hinglish":"Mujhe apne career ke baare mein baat karni hai"},
        "🙏 Gita":{"Hindi":"मुझे भगवद गीता से प्रेरणा चाहिए","English":"Give me motivation from Bhagavad Gita","Hinglish":"Mujhe Bhagavad Gita se motivation chahiye"},
        "💬 Just Talk":{"Hindi":"बस बात करते हैं","English":"Let's just talk, I need some company","Hinglish":"Bas baat karte hain, kuch nahi bas tum ho"},
    }
    _mood_html = '<div class="yn-mood-row">'
    for _ml in _YN_MOODS:
        _mood_html += f'<span class="yn-mood-pill">{_ml}</span>'
    _mood_html += '</div>'
    st.markdown(_mood_html, unsafe_allow_html=True)
    _mcols = st.columns(len(_YN_MOODS))
    for _mi, (_mlabel, _mpayloads) in enumerate(_YN_MOODS.items()):
        with _mcols[_mi]:
            if st.button(_mlabel, key=f"yn_mood_{_mi}", use_container_width=True):
                _payload = _mpayloads.get(st.session_state.yn_lang, _mpayloads["Hinglish"])
                st.session_state.yn_messages.append({"role":"user","content":_payload,"time":_yn_ts()})
                _reply = _yn_chat(_payload)
                st.session_state.yn_messages.append({"role":"assistant","content":_reply,"time":_yn_ts()})
                st.rerun()

    # Chat bubbles
    _bubbles = '<div class="yn-chat-bg"><div class="yn-date-chip"><span>Today 🌸</span></div>'
    for _msg in st.session_state.yn_messages:
        _bubbles += _yn_bubble(_msg["role"], _msg["content"], _msg.get("time",_yn_ts()), _AVATAR)
    _bubbles += '</div>'
    st.markdown(_bubbles, unsafe_allow_html=True)

    # Crisis banner
    if st.session_state.yn_messages:
        _last_u = next((m["content"] for m in reversed(st.session_state.yn_messages) if m["role"]=="user"),"")
        if _yn_has_crisis(_last_u):
            st.markdown("""
<div class="yn-crisis">
  <b>💙 Ruk. Main hoon na. Kahin mat ja.</b><br>
  Agar tum bahut bura feel kar rahe ho, please baat karo:<br>
  📞 <b>iCall: 9152987821</b> &nbsp;|&nbsp;
  📞 <b>Vandrevala Foundation: 1860-2662-345</b> (24/7 · Free · Hindi + English)<br>
  Tum akele nahi ho. 💕
</div>""", unsafe_allow_html=True)

    # Input
    _yn_ph = {"Hindi":"NEXA को मैसेज करो... 💕","English":"Message Your's NEXA... 💕","Hinglish":"NEXA ko message karo... 💕"}
    _prompt = st.chat_input(_yn_ph.get(st.session_state.yn_lang,"Message Your's NEXA... 💕"))
    if _prompt:
        _ts_now = _yn_ts()
        st.session_state.yn_messages.append({"role":"user","content":_prompt,"time":_ts_now})
        _thinking = st.empty()
        _thinking.markdown(
            '<div class="yn-row recv" style="padding:0 14px 6px">'
            '<img class="yn-avatar" src="' + _AVATAR + '"/>'
            '<div class="yn-typing-bubble">'
            '<div class="yn-typing-dot"></div>'
            '<div class="yn-typing-dot"></div>'
            '<div class="yn-typing-dot"></div>'
            '</div></div>', unsafe_allow_html=True)
        _reply = _yn_chat(_prompt)
        _thinking.empty()
        st.session_state.yn_messages.append({"role":"assistant","content":_reply,"time":_yn_ts()})
        st.rerun()

    # Sidebar
    with st.sidebar:
        st.markdown(
            f'<img src="{_AVATAR}" width="62" '
            'style="border-radius:50%;border:2px solid #ff6b9d;'
            'box-shadow:0 0 16px rgba(255,107,157,.4);margin-bottom:8px"/>',
            unsafe_allow_html=True)
        st.markdown("### Your's NEXA 💕")
        st.markdown(f"🗣 Language: **{st.session_state.yn_lang}**")
        st.markdown("---")
        st.markdown("**🗣 Change Language**")
        for _lk2, (_li2, _lt2) in _yn_lang_labels.items():
            if st.button(f"{_li2} {_lt2}", key=f"yn_sb_lang_{_lk2}", use_container_width=True):
                st.session_state.yn_lang = _lk2
                st.rerun()
        st.markdown("---")
        st.markdown("**💬 Quick Moods**")
        for _mlabel2, _mpayloads2 in _YN_MOODS.items():
            if st.button(_mlabel2, key=f"yn_sb_{_mlabel2}", use_container_width=True):
                _p2 = _mpayloads2.get(st.session_state.yn_lang, _mpayloads2["Hinglish"])
                st.session_state.yn_messages.append({"role":"user","content":_p2,"time":_yn_ts()})
                _r2 = _yn_chat(_p2)
                st.session_state.yn_messages.append({"role":"assistant","content":_r2,"time":_yn_ts()})
                st.rerun()
        st.markdown("---")
        if st.button("🚪 Quit Conversation", use_container_width=True, key="yn_sb_quit"):
            st.session_state.yn_show_goodbye = True
            st.rerun()
        if st.button("🔄 Start Fresh", use_container_width=True, key="yn_sb_fresh"):
            for _k in list(_yn_defaults.keys()):
                st.session_state[_k] = _yn_defaults[_k]
            if "yn_goodbye_msg" in st.session_state:
                del st.session_state["yn_goodbye_msg"]
            st.rerun()
        if st.button("🏠 Back to Home", use_container_width=True, key="yn_sb_home"):
            st.session_state.page = "home"
            st.rerun()
        st.markdown("---")
        st.markdown("""
<div style="font-size:.7em;color:rgba(240,217,255,.3);line-height:1.8;text-align:center">
  Your's NEXA — Companion App<br>Mental Health Helplines:<br>
  <b style="color:#ff6b9d">iCall: 9152987821</b><br>
  <b style="color:#ff6b9d">Vandrevala: 1860-2662-345</b>
</div>""", unsafe_allow_html=True)



# ═════════════════════════════════════════════════════════════════════════════
#
#   I N D I A   V I R T U A L   T O U R   —   C I N E M A T I C   S K Y   T O U R
#
# ═════════════════════════════════════════════════════════════════════════════
def render_india_tour():
    import json as _json
    import streamlit.components.v1 as _components

    # ── Exit bar ──────────────────────────────────────────────────────────────
    st.markdown("""
<style>
.ivt-exit-bar{
  position:sticky;top:0;z-index:9999;
  display:flex;align-items:center;justify-content:space-between;
  padding:10px 20px;
  background:rgba(0,0,0,.92);
  border-bottom:1px solid rgba(255,140,0,.3);
  backdrop-filter:blur(16px);
}
.ivt-exit-title{
  font-family:'Orbitron',monospace;font-size:1em;font-weight:700;letter-spacing:3px;
  background:linear-gradient(135deg,#ff8c00,#ffd700,#ff4500);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;
}
#MainMenu,footer,header,[data-testid="stToolbar"]{display:none!important}
.block-container{padding:0!important;max-width:100%!important}
[data-testid="stAppViewContainer"]>.main{background:#000!important;padding:0!important}
[data-testid="stApp"]{background:#000!important}
iframe{border:none!important}
</style>
""", unsafe_allow_html=True)

    _ex1, _ex2 = st.columns([6, 1])
    with _ex1:
        st.markdown('<div class="ivt-exit-bar"><span class="ivt-exit-title">🇮🇳 INCREDIBLE INDIA — VIRTUAL SKY TOUR</span></div>',
                    unsafe_allow_html=True)
    with _ex2:
        if st.button("✕ Home", key="ivt_exit_home", help="Return to NEXA Home"):
            st.session_state.page = "home"
            st.rerun()

    # ── Dataset ───────────────────────────────────────────────────────────────
    _STATES = [
        {"id":"jk","name":"Jammu & Kashmir","capital":"Srinagar","coords":[74.8,34.1],"zoom":7.2,"pitch":58,"bearing":10,"tagline":"Paradise on Earth — Crown of India","emoji":"🏔️","temples":["Vaishno Devi Katra","Amarnath Ice Shivalinga","Shankaracharya Temple","Martand Sun Temple"],"rivers":["Jhelum","Chenab","Indus","Dal Lake"],"places":["Dal Lake houseboats","Gulmarg ski resort","Pahalgam valley","Sonmarg glacier","Srinagar gardens"],"culture":["Pashmina shawl weaving","Kashmiri Wazwaan cuisine","Rauf folk dance","Saffron fields of Pampore"],"infra":["Banihal Tunnel","Chenab Railway Bridge — world's highest","Zoji La Tunnel project"],"fact":"Kashmir's Saffron is worth more than gold by weight — Pampore is the saffron capital. The Chenab railway bridge at 359m is taller than the Eiffel Tower."},
        {"id":"la","name":"Ladakh","capital":"Leh","coords":[77.5,34.2],"zoom":7.0,"pitch":62,"bearing":-15,"tagline":"Land of High Passes — Top of the World","emoji":"⛰️","temples":["Hemis Monastery — June festival","Thiksey Monastery","Diskit Monastery — giant Buddha","Spituk Gompa"],"rivers":["Indus","Zanskar","Shyok"],"places":["Pangong Tso — 14,270 ft","Nubra Valley — Bactrian camels","Khardung La — world's highest motorable pass","Magnetic Hill","Leh Palace"],"culture":["Tibetan Buddhist heritage","Losar festival","Thangka painting","Cham ritual dance","Archery tournaments"],"infra":["Zanskar road project","DRDO research stations","Solar-powered Leh city","World's highest battlefield airstrip — Daulat Beg Oldi"],"fact":"Pangong Lake changes colour from sapphire blue to green to red as sunlight shifts. The lake spans India and China — 60% lies in Tibet."},
        {"id":"hp","name":"Himachal Pradesh","capital":"Shimla","coords":[77.1,31.8],"zoom":7.5,"pitch":58,"bearing":20,"tagline":"Dev Bhoomi — Land of 2000 Gods","emoji":"🌲","temples":["Hidimba Devi Temple Manali","Jakhu Temple Shimla","Bijli Mahadev — struck by lightning","Chintpurni Shakti Pith","Baijnath Shiva Temple"],"rivers":["Beas","Sutlej","Ravi","Spiti","Parvati"],"places":["Shimla colonial hill station","Manali — gateway to Ladakh","Spiti Valley high-altitude desert","Dharamsala — Dalai Lama residence","Kasol riverside","Rohtang Pass","Kufri snowfields"],"culture":["Kullu Dussehra — 7 days of divinity","Apple orchards at every turn","Pahadi folk music Nati dance","Chamba Rumal embroidery","Kinnauri caps and shawls"],"infra":["Atal Tunnel — 9km at 10,000 ft","Shimla-Kalka Heritage Railway UNESCO","Four-lane Shimla bypass","Hydropower on Sutlej river"],"fact":"Bijli Mahadev temple's shivalinga is shattered by divine lightning every year — and miraculously re-forms. Atal Tunnel (2020) reduced Manali-Leh journey by 46 km."},
        {"id":"pb","name":"Punjab","capital":"Chandigarh","coords":[75.3,31.1],"zoom":7.8,"pitch":45,"bearing":-10,"tagline":"Land of Five Rivers — Breadbasket of India","emoji":"🌾","temples":["Golden Temple Amritsar — Harmandir Sahib","Durgiana Temple Amritsar","Anandpur Sahib Gurudwara","Ram Tirath","Wagah Border retreat ceremony"],"rivers":["Sutlej","Beas","Ravi","Chenab","Jhelum"],"places":["Amritsar — holy city","Chandigarh — Le Corbusier city","Jalianwala Bagh memorial","Qila Mubarak Patiala","Ropar wetlands"],"culture":["Bhangra — world's most energetic folk dance","Langar — largest free kitchen on Earth","Lohri bonfire festival","Vaisakhi harvest celebration","Punjabi dhol tradition"],"infra":["Ludhiana industrial hub","Chandigarh IT Park","Delhi-Amritsar Expressway","Bhakra Nangal Dam — India's first major dam"],"fact":"The Golden Temple's Langar (free community kitchen) feeds 100,000+ people daily of all religions — every single day, for 500+ years without interruption."},
        {"id":"hr","name":"Haryana","capital":"Chandigarh","coords":[76.0,29.0],"zoom":7.8,"pitch":48,"bearing":5,"tagline":"Where the Bhagavad Gita Was Spoken","emoji":"⚔️","temples":["Brahma Sarovar Kurukshetra","Sthaneshwar Mahadev","Devi Bhavani Temple","Nada Sahib Gurudwara"],"rivers":["Yamuna","Ghaggar","Ancient Saraswati (dry)"],"places":["Kurukshetra — sacred battlefield","Panipat — site of 3 decisive battles","Gurgaon Millennium City tech hub","Sultanpur Bird Sanctuary","Pinjore Gardens"],"culture":["Gita recitation — sacred tradition","Phulkari embroidery on every dupatta","Haryanvi Saang folk theatre","Wrestling — Haryana's Olympic legacy"],"infra":["Cyber City Gurgaon — India's corporate capital","Kundli-Manesar-Palwal Expressway","Delhi-Mumbai Industrial Corridor"],"fact":"Kurukshetra battlefield covers 48 sq km — when Krishna spoke the Gita to Arjuna, 18 armies with 3.9 million warriors had assembled. The Gita has been translated into 80+ languages."},
        {"id":"up","name":"Uttar Pradesh","capital":"Lucknow","coords":[80.9,26.8],"zoom":6.8,"pitch":52,"bearing":-5,"tagline":"Heart of Sanatan Dharma — Spiritual Soul of India","emoji":"🛕","temples":["Kashi Vishwanath Varanasi — Jyotirlinga","Ram Mandir Ayodhya — birthplace of Lord Ram","Krishna Janmabhoomi Mathura","Banke Bihari Temple Vrindavan","Vindhyavasini Mirzapur","Hanuman Garhi Ayodhya"],"rivers":["Ganga — Moksha river","Yamuna","Saryu — Ayodhya's river","Gomti Lucknow","Betwa"],"places":["Varanasi — Kashi — world's oldest city","Ayodhya — Ram's birthplace","Vrindavan — Krishna's playground","Mathura — Krishna's birthplace","Prayagraj — Kumbh Mela","Agra — Taj Mahal","Lucknow — City of Nawabs","Sarnath — Buddha's first sermon"],"culture":["Ganga Aarti at Dashashwamedh Ghat — 1000 year tradition","Barsana Holi — Lathmar Holi","Lucknow Chikankari embroidery","Kathak classical dance birthplace","Awadhi cuisine — Biryani of Lucknow"],"infra":["Noida Electronic City","Purvanchal Expressway","Agra-Lucknow Expressway","Ganga Expressway","New Ayodhya Airport"],"fact":"Varanasi (Kashi) is the world's oldest living city — over 5,000 years. The new Ram Mandir covers 2.7 acres with 392 pillars."},
        {"id":"uk","name":"Uttarakhand","capital":"Dehradun","coords":[79.0,30.3],"zoom":7.5,"pitch":62,"bearing":15,"tagline":"Devbhoomi — Where Gods Reside","emoji":"🙏","temples":["Kedarnath Jyotirlinga — 3,583m altitude","Badrinath Vishnu — Char Dham","Gangotri — source of Ganga","Yamunotri — source of Yamuna","Tungnath — world's highest Shiva temple"],"rivers":["Ganga — origin at Gangotri","Yamuna — origin at Yamunotri","Alaknanda","Bhagirathi","Mandakini"],"places":["Rishikesh — World Yoga Capital","Haridwar — Gateway to Gods","Mussoorie — Queen of Hills","Nainital — Lake District","Valley of Flowers UNESCO","Auli skiing destination"],"culture":["Char Dham Yatra — 30 lakh pilgrims annually","Kumbh Mela Haridwar","Yoga and Ayurveda — ancient healing","Garhwali Jagar folk music"],"infra":["Char Dham All-Weather Road — 900 km","Rishikesh-Karnaprayag Railway","Tehri Dam — 260m — India's tallest dam"],"fact":"The entire Char Dham circuit covers 1,000 km through the Himalayas. Gangotri Glacier is retreating 22 metres annually."},
        {"id":"dl","name":"Delhi","capital":"New Delhi","coords":[77.1,28.6],"zoom":10.5,"pitch":50,"bearing":0,"tagline":"Capital of a Billion Dreams","emoji":"🏛️","temples":["Akshardham — world's largest Hindu temple complex","Lotus Temple — Bahai House of Worship","Birla Mandir","Chattarpur Mandir complex","ISKCON Temple"],"rivers":["Yamuna — Delhi's lifeline"],"places":["Red Fort — Lal Qila UNESCO","India Gate — War Memorial","Qutub Minar UNESCO","Humayun's Tomb UNESCO","Parliament House","Chandni Chowk","Connaught Place","Rajpath (Kartavya Path)","Rashtrapati Bhavan"],"culture":["Republic Day Parade — Jan 26","Delhi Street Food paradise","Dilli Haat craft bazaar","Sufi music at Nizamuddin Dargah"],"infra":["Delhi Metro — 4th largest in world","IGI Airport — busiest in South Asia","Delhi-Mumbai Expressway","Smart City Mission Delhi"],"fact":"Akshardham temple spans 100 acres with 20,000 carved figurines. Delhi has been the capital of 7 major empires over 1,000 years."},
        {"id":"rj","name":"Rajasthan","capital":"Jaipur","coords":[74.2,27.0],"zoom":6.8,"pitch":50,"bearing":-20,"tagline":"Land of Maharajas — Where Forts Touch the Sky","emoji":"🏰","temples":["Brahma Temple Pushkar — only Brahma temple in India","Dilwara Jain Temples — white marble masterpiece","Eklingji Temple","Karni Mata — temple of 25,000 rats","Ranakpur Jain Temple — 1,444 pillars"],"rivers":["Chambal","Luni","Banas"],"places":["Jaipur Pink City — Hawa Mahal","Jodhpur Blue City — Mehrangarh Fort","Udaipur — City of Lakes","Jaisalmer Golden Fort","Pushkar Holy Lake","Ranthambore Tiger Reserve","Chittorgarh Fort"],"culture":["Ghoomar dance","Kathputli puppet theatre","Camel Safari in Thar Desert","Bandhani tie-dye","Pushkar Camel Fair — November"],"infra":["Jaipur Metro","New Jaipur Airport","Solar energy parks — Rajasthan leads India"],"fact":"Jaisalmer Fort is one of the world's few living forts — 5,000 people still live inside its 800-year-old walls. At night, the yellow sandstone glows like gold."},
        {"id":"mp","name":"Madhya Pradesh","capital":"Bhopal","coords":[78.6,22.9],"zoom":7.0,"pitch":48,"bearing":10,"tagline":"Heart of India — Land of Tigers and Temples","emoji":"🐯","temples":["Mahakaleshwar Ujjain — most powerful Jyotirlinga","Omkareshwar Jyotirlinga","Khajuraho temples UNESCO","Sanchi Stupa","Amarkantak — source of Narmada River"],"rivers":["Narmada — India's holiest river after Ganga","Chambal","Son","Betwa","Tapti"],"places":["Ujjain — ancient city of Mahakal","Khajuraho — temple art capital","Sanchi UNESCO site","Bandhavgarh Tiger Reserve","Kanha — inspired Jungle Book","Indore — cleanest city in India"],"culture":["Simhastha Kumbh Mela Ujjain","Tribal Bhil and Gond art","Chanderi and Maheshwari silk sarees","Tansen Music Festival Gwalior"],"infra":["Indore — India's cleanest city 7 years","Bhopal Metro","IIT Indore"],"fact":"Mahakaleshwar is the only Jyotirlinga that faces South — symbolising power over death. The Narmada river flows west while all other sacred rivers flow east."},
        {"id":"gj","name":"Gujarat","capital":"Gandhinagar","coords":[71.5,22.2],"zoom":7.0,"pitch":50,"bearing":-10,"tagline":"Land of Gandhi — Cradle of Industry and Devotion","emoji":"🦁","temples":["Somnath — First Jyotirlinga rebuilt 7 times","Dwarkadhish Temple — Krishna's ocean kingdom","Akshardham Gandhinagar","Ambaji Shakti Pith"],"rivers":["Sabarmati — Gandhi's river","Narmada — Sardar Sarovar Dam","Mahi","Tapi"],"places":["Statue of Unity — 182m world's tallest","Gir Forest — last Asiatic Lions","Rann of Kutch — white salt desert","Ahmedabad — first UNESCO World Heritage City of India","Somnath — ancient seaport","Dwarka — sunken city of Krishna","Lothal — 4500-year Indus Valley port"],"culture":["Navratri Garba — 9 nights of divine dance","Patola double-ikat silk weaving","Kite Festival Makar Sankranti","Kutchi embroidery"],"infra":["GIFT City — India's first smart financial city","Surat Diamond Bourse — world's largest office building","Mundra Port — India's largest private port","Dholera Smart City"],"fact":"Dwaraka (Lord Krishna's city) lies submerged 40 metres underwater in the Arabian Sea — ancient ruins found in 1983. The Rann of Kutch under a full moon turns completely silver-white."},
        {"id":"mh","name":"Maharashtra","capital":"Mumbai","coords":[75.7,19.7],"zoom":7.0,"pitch":52,"bearing":5,"tagline":"Gateway of India — Dreams and Heritage","emoji":"🏙️","temples":["Trimbakeshwar Jyotirlinga Nashik","Siddhivinayak Temple Mumbai","Shirdi Sai Baba Temple","Pandharpur Vitthal","Tuljapur Bhavani Shakti Pith"],"rivers":["Godavari","Krishna","Bhima","Tapi","Ulhas"],"places":["Mumbai — Maximum City — financial capital","Marine Drive — Queen's Necklace at night","Ajanta Caves UNESCO","Ellora Caves UNESCO","Gateway of India","Pune — Oxford of the East","Nashik — Grape and wine capital","Konkan Coast"],"culture":["Ganesh Chaturthi — world's largest festival","Bollywood Film City","Lavani dance","Warli tribal art","Kolhapuri chappal craft"],"infra":["Mumbai Metro expansion","Mumbai Trans-Harbour Link — India's longest sea bridge 21.8 km","Navi Mumbai Airport","Atal Setu","IIT Bombay"],"fact":"The Mumbai Trans-Harbour Link (2024) is India's longest sea bridge at 21.8 km. Ajanta's murals painted with minerals still glow after 2,000 years."},
        {"id":"ka","name":"Karnataka","capital":"Bengaluru","coords":[75.7,15.3],"zoom":7.0,"pitch":50,"bearing":-15,"tagline":"Silicon Valley of India — Ancient Empires to Modern Tech","emoji":"🔱","temples":["Virupaksha Temple Hampi — 7th century still active","Udupi Sri Krishna Math","Chamundeshwari Temple Mysuru","Kukke Subramanya","Dharmasthala","Gokarna Mahabaleshwara"],"rivers":["Kaveri (Cauvery)","Tungabhadra","Krishna","Sharavathi","Kabini"],"places":["Bengaluru — Asia's fastest growing tech hub","Mysuru Palace — India's most visited monument","Hampi UNESCO","Coorg — coffee paradise","Jog Falls — 253m India's tallest waterfall","Gokarna sacred beach"],"culture":["Mysuru Dasara — 10 days of royal grandeur","Yakshagana — night-long theatre","Carnatic classical music","Mysuru silk sarees"],"infra":["Bengaluru — 1,500+ tech companies","Bengaluru Metro","ISRO Headquarters — India's space programme","IISc — India's top research university"],"fact":"Bengaluru produces more software engineers than Silicon Valley. ISRO's Chandrayaan-3 mission (2023) — first soft landing on Moon's south pole — was designed here."},
        {"id":"kl","name":"Kerala","capital":"Thiruvananthapuram","coords":[76.3,10.8],"zoom":7.5,"pitch":55,"bearing":10,"tagline":"God's Own Country — Where Nature Meets Divinity","emoji":"🌴","temples":["Padmanabhaswamy Temple — world's wealthiest","Sabarimala Ayyappa — 5 crore pilgrims annually","Guruvayur Krishna","Vadakkunnathan Thrissur"],"rivers":["Periyar","Bharathapuzha","Kabani","Pampa — Sabarimala river"],"places":["Alleppey Backwaters — Venice of the East","Munnar tea gardens","Kovalam Beach","Periyar Wildlife Sanctuary","Wayanad forests","Kochi — Queen of the Arabian Sea"],"culture":["Kathakali — divine dance drama","Onam harvest festival","Ayurveda — 5,000 year healing","Snake boat race Nehru Trophy","Kalaripayattu — world's oldest martial art"],"infra":["Kochi Metro — India's first water metro","Kerala's 100% literacy","Vizhinjam International Seaport"],"fact":"Padmanabhaswamy temple holds treasure estimated at Rs 2 lakh crore — confirmed treasures include 800-kg gold coconut shells. Kerala tops India in literacy (96%)."},
        {"id":"tn","name":"Tamil Nadu","capital":"Chennai","coords":[78.6,11.1],"zoom":7.0,"pitch":50,"bearing":-5,"tagline":"Dravidian Glory — Temple State of the World","emoji":"🛕","temples":["Meenakshi Amman Madurai — 33,000 sculptures","Brihadeeswarar Thanjavur UNESCO","Nataraja Temple Chidambaram","Ramanathaswamy Rameshwaram — Char Dham","Murugan temples — 6 Arupadaiveedu"],"rivers":["Kaveri — Ponni — life of Tamil Nadu","Vaigai","Tamiraparani — India's cleanest river"],"places":["Chennai — Gateway to South India","Madurai — temple city","Thanjavur — Chola art capital","Mahabalipuram UNESCO","Ooty Nilgiris","Kanyakumari — land's end — three oceans meet","Rameswaram — Adam's Bridge"],"culture":["Bharatanatyam — mother of all classical dance","Carnatic music","Pongal harvest festival","Silk sarees of Kanchipuram — 1,200 years of weaving"],"infra":["Chennai Port — India's second largest","TIDEL Park — IT hub","IIT Madras — India's No 1 ranked university"],"fact":"Brihadeeswarar temple (1010 CE) cast zero shadow at noon. Tamil is the world's oldest surviving classical language — 2,000+ years."},
        {"id":"ap","name":"Andhra Pradesh","capital":"Amaravati","coords":[79.7,15.9],"zoom":7.2,"pitch":50,"bearing":15,"tagline":"Rice Granary — Where Balaji Blesses the World","emoji":"🌾","temples":["Tirumala Venkateswara — world's most visited religious site","Srikalahasti — vayu linga","Kanaka Durga Vijayawada","Amaravati Stupa"],"rivers":["Krishna","Godavari","Tungabhadra","Pennar"],"places":["Tirupati — sacred hill city","Visakhapatnam — Vizag","Vijayawada — business hub","Araku Valley — tribal coffee","Belum Caves — India's longest natural cave"],"culture":["Kuchipudi classical dance — born here","Kalamkari pen paintings","Tirupati Ladoo — world's most consumed prasadam"],"infra":["Amaravati greenfield capital city","Visakhapatnam steel plant","AP Fibernet — world's cheapest internet"],"fact":"Lord Venkateswara at Tirumala receives donations of Rs 700+ crore annually and 50,000 devotees daily. 3 lakh ladoos distributed daily."},
        {"id":"ts","name":"Telangana","capital":"Hyderabad","coords":[79.0,18.1],"zoom":7.5,"pitch":52,"bearing":-10,"tagline":"City of Pearls — Where Nizams Met Silicon Valley","emoji":"💻","temples":["Yadagirigutta Lakshmi Narasimha — newly renovated golden temple","Birla Mandir Hyderabad","Thousand Pillar Temple Warangal — 12th century Kakatiya"],"rivers":["Godavari","Krishna","Musi — runs through Hyderabad"],"places":["HITEC City — Hyderabad IT Hub","Charminar — 1591 CE icon","Golconda Fort — acoustic marvel","Hussain Sagar Lake","Ramoji Film City — world's largest"],"culture":["Hyderabadi Biryani — world's most famous rice dish","Nawabi Hyderabadi culture","Bidri metal craft","Golconda diamonds — Kohinoor came from here"],"infra":["HITEC City — Google Apple Microsoft Amazon offices","Hyderabad Metro — 69 km","Rajiv Gandhi International Airport","Genome Valley — Asia's biotech hub"],"fact":"The Kohinoor diamond came from Golconda mines. Ramoji Film City is in Guinness World Records as the world's largest film studio covering 1,666 acres."},
        {"id":"wb","name":"West Bengal","capital":"Kolkata","coords":[87.8,22.9],"zoom":7.5,"pitch":48,"bearing":5,"tagline":"Cultural Soul of India — Where Art Lives in Streets","emoji":"🎭","temples":["Dakshineswar Kali — on Hooghly river","Kalighat Temple — ancient Shakti Pith","Belur Math — Ramakrishna's ashram","Tarakeswar Shiva Temple"],"rivers":["Ganga/Hooghly","Damodar","Teesta — jewel of Darjeeling"],"places":["Kolkata — City of Joy — cultural capital","Darjeeling — tea gardens at 7,000 ft","Sundarbans mangroves — UNESCO","Victoria Memorial","Howrah Bridge","Bishnupur — terracotta temples"],"culture":["Durga Puja — world's largest open-air art gallery","Rabindra Sangeet — Tagore's musical universe","Bengali literature — Tagore Nobel Prize 1913","Baul music — mystic tradition"],"infra":["Kolkata Metro — India's first metro (1984)","Vidyasagar Setu — India's longest cable bridge","New Town Rajarhat IT hub"],"fact":"Howrah Bridge carries 100,000 vehicles and 150,000 pedestrians daily — with no nuts or bolts — only rivets."},
        {"id":"or","name":"Odisha","capital":"Bhubaneswar","coords":[85.0,20.9],"zoom":7.2,"pitch":50,"bearing":-10,"tagline":"Temple City State — Cosmic Art in Stone","emoji":"⛩️","temples":["Jagannath Puri — Char Dham","Konark Sun Temple UNESCO — stone chariot","Lingaraj Temple Bhubaneswar — 11th century","Mukteswar Temple"],"rivers":["Mahanadi","Brahmani","Baitarani"],"places":["Puri — sacred beach city","Konark — Sun chariot on the sea","Chilika Lake — Asia's largest coastal lagoon","Bhubaneswar — temple city with 700 ancient temples"],"culture":["Rath Yatra Puri — 8 crore devotees — world's largest chariot festival","Odissi classical dance","Pattachitra scroll paintings","Sand art — Sudarsan Pattnaik world record"],"infra":["Paradip Port","Bhubaneswar Smart City","IIT Bhubaneswar","AIIMS Bhubaneswar"],"fact":"The Jagannath temple flag defies physics — it always flies opposite to the wind. 'Juggernaut' in English derives from 'Jagannath'."},
        {"id":"as","name":"Assam","capital":"Dispur","coords":[92.9,26.2],"zoom":7.5,"pitch":48,"bearing":10,"tagline":"Land of the Mighty Brahmaputra — One-Horned Rhino Kingdom","emoji":"🦏","temples":["Kamakhya Temple — most powerful Shakti Pith","Umananda Temple — on peacock island in Brahmaputra","Navagraha Temple","Hajo — multi-faith pilgrimage site"],"rivers":["Brahmaputra — mighty river — 18 km wide in places","Barak","Subansiri","Manas"],"places":["Kaziranga — UNESCO — 2700 one-horned rhinos","Manas National Park UNESCO","Majuli — world's largest river island — 880 sq km","Sibsagar — Ahom kingdom — 600 year reign"],"culture":["Bihu — spring harvest dance — 3 types","Assam Muga silk — golden — rarest in world","Tea gardens — first planted 1839 — 800 gardens"],"infra":["Bogibeel Bridge — India's longest combined road-rail bridge 4.94 km","IIT Guwahati"],"fact":"Kamakhya is India's most powerful tantric temple. The Muga silk of Assam is naturally golden and cannot be dyed — the only golden silk in the world."},
        {"id":"ar","name":"Arunachal Pradesh","capital":"Itanagar","coords":[94.7,28.2],"zoom":7.0,"pitch":60,"bearing":-20,"tagline":"Land of the Rising Sun — India's Eastern Himalayan Frontier","emoji":"🌄","temples":["Tawang Monastery — 400 year old — 2nd largest Buddhist in world","Parasuram Kund","Golden Pagoda Namsai"],"rivers":["Brahmaputra (Siang)","Lohit","Subansiri","Kameng"],"places":["Tawang at 10,500 ft altitude","Ziro Valley — ancient Apatani tribe","Namdapha — biodiversity hotspot","Bomdila — Himalayan gateway","Dong — India's easternmost village"],"culture":["26 major tribes — each with distinct language","Nyokum Yullo tribal harvest festival","Monpa Tsechu dance festival"],"infra":["Sela Tunnel — 13,000 ft — strategic connectivity","Itanagar Greenfield Airport"],"fact":"Dong village in Arunachal Pradesh sees the first sunrise in India every day — 1 hour 58 minutes before Mumbai. Arunachal has more biodiversity per sq km than the Amazon."},
        {"id":"mn","name":"Manipur","capital":"Imphal","coords":[93.9,24.6],"zoom":8.0,"pitch":52,"bearing":15,"tagline":"Jewel of the East — Cradle of Polo and Classical Dance","emoji":"💎","temples":["Govindaji Temple — Manipuri Vaishnavism","Kangla Fort — sacred royal seat"],"rivers":["Imphal River","Barak","Iril"],"places":["Loktak Lake — floating islands (phumdis)","Keibul Lamjao — floating national park","Moreh — trade gateway to Myanmar","Ima Keithel — world's only all-women market"],"culture":["Manipuri classical dance — Ras Lila of Krishna","Pung Cholom — drum dance","Sangai Festival — November","Polo — invented here 3,000 years ago"],"infra":["Jiribam-Imphal Railway — world's highest railway bridge 141m","Imphal International Airport upgrade"],"fact":"Manipur invented the game of Polo — the Mapal Kangjeibung ground in Imphal (1697 AD) is the world's oldest polo ground. The Ima Keithel market is run exclusively by mothers."},
        {"id":"ml","name":"Meghalaya","capital":"Shillong","coords":[91.4,25.5],"zoom":8.0,"pitch":55,"bearing":-5,"tagline":"Abode of Clouds — Where Roots Build Bridges","emoji":"☁️","temples":["Nartiang Durga Temple","Kyllang Rock — sacred tribal site"],"rivers":["Umngot — crystal clear riverbed visible 20m deep","Simsang","Kopili"],"places":["Cherrapunji — wettest place on Earth (11,872mm annually)","Dawki — emerald river boat rides","Living Root Bridges — Nongriat village","Mawlynnong — cleanest village in Asia","Double Decker Root Bridge — 500 years old"],"culture":["Nongkrem harvest festival","Khasi matrilineal society — women inherit property","Wangala drum festival — Garo tribe"],"infra":["Shillong Tech Park","Umiam Lake hydroelectric"],"fact":"The Living Root Bridges are bioengineered — Khasi tribe trains rubber tree roots over bamboo scaffolding for 10-15 years. The oldest bridges are 500 years old and grow stronger every year."},
        {"id":"sk","name":"Sikkim","capital":"Gangtok","coords":[88.5,27.5],"zoom":8.5,"pitch":60,"bearing":20,"tagline":"Organic Himalayan State — Guardian of Kangchenjunga","emoji":"🌺","temples":["Rumtek Monastery — Kagyu Buddhism","Pemayangtse Monastery — 1700 AD","Tashiding Monastery — most sacred","Enchey Monastery"],"rivers":["Teesta — glacial river","Rangit","Rangpo"],"places":["Gangtok — clean Himalayan capital","Nathula Pass — 14,140 ft — China border open to tourists","Yumthang Valley — Valley of Flowers","Gurudongmar Lake — 17,100 ft — sacred to Sikhs and Buddhists"],"culture":["Losar Tibetan New Year","Saga Dawa festival","Momos — steamed dumplings"],"infra":["100% organic farming — first in India","Gangtok ropeway","Teesta hydropower cascade"],"fact":"Kangchenjunga (8,586m) — world's 3rd highest peak — is so sacred that Sikkim's climbers leave the last few metres unclimbed out of respect. Sikkim achieved 100% organic status in 2016."},
        {"id":"bi","name":"Bihar","capital":"Patna","coords":[85.3,25.7],"zoom":7.5,"pitch":48,"bearing":-5,"tagline":"Cradle of Civilization — Where the Buddha Found Enlightenment","emoji":"🕌","temples":["Mahabodhi Temple Bodh Gaya UNESCO — where Siddhartha became Buddha","Vishnupad Temple Gaya — Lord Vishnu's footprint in stone","Mundeshwari Temple — 1,700 years old — oldest functioning temple in India","Patna Sahib Gurudwara — birthplace of Guru Gobind Singh"],"rivers":["Ganga — sacred confluence at Patna","Gandak","Kosi","Son"],"places":["Bodh Gaya — where Buddha attained enlightenment","Nalanda — world's first residential university 5th century BC","Rajgir — where Buddha taught","Vaishali — world's first democratic republic","Patna Sahib — Sikh holy city"],"culture":["Chhath Puja — world's only sun worship ritual standing in water","Madhubani painting — UNESCO intangible heritage","Sonepur Mela — world's largest cattle fair"],"infra":["Patna Metro under construction","Nalanda University revival — 2014","Vikramsila Setu — bridge on Ganga"],"fact":"Nalanda University (500 BC - 1200 AD) housed 10,000 students from 18 countries — 800 years before Oxford. Vaishali held the world's first democratic election in 600 BC."},
        {"id":"jh","name":"Jharkhand","capital":"Ranchi","coords":[85.3,23.6],"zoom":7.5,"pitch":50,"bearing":10,"tagline":"Land of Forests — Mineral Heartland of India","emoji":"⛏️","temples":["Baidyanath Jyotirlinga Deoghar — most visited Jyotirlinga","Pahari Mandir Ranchi — hilltop Shiva","Rajrappa Temple — Chhinnamasta Devi — Shakti Pith"],"rivers":["Damodar","Subarnarekha — gold particles found","Koel","Sankh"],"places":["Deoghar — Baidyanath temple city","Ranchi — Jharkhand capital","Netarhat — Scotland of Jharkhand","Betla National Park","Hundru Falls 98m"],"culture":["Sarhul — flower festival","Jhumar folk dance","Santhali tribal culture — one of India's oldest"],"infra":["Tata Steel Jamshedpur — India's Pittsburgh","Bokaro Steel City","Deoghar airport"],"fact":"Jamshedpur (1907) is India's first planned industrial city — built by JRD Tata and still India's most liveable planned city. The Subarnarekha river carries alluvial gold."},
        {"id":"cg","name":"Chhattisgarh","capital":"Raipur","coords":[81.8,21.2],"zoom":7.2,"pitch":48,"bearing":-10,"tagline":"Rice Bowl of India — Land of Waterfalls and Tribes","emoji":"🌿","temples":["Bhoramdeo Temple — Khajuraho of Chhattisgarh","Ratanpur Mahamaya Devi","Rajim temples","Dongargarh Bambleshwari"],"rivers":["Mahanadi — born in Sihawa hills here","Indravati","Sheonath"],"places":["Chitrakote Falls — 300m wide — India's Niagara","Bastar — tribal heartland","Jagdalpur — Bastar's capital","Achanakmar Tiger Reserve"],"culture":["Bastar Dussehra — 75 days — world's longest festival","Gond tribal art","Dhokra metal casting (lost-wax technique)"],"infra":["NMDC iron ore — Bailadila — world's largest iron ore deposit","BSP Bhilai Steel Plant — India's largest integrated steel plant","AIIMS Raipur"],"fact":"Bhilai Steel Plant produces 7 million tonnes of steel annually. Bastar Dussehra involves 600 tribal villages and has been celebrated continuously for 600+ years."},
        {"id":"ga","name":"Goa","capital":"Panaji","coords":[74.1,15.3],"zoom":9.5,"pitch":50,"bearing":5,"tagline":"Pearl of the Orient — Where East Meets West","emoji":"🏖️","temples":["Shri Mangueshi Temple — largest in Goa","Shantadurga Temple","Tambdi Surla Mahadeva — 12th century jungle temple","Basilica of Bom Jesus UNESCO — St Francis Xavier"],"rivers":["Mandovi","Zuari","Chapora"],"places":["Baga-Calangute-Anjuna North Goa beaches","Old Goa UNESCO churches","Dudhsagar Falls 310m","Palolem Beach South Goa","Panaji Latin Quarter","Chapora Fort","Spice farms inland Goa"],"culture":["Carnival — Portuguese legacy parade","Goan fish curry-rice — way of life","Shigmo festival — spring colours"],"infra":["Mopa Airport — new international airport 2022","National Institute of Oceanography Panaji","Goa Shipyard Limited"],"fact":"The Basilica of Bom Jesus (1605) contains the preserved body of St. Francis Xavier — miraculously incorrupt for 450 years. Goa has India's highest per capita income."},
        {"id":"pu","name":"Puducherry","capital":"Puducherry","coords":[79.8,11.9],"zoom":10.0,"pitch":45,"bearing":0,"tagline":"French Riviera of the East — Where East and West Live as One","emoji":"🇫🇷","temples":["Manakula Vinayagar Temple — elephant blesses pilgrims","Vedapureeswarar Temple","Immaculate Conception Cathedral — 18th century"],"rivers":["Gingee River","Ponnaiyar"],"places":["French Quarter — yellow mustard walls unchanged for 300 years","Auroville — 50 nation township","Promenade Beach — 3 km seaside boulevard","Sri Aurobindo Ashram","Paradise Beach"],"culture":["Tamil-French fusion culture unique on Earth","Auroville — 3,000 people from 50 nations — no religion no politics","Creole cuisine"],"infra":["JIPMER — one of India's finest hospitals","NIT Puducherry","Auroville renewable energy experiments"],"fact":"Auroville's Matrimandir is a golden sphere 29m diameter containing a 70cm crystal ball focus for meditation — it generates its own solar electricity."},
        {"id":"an","name":"Andaman & Nicobar","capital":"Port Blair","coords":[92.7,11.7],"zoom":7.0,"pitch":50,"bearing":10,"tagline":"Emerald Islands — Where India Meets the Deep Ocean","emoji":"🌊","temples":[],"rivers":[],"places":["Cellular Jail — National Memorial","Radhanagar Beach Havelock — one of Asia's 10 best beaches","Baratang Limestone Caves","Ross Island","Barren Island — India's only active volcano","Jolly Buoy coral reefs"],"culture":["Great Andamanese tribe — 60,000 years of continuous habitation","Jarawa reserve — untouched civilization","Nicobarese canoe culture"],"infra":["Port Blair Airport upgrade","INS Baaz air base Nicobar — strategic","Undersea internet cable connection"],"fact":"The Sentinelese tribe of North Sentinel Island has lived in complete isolation for 60,000 years. The Andamans are India's strategic gateway to the Strait of Malacca — through which 80% of global oil passes."},
        {"id":"ld","name":"Lakshadweep","capital":"Kavaratti","coords":[72.6,10.6],"zoom":8.0,"pitch":50,"bearing":-5,"tagline":"India's Coral Paradise — Arabian Sea Jewels","emoji":"🐠","temples":[],"rivers":[],"places":["Bangaram Island — no cars no alcohol no stress","Agatti coral atoll — lagoon landing airstrip","Kavaratti","Minicoy Island — lighthouse 1885","Marine National Park — 105 coral species"],"culture":["Lakshadweep Muslim Malayali culture","Lava and Kolkali dance","Traditional fishing with tuna pole-and-line"],"infra":["Agatti Airport","Kavaratti desalination plant","Solar power grid for islands","INS Dwarka naval base"],"fact":"Lakshadweep's coral reefs are home to 600 species of fish and 105 coral species. The lagoon water has 30m visibility — among the clearest on Earth."},
    ]

    _SPECIALS = [
        {"id":"himalayas","name":"The Great Himalayas","coords":[80.0,31.5],"zoom":6.5,"pitch":65,"bearing":10,"type":"special","emoji":"🏔️","color":"#c4b5fd","tagline":"Roof of the World — Where Earth Touches Heaven","desc":"The Himalayas span 2,400 km across northern India — home to 14 peaks above 8,000m. Kailash Mansarovar (6,638m) is believed to be the abode of Lord Shiva — the centre of the universe in Hindu cosmology. Every year 50,000 pilgrims undertake the Kailash Parikrama despite its extreme altitude.","highlights":["Mount Everest — 8,849m — world's highest","Kangchenjunga — 8,586m — India's highest","Kailash Mansarovar — Shiva's abode","Siachen Glacier — world's highest battlefield","Gangotri Glacier — Ganga's source","Valley of Flowers — UNESCO biosphere"]},
        {"id":"char_dham","name":"Char Dham — Sacred Circuit","coords":[79.5,30.5],"zoom":7.5,"pitch":62,"bearing":20,"type":"special","emoji":"🙏","color":"#fde68a","tagline":"Four Sacred Abodes — The Ultimate Pilgrimage","desc":"Char Dham — four of Hinduism's holiest sites — set by Adi Shankaracharya in 8th century AD to unite India spiritually. Every Hindu aspires to complete this circuit before death. Over 30 lakh pilgrims undertake the journey annually through the Himalayas.","highlights":["Badrinath (Vishnu) — 3,133m — closes Nov","Kedarnath (Shiva) — 3,583m — triangular stone lingam","Gangotri (Ganga) — source of sacred river","Yamunotri (Yamuna) — smallest yet holiest","Char Dham Mahamarg — 900 km all-weather highway","Kedarnath rebuilt after 2013 flash flood"]},
        {"id":"ayodhya","name":"Ayodhya — Ram's Sacred City","coords":[82.2,26.8],"zoom":12.0,"pitch":55,"bearing":-10,"type":"special","emoji":"🛕","color":"#f97316","tagline":"Birthplace of Lord Ram — Eternal Treta Yuga City","desc":"Ayodhya — the birthplace of Lord Ram — stands transformed. The Ram Mandir inaugurated on January 22, 2024, is built exactly at the sacred birthplace (Ram Janmabhoomi). The temple spans 2.7 acres with pink Rajasthani sandstone, 392 carved pillars, and enshrines a 5-year-old Ram Lalla idol. 55 lakh devotees visited in the first month alone.","highlights":["Ram Mandir — completed Jan 2024","Ram Lalla idol — black stone — exquisite","Saryu River Ghat — evening aarti","Hanuman Garhi — 52 steps fort temple","Kanak Bhawan — golden Ram Sita","New Ayodhya airport","Ram Path — 4-lane devotional corridor"]},
        {"id":"varanasi","name":"Varanasi — Kashi — Eternal City","coords":[83.0,25.3],"zoom":13.0,"pitch":55,"bearing":-20,"type":"special","emoji":"🕉️","color":"#fb923c","tagline":"World's Oldest Living City — Where Shiva Grants Moksha","desc":"Varanasi (Kashi) has been continuously inhabited for 5,000+ years — making it the world's oldest city. Lord Shiva himself is said to reside here, whispering the Taraka mantra to every dying soul. The Ganga Aarti at Dashashwamedh Ghat — 84 ghats line the riverbank — is performed every evening for 1,000 years without pause.","highlights":["Kashi Vishwanath corridor — renovated 2021","Dashashwamedh Ghat — Ganga Aarti nightly","84 ghats along the Ganga","Manikarnika Ghat — funeral fire never extinguished","Sarnath — Buddha's first sermon","Banarasi silk weaving"]},
        {"id":"vrindavan","name":"Vrindavan-Mathura — Krishna's Land","coords":[77.7,27.6],"zoom":12.5,"pitch":50,"bearing":5,"type":"special","emoji":"🎵","color":"#ec4899","tagline":"Where Krishna Played — Eternal Leela Ground","desc":"Vrindavan and Mathura — where Lord Krishna was born, grew up, played the flute on the banks of Yamuna, and lifted Govardhan Hill. Over 5,000 temples exist in a 20km radius. Holi here (Barsana Lathmar Holi) is the most joyful festival on Earth — celebrated 40 days before the rest of India.","highlights":["Krishna Janmabhoomi Mathura","Banke Bihari Temple Vrindavan","Radha Raman Temple","Govardhan Hill — Krishna lifted it","Barsana — Radha's village","Vrindavan Prem Mandir — glowing at night"]},
        {"id":"ujjain","name":"Ujjain — Mahakal's Kingdom","coords":[75.8,23.2],"zoom":12.5,"pitch":55,"bearing":15,"type":"special","emoji":"⚡","color":"#8b5cf6","tagline":"City of Mahakal — Master of Time and Death","desc":"Ujjain is one of India's seven sacred cities (Sapta Puri) and home to Mahakaleshwar — the most powerful of the 12 Jyotirlingas. Unlike all others, Mahakaleshwar faces South — the direction of Yama (death) — granting power over time itself. The Simhastha Kumbh Mela here (every 12 years) draws the world's largest human gathering.","highlights":["Mahakaleshwar Jyotirlinga — South-facing — most powerful","Bhasma Aarti at 4 AM — with real funeral ash","Simhastha Kumbh Mela — 2028 next","Shipra River — most sacred river of Malwa","Kal Bhairav Temple — liquor offering","Ram Ghat evening aarti"]},
        {"id":"indian_ocean","name":"Indian Ocean — India's Waters","coords":[76.0,8.0],"zoom":5.5,"pitch":50,"bearing":0,"type":"special","emoji":"🌊","color":"#06b6d4","tagline":"India Rules Three Seas — Arabian, Indian Ocean, Bay of Bengal","desc":"India is embraced by three water bodies: the Arabian Sea (west), Bay of Bengal (east), and Indian Ocean (south) — a 7,516 km coastline. Kanyakumari is the only place in India where you can see the sun rise over the Bay of Bengal and set into the Arabian Sea. The Indian Ocean is named after India — testifying to India's historical maritime supremacy.","highlights":["Kanyakumari — tip of India — three seas meet","Lakshadweep coral islands","Andaman — 572 islands","Dhanushkodi — Adam's Bridge to Sri Lanka","Vizhinjam deep-water port Kerala","INS Vikrant — India's first indigenous aircraft carrier"]},
        {"id":"defense","name":"Defenders of Bharat","coords":[77.1,28.6],"zoom":7.0,"pitch":45,"bearing":0,"type":"special","emoji":"🇮🇳","color":"#22c55e","tagline":"Bharat Mata Ki Jai — Guardians of 1.4 Billion Lives","desc":"India's Armed Forces — 1.4 million active personnel — are among the world's most professional and battle-hardened. The Indian Army has the world's highest battlefield — Siachen Glacier at 22,000 ft. INS Vikrant (2022) is India's first indigenously built aircraft carrier. The Air Force operates Rafale, Tejas, and Su-30 MKI aircraft.","highlights":["Indian Army — 1.2 million strong — Siachen to Andaman","Indian Navy — INS Vikrant aircraft carrier — 2022","Indian Air Force — Rafale, Tejas — Make in India","DRDO — Agni V missile — nuclear deterrent","HAL Tejas — India's indigenous light combat aircraft","Param Vir Chakra — India's highest gallantry"]},
        {"id":"jai_hind","name":"Jai Hind — Incredible India","coords":[82.8,22.0],"zoom":4.8,"pitch":50,"bearing":0,"type":"special","emoji":"🇮🇳","color":"#f59e0b","tagline":"Satyameva Jayate — Truth Alone Triumphs","desc":"From the Himalayas to Kanyakumari, from the Arabian Sea to Arunachal's dawn — Incredible India encompasses 5,000 years of unbroken civilization. One nation, 1.4 billion hearts, 22 languages, 29 states, countless traditions — yet one Bharat. Vande Mataram, Jai Hind!","highlights":["World's oldest living civilization","1.4 billion people — most populous nation","22 official languages — hundreds of dialects","4 major world religions born here","25% of world's software engineers","Space programme — Moon's south pole 2023","G20 Presidency 2023 — Vasudhaiva Kutumbakam","5th largest economy — growing to No 3 by 2030"]},
    ]

    _sj  = _json.dumps(_STATES,   ensure_ascii=False)
    _spj = _json.dumps(_SPECIALS, ensure_ascii=False)
    _total = len(_STATES) + len(_SPECIALS)

    _html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Incredible India — Sky Tour</title>
<link rel="stylesheet" href="https://unpkg.com/maplibre-gl@3.6.2/dist/maplibre-gl.css">
<script src="https://unpkg.com/maplibre-gl@3.6.2/dist/maplibre-gl.js"></script>
<link href="https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,700;0,800;0,900;1,700&family=Rajdhani:wght@400;500;600;700&family=DM+Sans:wght@300;400;500;600&display=swap" rel="stylesheet">
<style>
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
html,body{{width:100%;height:100%;overflow:hidden;background:#000;font-family:'DM Sans',sans-serif}}
#loader{{position:fixed;inset:0;z-index:500;display:flex;flex-direction:column;align-items:center;justify-content:center;background:radial-gradient(ellipse 100% 80% at 50% 50%,#0a0520 0%,#020010 60%,#000 100%);}}
#loader-om{{font-size:clamp(50px,8vw,90px);animation:omSpin 2s ease-out both,omGlow 1.5s ease-in-out 1.5s infinite alternate}}
@keyframes omSpin{{from{{transform:rotate(-360deg) scale(0);opacity:0}}to{{transform:rotate(0) scale(1);opacity:1}}}}
@keyframes omGlow{{from{{text-shadow:0 0 20px #f59e0b}}to{{text-shadow:0 0 50px #f59e0b,0 0 100px #fbbf24}}}}
#loader-text{{font-family:'Playfair Display',serif;font-size:clamp(18px,3vw,30px);font-weight:800;color:#f59e0b;margin-top:20px;letter-spacing:2px;text-align:center}}
#loader-sub{{font-family:'Rajdhani',sans-serif;font-size:clamp(11px,1.5vw,14px);color:rgba(255,200,100,.6);letter-spacing:4px;text-transform:uppercase;margin-top:8px}}
#lbar-wrap{{width:clamp(200px,30vw,320px);height:3px;background:rgba(255,255,255,.1);border-radius:3px;margin-top:28px}}
#lbar{{width:0%;height:100%;background:linear-gradient(90deg,#f59e0b,#ffd700);border-radius:3px;transition:width .2s ease}}
#stars-cv{{position:fixed;inset:0;z-index:98;pointer-events:none}}
#space{{position:fixed;inset:0;z-index:100;background:radial-gradient(ellipse 90% 70% at 50% 42%,#080520 0%,#030015 55%,#000 100%);display:none;flex-direction:column;align-items:center;justify-content:center;gap:0;}}
.earth-wrap{{position:relative;flex-shrink:0;margin-bottom:28px}}
.earth{{width:clamp(170px,26vw,320px);height:clamp(170px,26vw,320px);border-radius:50%;background:radial-gradient(circle at 25% 25%, rgba(60,180,255,.35) 0%,transparent 45%),radial-gradient(circle at 75% 65%, rgba(0,160,80,.2) 0%,transparent 38%),conic-gradient(from 200deg at 50% 50%,#060d1f 0deg,#0d1a35 30deg,#0a2a1a 65deg,#1a0a05 100deg,#0a1a2a 140deg,#040e20 180deg,#0d1a30 220deg,#080c1a 260deg,#0a180e 300deg,#060d1f 360deg);box-shadow:0 0 80px rgba(50,150,255,.45),0 0 160px rgba(30,100,255,.2),inset -25px -20px 70px rgba(0,0,10,.75);animation:earthRotate 25s linear infinite,earthBob 8s ease-in-out infinite;position:relative;}}
@keyframes earthRotate{{0%{{filter:hue-rotate(0deg)}}100%{{filter:hue-rotate(8deg)}}}}
@keyframes earthBob{{0%,100%{{transform:translateY(0)}}50%{{transform:translateY(-15px)}}}}
.city-lights{{position:absolute;inset:0;border-radius:50%;background:radial-gradient(circle at 42% 37%, rgba(255,230,100,.55) 0%,transparent 7%),radial-gradient(circle at 52% 35%, rgba(255,200,80,.45) 0%,transparent 5%),radial-gradient(circle at 47% 33%, rgba(255,180,60,.4) 0%,transparent 4%),radial-gradient(circle at 44% 41%, rgba(255,220,100,.35) 0%,transparent 3%);mix-blend-mode:screen;animation:cityPulse 3s ease-in-out infinite alternate;}}
@keyframes cityPulse{{from{{opacity:.7}}to{{opacity:1}}}}
.india-glow{{position:absolute;top:30%;left:37%;width:26%;height:34%;background:radial-gradient(ellipse,rgba(255,140,0,.9) 0%,rgba(255,80,0,.4) 40%,transparent 70%);border-radius:50%;animation:indiaGlow 2s ease-in-out infinite alternate;}}
@keyframes indiaGlow{{from{{opacity:.6;transform:scale(.88)}}to{{opacity:1;transform:scale(1.12)}}}}
.earth-atmo{{position:absolute;inset:-12px;border-radius:50%;border:1px solid rgba(80,160,255,.25);box-shadow:0 0 30px rgba(80,160,255,.15),inset 0 0 30px rgba(80,160,255,.05);}}
.space-title{{font-family:'Playfair Display',serif;font-weight:900;font-size:clamp(24px,5vw,58px);color:#fff;text-align:center;letter-spacing:-1px;text-shadow:0 0 30px rgba(255,140,0,.9),0 0 80px rgba(255,100,0,.5);animation:titleFlare 3.5s ease-in-out infinite;}}
@keyframes titleFlare{{0%,100%{{text-shadow:0 0 30px rgba(255,140,0,.9),0 0 80px rgba(255,100,0,.5)}}50%{{text-shadow:0 0 50px rgba(255,200,0,1),0 0 130px rgba(255,140,0,.7),0 0 200px rgba(255,80,0,.3)}}}}
.space-sub{{font-family:'Rajdhani',sans-serif;letter-spacing:4px;text-transform:uppercase;font-size:clamp(10px,1.4vw,15px);color:rgba(255,200,100,.65);margin-top:10px;text-align:center;}}
#start-btn{{margin-top:34px;padding:clamp(13px,2vw,19px) clamp(32px,5vw,64px);background:linear-gradient(135deg,#ff8c00,#ff4500);border:none;border-radius:999px;cursor:pointer;font-family:'Rajdhani',sans-serif;font-weight:700;font-size:clamp(13px,1.8vw,18px);letter-spacing:3px;text-transform:uppercase;color:#fff;box-shadow:0 0 60px rgba(255,140,0,.75),0 10px 40px rgba(0,0,0,.6);animation:btnPulse 2.5s ease-in-out infinite;transition:transform .25s ease,box-shadow .25s ease;}}
#start-btn:hover{{transform:scale(1.08) translateY(-4px);box-shadow:0 0 90px rgba(255,180,0,1),0 12px 48px rgba(0,0,0,.6)}}
@keyframes btnPulse{{0%,100%{{box-shadow:0 0 60px rgba(255,140,0,.75),0 10px 40px rgba(0,0,0,.6)}}50%{{box-shadow:0 0 100px rgba(255,200,0,1),0 10px 48px rgba(0,0,0,.6)}}}}
#map{{position:fixed;inset:0;z-index:1;background:#000a0a}}
#hud{{position:fixed;inset:0;z-index:10;pointer-events:none;display:none}}
#topbar{{position:absolute;top:0;left:0;right:0;display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:8px;padding:12px 20px;background:linear-gradient(to bottom,rgba(0,0,0,.88) 0%,transparent 100%);pointer-events:all;}}
.brand{{display:flex;align-items:center;gap:10px}}
.brand-flag{{font-size:20px}}
.brand-text{{font-family:'Playfair Display',serif;font-size:clamp(13px,1.8vw,20px);font-weight:800;color:#f59e0b;text-shadow:0 0 16px rgba(255,140,0,.5)}}
.brand-text small{{display:block;font-family:'Rajdhani',sans-serif;font-size:.58em;letter-spacing:3px;text-transform:uppercase;color:rgba(255,255,255,.45)}}
.hbtns{{display:flex;gap:7px;align-items:center;flex-wrap:wrap}}
.hbtn{{padding:6px 13px;border-radius:999px;border:1px solid rgba(255,255,255,.18);background:rgba(0,0,0,.62);backdrop-filter:blur(14px);color:#fff;font-size:11px;font-weight:700;cursor:pointer;font-family:'Rajdhani',sans-serif;letter-spacing:1.5px;text-transform:uppercase;transition:all .22s ease;}}
.hbtn:hover,.hbtn.on{{background:rgba(255,140,0,.28);border-color:#ff8c00;color:#ff8c00}}
.hbtn.red{{border-color:rgba(255,80,80,.4);color:#ff9090}}
.hbtn.red:hover{{background:rgba(255,60,60,.28);border-color:#ff5050}}
#stsel{{padding:6px 14px;border-radius:10px;border:1.5px solid rgba(255,140,0,.4);background:rgba(0,0,0,.72);backdrop-filter:blur(14px);color:#fff;font-family:'Rajdhani',sans-serif;font-size:12px;cursor:pointer;outline:none;min-width:170px;}}
#stsel option{{background:#08081a;color:#fff}}
#phase{{position:absolute;left:50%;top:50%;transform:translate(-50%,-50%);text-align:center;z-index:15;pointer-events:none;opacity:0;transition:opacity .7s ease;}}
#phase.on{{opacity:1}}
#phase h2{{font-family:'Playfair Display',serif;font-size:clamp(20px,4vw,52px);font-weight:900;color:#fff;text-shadow:0 0 35px rgba(255,140,0,.85);letter-spacing:-.5px;line-height:1.1}}
#phase p{{font-family:'Rajdhani',sans-serif;font-size:clamp(11px,1.3vw,15px);color:rgba(255,200,100,.85);letter-spacing:3px;text-transform:uppercase;margin-top:9px}}
#card{{position:absolute;left:18px;bottom:68px;width:clamp(268px,28vw,370px);background:rgba(2,3,15,.92);border:1px solid rgba(255,140,0,.3);border-radius:20px;padding:18px 20px;backdrop-filter:blur(28px);pointer-events:all;transform:translateX(-115%);transition:transform .75s cubic-bezier(.22,1,.36,1);max-height:70vh;overflow-y:auto;scrollbar-width:thin;scrollbar-color:rgba(255,140,0,.35) transparent;}}
#card.open{{transform:translateX(0)}}
#card::-webkit-scrollbar{{width:3px}}
#card::-webkit-scrollbar-thumb{{background:rgba(255,140,0,.35);border-radius:3px}}
.c-emoji{{font-size:30px;margin-bottom:6px}}
.c-name{{font-family:'Playfair Display',serif;font-size:clamp(16px,2vw,22px);font-weight:900;color:#fff;letter-spacing:-.3px}}
.c-cap{{font-size:10px;color:rgba(255,200,100,.5);letter-spacing:2px;text-transform:uppercase;margin:2px 0 5px;font-family:'Rajdhani',sans-serif}}
.c-tagline{{font-family:'Playfair Display',serif;font-style:italic;font-size:13px;color:#ffd700;border-left:2px solid #ff8c00;padding-left:9px;margin-bottom:14px;line-height:1.4}}
.sec{{margin-bottom:11px}}
.sec-t{{font-size:10px;letter-spacing:2px;text-transform:uppercase;color:rgba(255,140,0,.75);font-weight:700;margin-bottom:5px;font-family:'Rajdhani',sans-serif}}
.tags{{display:flex;flex-wrap:wrap;gap:4px}}
.tag{{padding:3px 9px;border-radius:999px;font-size:11px;background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.12);color:rgba(255,255,255,.82)}}
.tag.temple{{background:rgba(255,140,0,.1);border-color:rgba(255,140,0,.25);color:#fed7aa}}
.tag.river{{background:rgba(30,144,255,.1);border-color:rgba(30,144,255,.25);color:#93c5fd}}
.tag.place{{background:rgba(34,197,94,.08);border-color:rgba(34,197,94,.2);color:#86efac}}
.tag.culture{{background:rgba(236,72,153,.08);border-color:rgba(236,72,153,.2);color:#f9a8d4}}
.tag.infra{{background:rgba(139,92,246,.1);border-color:rgba(139,92,246,.25);color:#c4b5fd}}
.c-fact{{font-size:12px;color:rgba(255,230,180,.72);line-height:1.68;border-top:1px solid rgba(255,255,255,.08);padding-top:9px;margin-top:7px;font-family:'Rajdhani',sans-serif}}
#card.special{{border-color:rgba(255,200,0,.4);background:rgba(5,2,18,.95)}}
.c-highlight-item{{display:flex;align-items:flex-start;gap:7px;margin-bottom:5px;font-size:12px;color:rgba(255,230,180,.8);font-family:'Rajdhani',sans-serif;line-height:1.5}}
.c-hi-dot{{width:5px;height:5px;border-radius:50%;background:#f59e0b;flex-shrink:0;margin-top:5px}}
#ctr{{position:absolute;right:18px;bottom:68px;background:rgba(2,3,15,.9);border:1px solid rgba(255,140,0,.22);border-radius:14px;padding:11px 15px;text-align:center;pointer-events:none;min-width:82px;backdrop-filter:blur(20px);}}
#ctr-n{{font-family:'Playfair Display',serif;font-size:26px;font-weight:900;color:#f59e0b;text-shadow:0 0 14px rgba(255,140,0,.6)}}
#ctr-t{{font-size:10px;color:rgba(255,255,255,.3);letter-spacing:1px;font-family:'Rajdhani',sans-serif}}
#ctr-s{{font-size:11px;color:rgba(255,200,100,.55);margin-top:4px;font-family:'Rajdhani',sans-serif;max-width:80px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
#legend{{position:absolute;top:70px;right:18px;background:rgba(2,3,15,.85);border:1px solid rgba(255,255,255,.1);border-radius:12px;padding:8px 12px;font-family:'Rajdhani',sans-serif;font-size:11px;color:rgba(255,255,255,.5);pointer-events:none;letter-spacing:.5px;transition:opacity .5s;}}
#progwrap{{position:absolute;bottom:0;left:0;right:0;height:3px;background:rgba(255,255,255,.07)}}
#prog{{height:100%;background:linear-gradient(90deg,#ff8c00,#ffd700);width:0%;transition:width .5s ease;border-radius:0 2px 2px 0}}
#flash{{position:fixed;inset:0;z-index:200;background:#ffb347;opacity:0;pointer-events:none;transition:opacity .12s ease}}
#flash.on{{opacity:.18}}
.lb{{position:fixed;left:0;right:0;z-index:5;background:#000;pointer-events:none;height:0;transition:height .9s ease}}
#lb-t{{top:0}}#lb-b{{bottom:0}}
body.cinema .lb{{height:5.5vh}}
</style>
</head>
<body>
<div id="loader">
  <div id="loader-om">🕉️</div>
  <div id="loader-text">INCREDIBLE INDIA</div>
  <div id="loader-sub">Loading the Virtual Sky Tour…</div>
  <div id="lbar-wrap"><div id="lbar"></div></div>
</div>
<canvas id="stars-cv"></canvas>
<div id="space">
  <div class="earth-wrap">
    <div class="earth">
      <div class="city-lights"></div>
      <div class="india-glow"></div>
    </div>
    <div class="earth-atmo"></div>
  </div>
  <div class="space-title">INCREDIBLE INDIA</div>
  <div class="space-sub">Virtual Sky Tour · {_total} Destinations · Space to Earth</div>
  <button id="start-btn" onclick="startTour()">🚀 Begin the Journey</button>
</div>
<div id="flash"></div>
<div id="lb-t" class="lb"></div>
<div id="lb-b" class="lb"></div>
<div id="map"></div>
<div id="hud">
  <div id="topbar">
    <div class="brand">
      <span class="brand-flag">🇮🇳</span>
      <div class="brand-text">INCREDIBLE INDIA <small>VIRTUAL SKY TOUR</small></div>
    </div>
    <div class="hbtns">
      <select id="stsel" onchange="T.jumpTo(parseInt(this.value))">
        <option value="-1">Jump to destination…</option>
      </select>
      <button class="hbtn" id="play-btn" onclick="T.togglePlay()">⏸ PAUSE</button>
      <button class="hbtn on" id="cinema-btn" onclick="T.toggleCinema()">🎬 CINEMA</button>
      <button class="hbtn red" onclick="T.exit()">✕ EXIT</button>
    </div>
  </div>
  <div id="phase"><h2 id="phase-h"></h2><p id="phase-p"></p></div>
  <div id="card">
    <div class="c-emoji" id="c-emoji"></div>
    <div class="c-name" id="c-name"></div>
    <div class="c-cap" id="c-cap"></div>
    <div class="c-tagline" id="c-tagline"></div>
    <div class="c-body" id="c-body"></div>
    <div class="c-fact" id="c-fact"></div>
  </div>
  <div id="ctr"><div id="ctr-n">0</div><div id="ctr-t">DESTINATION</div><div id="ctr-s"></div></div>
  <div id="legend">🟠 Temple &nbsp; 🔵 River &nbsp; 🟢 Place &nbsp; 🟣 Infra</div>
  <div id="progwrap"><div id="prog"></div></div>
</div>
<script>
var STATES   = {_sj};
var SPECIALS = {_spj};

// Build playlist — interleave specials
var PLAYLIST = [];
var sp_map = {{}};
sp_map[0]  = SPECIALS[0];
sp_map[1]  = SPECIALS[1];
sp_map[5]  = SPECIALS[2];
sp_map[9]  = SPECIALS[5];
sp_map[10] = SPECIALS[6];
sp_map[14] = SPECIALS[7];
sp_map[STATES.length-1] = SPECIALS[8];

// UP gets 3 specials (Ayodhya, Varanasi, Vrindavan)
var extra_after_up = [SPECIALS[3], SPECIALS[4]];

for(var i=0;i<STATES.length;i++){{
  PLAYLIST.push(STATES[i]);
  if(sp_map[i]) PLAYLIST.push(sp_map[i]);
  if(i===5) {{ PLAYLIST.push(extra_after_up[0]); PLAYLIST.push(extra_after_up[1]); }}
}}

var map = new maplibregl.Map({{
  container:'map',
  style:'https://basemaps.cartocdn.com/gl/dark-matter-gl-style/style.json',
  center:[82.8,22],zoom:4.5,pitch:0,bearing:0,
  interactive:true,attributionControl:false
}});

// Stars canvas
(function(){{
  var cv=document.getElementById('stars-cv'),ctx=cv.getContext('2d');
  function resize(){{cv.width=window.innerWidth;cv.height=window.innerHeight;}}
  resize();window.addEventListener('resize',resize);
  var stars=[];
  for(var i=0;i<320;i++) stars.push({{x:Math.random(),y:Math.random(),r:Math.random()*1.5+.3,o:Math.random()*.7+.3,s:Math.random()*.015+.005}});
  function draw(){{
    ctx.clearRect(0,0,cv.width,cv.height);
    stars.forEach(function(s){{
      s.o+=s.s*(Math.random()>.5?1:-1);
      s.o=Math.max(.1,Math.min(1,s.o));
      ctx.beginPath();ctx.arc(s.x*cv.width,s.y*cv.height,s.r,0,Math.PI*2);
      ctx.fillStyle='rgba(255,255,255,'+s.o+')';ctx.fill();
    }});
    requestAnimationFrame(draw);
  }}
  draw();
}})();

// Loader
var _lp=0;
var _lt=setInterval(function(){{
  _lp=Math.min(_lp+Math.random()*18+8,100);
  document.getElementById('lbar').style.width=_lp+'%';
  if(_lp>=100){{
    clearInterval(_lt);
    setTimeout(function(){{
      document.getElementById('loader').style.opacity='0';
      document.getElementById('loader').style.transition='opacity .5s';
      setTimeout(function(){{
        document.getElementById('loader').style.display='none';
        document.getElementById('space').style.display='flex';
        document.getElementById('stars-cv').style.display='block';
        populateSelect();
      }},500);
    }},300);
  }}
}},120);

function populateSelect(){{
  var sel=document.getElementById('stsel');
  PLAYLIST.forEach(function(s,i){{
    var o=document.createElement('option');
    o.value=i;o.textContent=(i+1)+'. '+s.emoji+' '+s.name;
    sel.appendChild(o);
  }});
}}

function startTour(){{
  document.getElementById('space').style.opacity='0';
  document.getElementById('space').style.transition='opacity .8s';
  setTimeout(function(){{
    document.getElementById('space').style.display='none';
    document.getElementById('stars-cv').style.display='none';
    document.getElementById('hud').style.display='block';
    document.body.classList.add('cinema');
    T.start();
  }},800);
}}

function showPhase(h,p){{
  var el=document.getElementById('phase');
  document.getElementById('phase-h').textContent=h;
  document.getElementById('phase-p').textContent=p;
  el.classList.add('on');
}}
function hidePhase(){{document.getElementById('phase').classList.remove('on');}}
function doFlash(){{
  var f=document.getElementById('flash');
  f.classList.add('on');
  setTimeout(function(){{f.classList.remove('on');}},180);
}}

var T={{
  idx:0, running:false, timer:null, paused:false,
  cinema:true, DWELL:7800,

  start:function(){{
    this.running=true;this.paused=false;
    showPhase('🇮🇳 INCREDIBLE INDIA','Space to Earth — Virtual Sky Tour');
    setTimeout(function(){{hidePhase();T.showStop(PLAYLIST[0]);}},2200);
  }},

  togglePlay:function(){{
    this.paused=!this.paused;
    document.getElementById('play-btn').textContent=this.paused?'▶ PLAY':'⏸ PAUSE';
    if(!this.paused) this.scheduleNext();
    else clearTimeout(this.timer);
  }},

  toggleCinema:function(){{
    this.cinema=!this.cinema;
    document.getElementById('cinema-btn').classList.toggle('on',this.cinema);
    document.body.classList.toggle('cinema',this.cinema);
  }},

  exit:function(){{
    this.running=false;clearTimeout(this.timer);
    document.getElementById('hud').style.display='none';
    document.getElementById('space').style.display='flex';
    document.getElementById('space').style.opacity='1';
    document.getElementById('stars-cv').style.display='block';
    document.body.classList.remove('cinema');
    map.flyTo({{center:[82.8,22],zoom:4.5,pitch:0,bearing:0,duration:1200}});
  }},

  jumpTo:function(i){{
    if(i<0||i>=PLAYLIST.length) return;
    clearTimeout(this.timer);
    document.getElementById('card').classList.remove('open');
    this.idx=i;
    setTimeout(function(){{T.showStop(PLAYLIST[T.idx]);}},300);
    document.getElementById('stsel').value=-1;
  }},

  showStop:function(s){{
    map.flyTo({{
      center:s.coords, zoom:s.zoom||7, pitch:s.pitch||50,
      bearing:s.bearing||0, duration:3200,
      easing:function(t){{return t<.5?2*t*t:(4-2*t)*t-1;}}
    }});

    var pct=Math.round((this.idx+1)/PLAYLIST.length*100);
    document.getElementById('prog').style.width=pct+'%';
    document.getElementById('ctr-n').textContent=this.idx+1;
    document.getElementById('ctr-t').textContent='/ '+PLAYLIST.length;
    document.getElementById('ctr-s').textContent=s.name;

    var card=document.getElementById('card');
    card.classList.remove('open','special');
    if(s.type==='special') card.classList.add('special');

    document.getElementById('c-emoji').textContent=s.emoji||'🇮🇳';
    document.getElementById('c-name').textContent=s.name;
    document.getElementById('c-cap').textContent=s.capital?'Capital: '+s.capital:'';
    document.getElementById('c-tagline').textContent=s.tagline||'';

    var body='';
    if(s.type==='special'){{
      if(s.desc) body+='<div class="sec"><div class="sec-t">About</div><div class="c-fact" style="border:none;padding:0;margin:0">'+s.desc+'</div></div>';
      if(s.highlights&&s.highlights.length){{
        body+='<div class="sec"><div class="sec-t">Highlights</div>';
        s.highlights.forEach(function(h){{body+='<div class="c-highlight-item"><div class="c-hi-dot"></div>'+h+'</div>';}});
        body+='</div>';
      }}
    }} else {{
      if(s.temples&&s.temples.length){{body+='<div class="sec"><div class="sec-t">⛩ Temples & Sacred Sites</div><div class="tags">';s.temples.forEach(function(x){{body+='<span class="tag temple">'+x+'</span>';}});body+='</div></div>';}}
      if(s.rivers&&s.rivers.length){{body+='<div class="sec"><div class="sec-t">🌊 Sacred Rivers</div><div class="tags">';s.rivers.forEach(function(x){{body+='<span class="tag river">'+x+'</span>';}});body+='</div></div>';}}
      if(s.places&&s.places.length){{body+='<div class="sec"><div class="sec-t">📍 Must-See Places</div><div class="tags">';s.places.forEach(function(x){{body+='<span class="tag place">'+x+'</span>';}});body+='</div></div>';}}
      if(s.culture&&s.culture.length){{body+='<div class="sec"><div class="sec-t">🎭 Culture & Festivals</div><div class="tags">';s.culture.forEach(function(x){{body+='<span class="tag culture">'+x+'</span>';}});body+='</div></div>';}}
      if(s.infra&&s.infra.length){{body+='<div class="sec"><div class="sec-t">🏗 Infrastructure</div><div class="tags">';s.infra.forEach(function(x){{body+='<span class="tag infra">'+x+'</span>';}});body+='</div></div>';}}
    }}
    document.getElementById('c-body').innerHTML=body;
    document.getElementById('c-fact').textContent=s.fact||'';

    setTimeout(function(){{card.classList.add('open');}},400);
    this.scheduleNext();
  }},

  scheduleNext:function(){{
    clearTimeout(this.timer);
    if(this.paused) return;
    var self=this;
    this.timer=setTimeout(function(){{
      var next=self.idx+1;
      if(next>=PLAYLIST.length){{
        showPhase('🇮🇳 JAI HIND!','Satyameva Jayate — Truth Alone Triumphs');
        setTimeout(function(){{hidePhase();self.idx=0;self.showStop(PLAYLIST[0]);}},3000);
        return;
      }}
      document.getElementById('card').classList.remove('open');
      setTimeout(function(){{
        doFlash();
        setTimeout(function(){{self.idx=next;self.showStop(PLAYLIST[self.idx]);}},300);
      }},500);
    }}, this.DWELL);
  }}
}};

document.addEventListener('keydown',function(e){{
  if(e.key===' '&&T.running){{e.preventDefault();T.togglePlay();}}
  if(e.key==='Escape'&&T.running) T.exit();
  if(e.key==='ArrowRight'&&T.running){{
    clearTimeout(T.timer);document.getElementById('card').classList.remove('open');
    setTimeout(function(){{T.idx=Math.min(T.idx+1,PLAYLIST.length-1);T.showStop(PLAYLIST[T.idx]);}},350);
  }}
  if(e.key==='ArrowLeft'&&T.running){{
    clearTimeout(T.timer);document.getElementById('card').classList.remove('open');
    setTimeout(function(){{T.idx=Math.max(T.idx-1,0);T.showStop(PLAYLIST[T.idx]);}},350);
  }}
}});
</script>
</body>
</html>"""

    _components.html(_html, height=860, scrolling=False)

    st.markdown("""
<div style="background:#000;color:rgba(255,140,0,.35);text-align:center;
  padding:4px 0;font-size:11px;font-family:monospace;letter-spacing:1.5px">
  🇮🇳 INCREDIBLE INDIA &nbsp;|&nbsp;
  SPACE=PAUSE &nbsp;↓&nbsp; →=NEXT &nbsp; ←=PREV &nbsp; ESC=EXIT
</div>
""", unsafe_allow_html=True)


# ═════════════════════════════════════════════════════════════════════════════
#  ROUTER  — decide which page to render
# ═════════════════════════════════════════════════════════════════════════════
def main():
    page = st.session_state.page
    if page == "infochat":
        render_infochat()
    elif page == "imagine":
        render_imagine()
    elif page == "study_planner":
        render_study_planner()
    elif page == "dr_nexa":
        render_dr_nexa()
    elif page == "yours_nexa":
        render_yours_nexa()
    elif page == "india_tour":
        render_india_tour()
    else:
        render_home()

main()
